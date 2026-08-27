#!/usr/bin/env python3
"""Build a printable vocabulary flashcard deck: Hebrew, or Greek volume 1 or 2.

The sheet follows the English tutoring deck already in use: A4 landscape, eight
cards to a page, front sheet then back sheet, duplex.  Two things differ.

No cutting lines are printed.  A hairline that prints a millimetre off shows up
on the finished card as a crooked edge, so instead the grid is even and the
cover sheet gives the two measurements a guillotine needs.

The back sheet mirrors the column order (4-3-2-1 instead of 1-2-3-4).  Printed
duplex, that is what puts each meaning behind its own word; without it every
card would carry a neighbour's translation.

Pictures are optional by design: a card with no honest picture prints without
one rather than borrowing an approximate image.

The Hebrew master carries a part of speech per word.  The Greek master carries
none, so it is worked out in ``flashcard_pos`` and left blank where the citation
form is ambiguous — a blank line costs nothing, a wrong label is learned as fact.
"""

from __future__ import annotations

import argparse
import json
from functools import lru_cache
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))

from flashcard_pos import greek_part_of_speech
from greek_citation_form import card_headword  # noqa: E402

# The Latin deck reads the identity and the part of speech the reader's own
# builders already worked out, rather than deriving either a second time here.
import latin_source_texts as _latin_text  # noqa: E402
from build_latin_full_reader import short_pos as latin_pos  # noqa: E402


def latin_key(entry: dict) -> str:
    return _latin_text.fold(entry.get("forms") or entry["headword"])

ROOT = Path(__file__).resolve().parents[1]
CACHE = ROOT / "output/source-cache"
IMAGE_DIR = CACHE / "flashcards/openmoji-618"
# 第二層圖庫（game-icons／Phosphor／MDI／Tabler，見 iconify_card_images.py）：
# 只給「跟不相干的詞共用同一張 emoji」的卡用，補 OpenMoji 沒有的概念。
ICON_DIR = CACHE / "flashcards/iconify"
OUTPUT_DIR = ROOT / "output/flashcards"

PAGE_W_MM = 297.0
PAGE_H_MM = 210.0
COLS = 4
ROWS = 2
CARD_W_MM = PAGE_W_MM / COLS      # 74.25
# Both figures are measured, not derived.  The renderer reserves more vertical
# space than the declared margins account for, and it moves the second row to a
# page of its own well before the arithmetic says the rows should stop fitting;
# a 10 mm margin fails where 5 mm holds, and 97 mm rows fail where 94 mm hold.
CARD_H_MM = 94.0
MARGIN_V_MM = 5.0

FONT_ZH = "MingLiU"
FONT_UI = "MingLiU"
INK = "1B1B1B"
MUTED = "8A8A8A"

POS_PT = 12
LESSON_PT = 10
IMAGE_MM = 32
# 卡框：比照使用者桌面那副《家教單字卡.pdf》—— 量出來線寬 2.12 mm、緋紅
# #ED0A3F、正反面都有框。框往內縮 FRAME_INSET_MM，裁歪一兩毫米只會讓白邊不
# 等寬，不會把框裁掉；框畫在裁切線上就會缺一邊、還帶進隔壁卡的框。
# 課次配色十色輪：第 11 課回到第 1 課的紅。使用者指定的順序是
# 紅橙黃綠藍紫棕粉深灰深綠，色值挑印得出來、彼此分得開的濃度。
FRAME_COLORS = (
    "ED0A3F",  # 紅
    "F07C1E",  # 橙
    "E8B10A",  # 黃
    "3AA655",  # 綠
    "1E6FD9",  # 藍
    "7B3FA0",  # 紫
    "8B5A2B",  # 棕
    "F080B4",  # 粉
    "4A4A4A",  # 深灰
    "1F6B4A",  # 深綠
)
FRAME_RADIUS_MM = 3.75      # 圓角半徑，量參考卡量出來的
FRAME_INSET_MM = 3.0
FRAME_EIGHTHS = 48          # w:sz 以八分之一點計，48 ＝ 6 pt ＝ 2.1 mm
# 框線是往外畫的：宣告的寬高是框內緣，成品外緣還要加上兩條框線。
FRAME_LINE_MM = FRAME_EIGHTHS / 8 / 72 * 25.4
# 卡框連同格線邊距要留得比卡片矮一點：擠滿的話 EXACTLY 的列高會被撐開，
# 第二排整排往下移，框就壓過裁切線。留 2 mm 也讓垂直置中重新生效。
FRAME_SLACK_MM = 2.0
FRAME_PAD_DXA = 110         # 框內再留 2 mm，字不要貼著框
# 卡寬 74.25 mm 扣掉兩側格線邊距，再留一點安全量；超過這個寬度 LibreOffice 會把
# 字頭的最後一個字母折到第二行，把下面的課次擠掉。
HEADWORD_MAX_MM = 54

# The glosses run from two characters to twenty-seven, so the meaning line is
# sized to fit the card rather than set at one size and allowed to overflow.
ZH_STEPS = ((6, 24), (10, 20), (16, 16), (22, 13), (99, 11))

POS_ZH = {
    "noun": "名詞",
    "verb": "動詞",
    "adjective": "形容詞",
    "adverb": "副詞",
    "pronoun": "代名詞",
    "preposition": "介系詞",
    "conjunction": "連接詞",
    "particle": "質詞",
    "proper_name": "專名",
    "particle_or_preposition": "質詞／介系詞",
    "interrogative_particle": "疑問質詞",
    "prepositional_phrase": "介系詞片語",
    "adverbial_phrase": "副詞片語",
    "conjunction_phrase": "連接詞片語",
}

# The Latin master already labels its own parts of speech, one character each;
# the card prints the two-character form the other decks use.
POS_LATIN = {
    "名": "名詞", "動": "動詞", "形": "形容詞", "副": "副詞", "介": "介系詞",
    "連": "連接詞", "代": "代名詞", "數": "數詞", "嘆": "感嘆詞", "不變": "不變詞",
}

def hebrew_part_of_speech(entry: dict, grammar: dict) -> str:
    """名詞印到性別，重複的字形再印數與狀態：名詞‧陰性‧複數。

    性別、數與狀態都讀 OSHB 標註（`scripts/hebrew_card_grammar.py` 產的檔），查
    無標註就只印詞性 —— 希伯來文的性別看不出字尾：אֶ֫רֶץ、עִיר、יָד 沒有陰性
    字尾卻是陰性，דֶּ֫רֶךְ、רוּחַ 兩性都用。標錯會被當事實背起來。
    """

    label = POS_ZH.get(entry["partOfSpeech"], entry["partOfSpeech"])
    parts = [label]
    if entry["partOfSpeech"] == "noun" and grammar.get("gender"):
        parts.append(grammar["gender"])
    for field in ("number", "state"):
        if grammar.get(field):
            parts.append(grammar[field])
    return "‧".join(parts)


DECKS: dict[str, dict] = {
    "hbo": {
        "title": "聖經希伯來文單字卡",
        "vocab": ROOT / "data/originalReaders/vocabulary/hebrew-1000.json",
        "glosses": CACHE / "original-readers/hebrew-full/hebrew-gloss-zh-reviewed-by-lemma.json",
        "images": CACHE / "flashcards/hebrew-card-images.json",
        "icons": CACHE / "flashcards/hebrew-card-icons.json",
        "grammar": CACHE / "flashcards/hebrew-card-grammar.json",
        "output": "hebrew-flashcards-1000.docx",
        "font": "Noto Serif Hebrew",
        "fontFile": "C:/Windows/Fonts/NotoSerifHebrew-Regular.ttf",
        "rtl": True,
        "headword_pt": 54,
        "sources": [
            "詞表：Pratico–Van Pelt《Basics of Biblical Hebrew》2/e 詞序，其後接語料頻率延伸。",
            "原文與出現次數：Westminster Leningrad Codex（OSHB）。",
            "人名、地名與民族國名不在本套卡內，另收於讀本的分類專名表。",
        ],
    },
    "grc1": {
        "title": "通用希臘文單字卡・上冊",
        "vocab": ROOT / "data/originalReaders/vocabulary/greek-2000.json",
        "glosses": CACHE / "original-readers/greek-full/greek-2000-gloss-zh-by-lemma.json",
        "images": CACHE / "flashcards/greek-card-images.json",
        "icons": CACHE / "flashcards/greek-card-icons.json",
        "output": "greek-flashcards-volume-1.docx",
        "font": "Palatino Linotype",
        "rtl": False,
        "headword_pt": 34,
        "volume": 1,
        "sources": [
            "詞表：Mounce《Basics of Biblical Greek》詞序，其後接新約與七十士譯本語料頻率延伸。",
            "正面為字典引用形，附詞尾與冠詞。",
        ],
    },
    "grc2": {
        "title": "通用希臘文單字卡・下冊",
        "vocab": ROOT / "data/originalReaders/vocabulary/greek-2000.json",
        "glosses": CACHE / "original-readers/greek-full/greek-2000-gloss-zh-by-lemma.json",
        "images": CACHE / "flashcards/greek-card-images.json",
        "icons": CACHE / "flashcards/greek-card-icons.json",
        "output": "greek-flashcards-volume-2.docx",
        "font": "Palatino Linotype",
        "rtl": False,
        "headword_pt": 34,
        "volume": 2,
        "sources": [
            "詞表：教父希臘文語料頻率延伸，與上冊不重複。",
            "正面為字典引用形，附詞尾與冠詞。",
        ],
    },
    "lat1": {
        "title": "教會拉丁文單字卡・上冊",
        "vocab": ROOT / "data/originalReaders/vocabulary/latin-2000.json",
        "images": CACHE / "flashcards/latin-card-images.json",
        "icons": CACHE / "flashcards/latin-card-icons.json",
        "output": "latin-flashcards-volume-1.docx",
        "font": "Noto Serif",
        "rtl": False,
        "headword_pt": 30,
        "volumeName": "上冊",
        "sources": [
            "詞表：Collins《A Primer of Ecclesiastical Latin》原書詞序。",
            "正面為字典引用形，動詞列四個主要部分，名詞列主格、屬格與性。",
            "長音符號依原書標示；發音為羅馬式教會發音。",
        ],
    },
    "lat2": {
        "title": "教會拉丁文單字卡・下冊",
        "vocab": ROOT / "data/originalReaders/vocabulary/latin-2000.json",
        "images": CACHE / "flashcards/latin-card-images.json",
        "icons": CACHE / "flashcards/latin-card-icons.json",
        "output": "latin-flashcards-volume-2.docx",
        "font": "Noto Serif",
        "rtl": False,
        "headword_pt": 30,
        "volumeName": "下冊",
        "sources": [
            "詞表：教父、中世紀與近現代教廷語料詞頻，與上冊不重複。",
            "詞形主要部分取自 Whitaker's WORDS。",
            "正面為字典引用形；發音為羅馬式教會發音。",
        ],
    },
}


def set_rfonts(run, font: str) -> None:
    """Pin every Word script slot so the headword never falls back to a UI font."""

    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{slot}"), font)
    for theme in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        fonts.attrib.pop(qn(f"w:{theme}"), None)


def write(paragraph, text: str, font: str, size: float, *, color: str = INK,
          bold: bool = False, rtl: bool = False):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    set_rfonts(run, font)
    # python-docx 只寫 w:sz，那管的是拉丁字。希伯來是複合語系，排版程式看的是
    # w:szCs；不一起設，整副卡的字頭都會用樣式預設的 11 pt 印出來 —— 宣告 54 pt
    # 也一樣，而且不會有任何錯誤訊息。
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(int(round(size * 2))))
    run._element.get_or_add_rPr().append(size_cs)
    if bold:
        bold_cs = OxmlElement("w:bCs")
        bold_cs.set(qn("w:val"), "1")
        run._element.get_or_add_rPr().append(bold_cs)
    if rtl:
        r_pr = run._element.get_or_add_rPr()
        mark = OxmlElement("w:rtl")
        mark.set(qn("w:val"), "1")
        r_pr.append(mark)
        p_pr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        p_pr.append(bidi)
    return run


def blank(cell, points: float) -> None:
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    write(paragraph, "", FONT_UI, points)


def strip_borders(table) -> None:
    """Print no rules at all; the cut is measured, not traced."""

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        borders.append(element)
    table._tbl.tblPr.append(borders)


def new_grid(document: Document):
    table = document.add_table(rows=ROWS, cols=COLS)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 外層格線也要鎖死寬度。裡面放了巢狀表格（卡框）之後，沒鎖的欄寬會被重算，
    # 八張卡整排縮成 61 mm 一張並往右擠出頁面。
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), str(int(Mm(CARD_W_MM * COLS).twips)))
    table_width.set(qn("w:type"), "dxa")
    table._tbl.tblPr.append(table_width)
    for grid_col in table._tbl.findall(qn("w:tblGrid") + "/" + qn("w:gridCol")):
        grid_col.set(qn("w:w"), str(int(Mm(CARD_W_MM).twips)))
    strip_borders(table)
    for row in table.rows:
        # One height declaration only.  Setting row.height and then appending a
        # second w:trHeight leaves two competing rules in the XML.
        row.height = Mm(CARD_H_MM)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            cell.width = Mm(CARD_W_MM)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 外層格子不留邊距：內縮由卡框自己的尺寸決定，框才置得中；靠邊距擠
            # 會把列撐高，第二排整排下移、框壓過裁切線。
            margins = OxmlElement("w:tcMar")
            for edge, value in (("top", 0), ("start", 0), ("bottom", 0), ("end", 0)):
                node = OxmlElement(f"w:{edge}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            cell._tc.get_or_add_tcPr().append(margins)
    return table


def clear(cell) -> None:
    first = cell.paragraphs[0]._p
    first.getparent().remove(first)


DRAWING_NS = {
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
}
EMU_PER_MM = 36000


def frame_shape(cell, color: str, left_mm: float, top_mm: float, shape_id: int) -> None:
    """在卡片格裡放一個圓角矩形當卡框，位置寫死在頁面座標上。

    表格框線沒有圓角，所以框得用圖形做。圖形錨在**頁面**而不是段落或格子：
    格內錨點的 y 會跟著內容高度浮動（內容是垂直置中的），八張卡的框就會各自
    高低不一。版面是固定格線，每張卡在頁面上的座標算得出來，直接寫死最穩。
    `layoutInCell="0"` 一定要有，否則位置又會被格子夾回去。
    """

    width = CARD_W_MM - 2 * FRAME_INSET_MM
    height = CARD_H_MM - 2 * FRAME_INSET_MM
    adjust = int(FRAME_RADIUS_MM / width * 100000)
    xml = f"""<w:r {" ".join(f'xmlns:{k}="{v}"' for k, v in DRAWING_NS.items())}
                   xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:drawing>
        <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
                   relativeHeight="{shape_id}" behindDoc="1" locked="0"
                   layoutInCell="0" allowOverlap="1">
          <wp:simplePos x="0" y="0"/>
          <wp:positionH relativeFrom="page">
            <wp:posOffset>{int(left_mm * EMU_PER_MM)}</wp:posOffset>
          </wp:positionH>
          <wp:positionV relativeFrom="page">
            <wp:posOffset>{int(top_mm * EMU_PER_MM)}</wp:posOffset>
          </wp:positionV>
          <wp:extent cx="{int(width * EMU_PER_MM)}" cy="{int(height * EMU_PER_MM)}"/>
          <wp:effectExtent l="0" t="0" r="0" b="0"/>
          <wp:wrapNone/>
          <wp:docPr id="{shape_id}" name="card-frame-{shape_id}"/>
          <wp:cNvGraphicFramePr/>
          <a:graphic>
            <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
              <wps:wsp>
                <wps:cNvSpPr/>
                <wps:spPr>
                  <a:xfrm>
                    <a:off x="0" y="0"/>
                    <a:ext cx="{int(width * EMU_PER_MM)}" cy="{int(height * EMU_PER_MM)}"/>
                  </a:xfrm>
                  <a:prstGeom prst="roundRect">
                    <a:avLst><a:gd name="adj" fmla="val {adjust}"/></a:avLst>
                  </a:prstGeom>
                  <a:noFill/>
                  <a:ln w="{int(FRAME_LINE_MM * EMU_PER_MM)}">
                    <a:solidFill><a:srgbClr val="{color}"/></a:solidFill>
                  </a:ln>
                </wps:spPr>
                <wps:bodyPr/>
              </wps:wsp>
            </a:graphicData>
          </a:graphic>
        </wp:anchor>
      </w:drawing>
    </w:r>"""
    cell.paragraphs[0]._p.append(parse_xml(xml))


def framed(cell, lesson: int, place: tuple[int, int, int]):
    """在卡片格內再放一張單格表當卡框，回傳要填內容的那一格。

    表格框線畫在格子邊上，所以框要用「內縮一圈的巢狀表格」做，不能直接給外層
    格子加邊框 —— 那條線正好落在裁切線上。
    """

    column, row, shape_id = place
    frame_shape(
        cell,
        FRAME_COLORS[(lesson - 1) % len(FRAME_COLORS)],
        column * CARD_W_MM + FRAME_INSET_MM,
        MARGIN_V_MM + row * CARD_H_MM + FRAME_INSET_MM,
        shape_id,
    )
    inner = cell.add_table(rows=1, cols=1)
    inner.autofit = False
    inner.alignment = WD_TABLE_ALIGNMENT.CENTER
    width_dxa = str(int(Mm(CARD_W_MM - 2 * (FRAME_INSET_MM + FRAME_LINE_MM)).twips))
    # 巢狀表格不給明確寬度就會照內容撐開，把框撐出格子外、還把整個版面推歪。
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    inner._tbl.tblPr.append(layout)
    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), width_dxa)
    table_width.set(qn("w:type"), "dxa")
    inner._tbl.tblPr.append(table_width)
    for grid_col in inner._tbl.findall(qn("w:tblGrid") + "/" + qn("w:gridCol")):
        grid_col.set(qn("w:w"), width_dxa)
    strip_borders(inner)

    row = inner.rows[0]
    row.height = Mm(CARD_H_MM - 2 * (FRAME_INSET_MM + FRAME_LINE_MM) - FRAME_SLACK_MM)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    target = inner.cell(0, 0)
    target.width = Mm(CARD_W_MM - 2 * (FRAME_INSET_MM + FRAME_LINE_MM))
    target.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    margins = OxmlElement("w:tcMar")
    for edge in ("top", "start", "bottom", "end"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(FRAME_PAD_DXA))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    target._tc.get_or_add_tcPr().append(margins)

    # Word 規定表格後面一定要有段落，那個段落會佔高度並把框頂出格子；壓到 1 pt。
    tail = cell.paragraphs[-1]
    tail.paragraph_format.space_before = Pt(0)
    tail.paragraph_format.space_after = Pt(0)
    tail.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    tail.paragraph_format.line_spacing = Pt(1)
    tail.add_run().font.size = Pt(1)
    clear(target)
    return target


def zh_size(gloss: str) -> float:
    for limit, size in ZH_STEPS:
        if len(gloss) <= limit:
            return size
    return ZH_STEPS[-1][1]


def consonants(text: str) -> int:
    """希伯來字頭有幾個字母。母音點、達格什、重音各自都是一個碼點，數不得。"""

    return sum(1 for ch in text if "א" <= ch <= "ת")


@lru_cache(maxsize=None)
def face(font_file: str):
    """量字寬用的字型；量不到就回 None，改走子音數階梯。"""

    try:
        from PIL import ImageFont  # 只有希伯來牌組要量字寬，其餘四副不必為它裝 Pillow

        return ImageFont.truetype(font_file, 100)
    except Exception:
        return None


def headword_size(deck: dict, text: str) -> float:
    """Greek citation forms are long (ἄγγελος, -ου, ὁ); shrink so they fit."""

    base = deck["headword_pt"]
    if deck.get("rtl"):
        # 希伯來字頭的寬度跟字母數不成比例：מִשְׁפָּחָה 五個字母比 מַלְכוּת 寬得多。
        # 所以直接量字寬，取「塞得進卡片」的字級 —— 量出來的值與 LibreOffice 排出
        # 來的差 0.1 mm 以內。量不到字型檔才退回字母數階梯。
        measured = face(deck.get("fontFile", ""))
        if measured is not None:
            em = measured.getlength(text) / 100
            if em > 0:
                fits = HEADWORD_MAX_MM / 25.4 * 72 / em
                return min(base, round(fits * 2) / 2)
        for limit, factor in ((5, 1.0), (7, 0.86), (9, 0.72), (11, 0.60), (99, 0.50)):
            if consonants(text) <= limit:
                return base * factor
    # Latin citation forms are longer than any other deck's -- four principal
    # parts run past forty characters -- so they get their own ladder and are
    # allowed to wrap rather than being shrunk to illegibility.
    if deck.get("volumeName"):
        for limit, factor in ((12, 1.0), (20, 0.80), (28, 0.66), (38, 0.55), (99, 0.46)):
            if len(text) <= limit:
                return base * factor
    if len(text) <= 8:
        return base
    if len(text) <= 14:
        return base * 0.78
    if len(text) <= 20:
        return base * 0.62
    return base * 0.52


def fill_front(cell, card: dict, deck: dict, place: tuple[int, int, int]) -> None:
    cell = framed(cell, card["lesson"], place)
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    write(paragraph, card["headword"], deck["font"],
          headword_size(deck, card["headword"]), rtl=deck["rtl"])
    blank(cell, 20)
    footer = cell.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    write(footer, f"第 {card['lesson']} 課", FONT_UI, LESSON_PT, color=MUTED)


def fill_back(cell, card: dict, picture: Path | None, place: tuple[int, int, int]) -> None:
    cell = framed(cell, card["lesson"], place)
    if picture:
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run().add_picture(str(picture), width=Mm(IMAGE_MM))
    else:
        blank(cell, 20)
    meaning = cell.add_paragraph()
    meaning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meaning.paragraph_format.space_after = Pt(3)
    write(meaning, card["glossZh"], FONT_ZH, zh_size(card["glossZh"]))
    if card["pos"]:
        pos = cell.add_paragraph()
        pos.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pos.paragraph_format.space_after = Pt(0)
        write(pos, card["pos"], FONT_UI, POS_PT, color=MUTED)
    blank(cell, 10)
    footer = cell.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    write(footer, f"第 {card['lesson']} 課", FONT_UI, LESSON_PT, color=MUTED)


def configure(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(PAGE_W_MM)
    section.page_height = Mm(PAGE_H_MM)
    for attribute in ("left_margin", "right_margin", "header_distance", "footer_distance", "gutter"):
        setattr(section, attribute, Mm(0))
    section.top_margin = Mm(MARGIN_V_MM)
    section.bottom_margin = Mm(MARGIN_V_MM)
    style = document.styles["Normal"]
    style.font.name = FONT_UI
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0


def page_break(document: Document) -> None:
    """A full-height paragraph here would push the bottom row onto its own page."""

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    spacer.paragraph_format.line_spacing = Pt(1)
    run = spacer.add_run()
    run.font.size = Pt(1)
    run.add_break(WD_BREAK.PAGE)


def add_cover(document: Document, deck: dict, total: int, sheets: int, with_picture: int) -> None:
    """One sheet of instructions, then a blank back so the duplex pairing holds."""

    cuts_h = "、".join(f"{MARGIN_V_MM + CARD_H_MM * index:.0f}" for index in range(3))
    cuts_v = "、".join(f"{CARD_W_MM * index:.2f}".rstrip("0").rstrip(".") for index in range(1, 4))
    lines = [
        (deck["title"], 30, INK, 14),
        (f"{total} 張・{sheets} 組雙面・每頁 8 張", 13, MUTED, 22),
        ("列印：A4 橫式，雙面列印選「沿長邊翻頁」，縮放設為 100%（不要選「符合頁面大小」）。", 11, INK, 6),
        (f"裁切：不印裁切線。自紙張上緣量，橫向切在 {cuts_h} mm；自左緣量，縱向切在 {cuts_v} mm。"
         f"成品每張 {CARD_W_MM:.2f}×{CARD_H_MM:.0f} mm。"
         f"卡框比裁切線內縮 {FRAME_INSET_MM:.0f} mm，裁歪一兩毫米只會讓白邊不等寬，不會切到框。", 11, INK, 6),
        ("框色按課次十色輪替：紅橙黃綠藍紫棕粉深灰深綠，第十一課回到紅色。同一課的卡同色，方便整理與抽考。",
         11, INK, 6),
        ("正面：原文與課次。背面：繁體中文詞義、詞性與課次。", 11, INK, 6),
        (f"插圖：{with_picture} 張有圖，其餘留白。多義詞取最常見的義項；找不到誠實對應的圖就不放，"
         "不以近似圖充數。", 11, INK, 22),
        ("圖片來源：OpenMoji 17.0.0（openmoji.org），CC BY-SA 4.0。", 9.5, MUTED, 4),
        ("OpenMoji 沒有的概念另取自 game-icons.net（CC BY 3.0）、Phosphor Icons（MIT）、"
         "Material Design Icons（Apache 2.0）與 Tabler Icons（MIT）。", 9.5, MUTED, 4),
    ]
    lines.extend((text, 9.5, MUTED, 4) for text in deck["sources"])
    for text, size, color, after in lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.space_before = Pt(0)
        write(paragraph, text, FONT_UI, size, color=color, bold=size >= 30)
    page_break(document)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write(note, "（此頁留白，供雙面列印對齊）", FONT_UI, 9.5, color=MUTED)


def load_cards(deck: dict) -> list[dict]:
    entries = json.loads(deck["vocab"].read_text(encoding="utf-8"))
    gloss_payload = (json.loads(deck["glosses"].read_text(encoding="utf-8"))
                     if "glosses" in deck else {})
    images = json.loads(deck["images"].read_text(encoding="utf-8"))["images"]
    icons = (json.loads(deck["icons"].read_text(encoding="utf-8"))["cards"]
             if "icons" in deck and deck["icons"].exists() else {})
    cards: list[dict] = []

    def picture(key: str) -> Path | None:
        record = icons.get(key)
        if record:
            return ICON_DIR / record["file"]
        record = images.get(key)
        return IMAGE_DIR / record["file"] if record else None

    if "volumeName" in deck:  # Latin
        rows = [item for item in entries["entries"] if item["volume"] == deck["volumeName"]]
        for entry in sorted(rows, key=lambda item: item["ordinal"]):
            gloss = (entry.get("glossZh") or "").strip()
            if not gloss:
                raise SystemExit(f"{entry['headword']} 缺繁中詞義")
            printed = (entry.get("forms") or entry["headword"]).strip()
            cards.append({
                "headword": printed,
                "glossZh": gloss,
                "pos": POS_LATIN.get(latin_pos(entry), ""),
                "lesson": entry["lesson"],
                "picture": picture(latin_key(entry)),
            })
        return cards

    if "volume" in deck:  # Greek
        entries = [item for item in entries["entries"] if item["volume"] == deck["volume"]]
        glosses = {lemma: record["glossZh"] for lemma, record in gloss_payload["glosses"].items()}
        for entry in sorted(entries, key=lambda item: item["ordinal"]):
            gloss = glosses.get(entry["lemma"], "").strip()
            if not gloss:
                raise SystemExit(f"{entry['lemma']} 缺繁中詞義")
            part_of_speech = greek_part_of_speech(entry, gloss)
            cards.append({
                # 讀本印完整詞典形；卡片只在推不出來時才印，其餘只印詞頭。
                "headword": card_headword(entry, part_of_speech),
                "glossZh": gloss,
                "pos": part_of_speech,
                "lesson": entry["lesson"],
                "picture": picture(entry["lemma"]),
            })
        return cards

    glosses = {
        (item["strong"], item["pointed"]): item["glossZh"]
        for item in gloss_payload["items"]
    }
    grammar = (json.loads(deck["grammar"].read_text(encoding="utf-8"))["cards"]
               if "grammar" in deck else {})
    for entry in sorted(entries, key=lambda item: item["ordinal"]):
        key = (entry["strong"], entry["pointed"])
        if key not in glosses:
            raise SystemExit(f"{entry['pointed']} 缺繁中詞義")
        card_key = f"{entry['strong']}|{entry['pointed']}"
        cards.append({
            "headword": entry["pointed"],
            "glossZh": glosses[key],
            "pos": hebrew_part_of_speech(entry, grammar.get(card_key, {})),
            "lesson": entry["lesson"],
            "picture": picture(card_key),
        })
    return cards


def build(deck: dict, cards: list[dict]) -> Path:
    document = Document()
    configure(document)
    sheets = -(-len(cards) // (COLS * ROWS))
    with_picture = sum(1 for card in cards if card["picture"])
    add_cover(document, deck, len(cards), sheets, with_picture)

    for start in range(0, len(cards), COLS * ROWS):
        page = cards[start : start + COLS * ROWS]
        for side in ("front", "back"):
            page_break(document)
            table = new_grid(document)
            for index, card in enumerate(page):
                row, column = divmod(index, COLS)
                # Duplex alignment: the back sheet runs right to left.
                target = table.cell(row, column if side == "front" else COLS - 1 - column)
                target_column = column if side == "front" else COLS - 1 - column
                # 圖形 id 在整份文件裡必須唯一，否則 Word 只認第一個。
                place = (target_column, row, 1000 + start * 2 + index * 2 + (side == "back"))
                if side == "front":
                    fill_front(target, card, deck, place)
                else:
                    fill_back(target, card, card["picture"], place)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / deck["output"]
    document.save(path)
    return path


def main() -> None:
    parser = argparse.ArgumentParser(description="產生單字卡（A4 橫式、每頁 8 張、雙面）")
    parser.add_argument("--deck", choices=sorted(DECKS), default="hbo")
    parser.add_argument("--limit", type=int, default=0, help="只做前 N 張（試印用）")
    args = parser.parse_args()

    deck = DECKS[args.deck]
    cards = load_cards(deck)
    if args.limit:
        cards = cards[: args.limit]
    path = build(deck, cards)

    sheets = -(-len(cards) // (COLS * ROWS))
    with_picture = sum(1 for card in cards if card["picture"])
    no_pos = sum(1 for card in cards if not card["pos"])
    print(f"  {deck['title']}：{len(cards)} 張，正反共 {sheets * 2} 頁"
          f"（每頁 8 張，{CARD_W_MM:.2f}×{CARD_H_MM:.0f} mm）")
    print(f"  有圖 {with_picture}，留白 {len(cards) - with_picture}；未標詞性 {no_pos}")
    print(path)


if __name__ == "__main__":
    main()
