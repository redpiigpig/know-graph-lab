# -*- coding: utf-8 -*-
"""城中週報 raw extractor — walk 周報 year folders, pull paragraphs + tables + filename date.
.docx -> python-docx ; .doc -> Word COM ; .htm -> bs4. Output JSONL (one row per file)."""
import sys, os, re, json, glob, traceback
sys.stdout.reconfigure(encoding='utf-8')

ROOT = os.path.abspath(r"pong-archive/stores/城中週報")
OUT = r"c:/tmp/cz_bulletins"
os.makedirs(OUT, exist_ok=True)

# Year folders that hold worship bulletins (周報). Exclude 講章/講道/錄音 subfolders.
YEAR_DIRS = sorted(glob.glob(os.path.join(ROOT, "20*周報")) +
                   glob.glob(os.path.join(ROOT, "20*週報午堂")))
LOOSE_2026 = sorted(glob.glob(os.path.join(ROOT, "26-*.docx")))  # 2026 not yet foldered

DATE_RE = re.compile(r"(?<!\d)(\d{2})-(\d{2})-(\d{2})(?!\d)")        # 24-01-07
DATE8_RE = re.compile(r"(?<!\d)(20\d{2})(\d{2})(\d{2})(?!\d)")        # 20030202

def parse_date(fname):
    m = DATE_RE.search(fname)
    if m:
        yy, mm, dd = m.groups()
        return f"20{yy}-{mm}-{dd}"
    m = DATE8_RE.search(fname)
    if m:
        return f"{m.group(1)}-{m.group(2)}-{m.group(3)}"
    return None

NONBULLETIN = ("經課", "索引", "index", "通訊", "代禱", "月曆", "readme", "目錄", "名單", "報名")

def is_bulletin_name(fname):
    low = fname.lower()
    if fname.startswith("~$") or fname.endswith(".tmp") or low in ("desktop.ini", "thumbs.db"):
        return False
    if "崇拜" in fname or "程序" in fname or "週報" in fname or "周報" in fname:
        return True
    if any(x in low for x in NONBULLETIN):
        return False
    # date-named files inside 2003 zips (after excluding lectionary/index above)
    return parse_date(fname) is not None

def clean(t):
    return t.replace("\r", " ").replace("\x07", " ").replace("\x0b", " ").replace("\xa0", " ").strip()

# ---- format readers ----
def read_docx(path):
    import docx
    d = docx.Document(path)
    paras = [clean(p.text) for p in d.paragraphs]
    paras = [p for p in paras if p]
    tables = []
    for t in d.tables:
        rows = []
        for r in t.rows:
            rows.append([clean(c.text) for c in r.cells])
        tables.append(rows)
    return paras, tables

def read_htm(path):
    from bs4 import BeautifulSoup
    for enc in ("utf-8", "big5", "cp950", "gb18030"):
        try:
            html = open(path, "r", encoding=enc, errors="strict").read(); break
        except Exception:
            html = None
    if html is None:
        html = open(path, "r", encoding="utf-8", errors="ignore").read()
    soup = BeautifulSoup(html, "html.parser")
    tables = []
    for t in soup.find_all("table"):
        rows = []
        for tr in t.find_all("tr"):
            rows.append([clean(td.get_text(" ")) for td in tr.find_all(["td","th"])])
        if rows:
            tables.append(rows)
    paras = [clean(x) for x in soup.get_text("\n").split("\n")]
    paras = [p for p in paras if p]
    return paras, tables

_word = None
def get_word():
    global _word
    if _word is None:
        import win32com.client as win32
        _word = win32.gencache.EnsureDispatch("Word.Application")
        _word.Visible = False
        _word.DisplayAlerts = False
    return _word

def read_doc(path):
    word = get_word()
    doc = word.Documents.Open(path, ReadOnly=True, AddToRecentFiles=False, Visible=False)
    try:
        # one COM call for all flow text; cell mark = \x07, row/para end = \r
        content = doc.Content.Text
        paras = [clean(x) for x in re.split(r"[\r\n]", content)]
        paras = [p for p in paras if p]
        tables = []
        for ti in range(1, doc.Tables.Count + 1):
            txt = doc.Tables(ti).Range.Text          # one COM call per table
            cells = [clean(c) for c in re.split(r"[\x07\r\n]", txt)]
            cells = [c for c in cells if c]
            if cells:
                tables.append([cells])               # single flat row; downstream flattens
        return paras, tables
    finally:
        doc.Close(False)

def read_any(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return read_docx(path)
    if ext == ".doc":
        return read_doc(path)
    if ext in (".htm", ".html"):
        return read_htm(path)
    raise ValueError("unsupported " + ext)

EXCLUDE_BRANCHES = ("講章", "講道", "錄音", "主日錄音", "崇拜錄音", "_原始壓縮檔")

def _is_bulletin_branch(rel):
    """True if rel path belongs to a 周報 branch (legacy 20XX周報/ or reorganized 周報/)."""
    top = rel.replace("\\", "/").split("/")[0]
    if top in EXCLUDE_BRANCHES:
        return False
    return top == "周報" or re.match(r"^20\d{2}周報$", top) or top == "2007週報午堂" \
        or re.match(r"^26-\d{2}-\d{2}", top)

def collect_files():
    files = []
    for dp, dn, fn in os.walk(ROOT):
        if dp.endswith(".files") or os.path.basename(dp) in ("images", "Folder Settings"):
            continue
        rel = os.path.relpath(dp, ROOT)
        if rel != "." and not _is_bulletin_branch(rel):
            continue
        for f in fn:
            if os.path.splitext(f)[1].lower() not in (".doc", ".docx", ".htm", ".html"):
                continue
            if rel == "." and not re.match(r"^26-\d{2}-\d{2}", f):
                continue  # at ROOT, only loose 2026 bulletins
            if not is_bulletin_name(f):
                continue
            files.append(os.path.join(dp, f))
    return sorted(set(files))

def main():
    files = collect_files()
    print(f"Collected {len(files)} candidate bulletin files")
    out_path = os.path.join(OUT, "raw.jsonl")
    n_ok = n_err = 0
    with open(out_path, "w", encoding="utf-8") as out:
        for i, path in enumerate(files):
            rel = os.path.relpath(path, ROOT)
            fname = os.path.basename(path)
            date = parse_date(fname) or parse_date(rel)
            rec = {"file": rel.replace("\\", "/"), "fname": fname, "date": date,
                   "ext": os.path.splitext(fname)[1].lower(),
                   "mtime": int(os.path.getmtime(path))}
            try:
                paras, tables = read_any(path)
                rec["paras"] = paras
                rec["tables"] = tables
                rec["nchars"] = sum(len(p) for p in paras)
                n_ok += 1
            except Exception as e:
                rec["error"] = f"{type(e).__name__}: {e}"
                rec["trace"] = traceback.format_exc()[-500:]
                n_err += 1
            out.write(json.dumps(rec, ensure_ascii=False) + "\n")
            if (i + 1) % 50 == 0:
                print(f"  {i+1}/{len(files)}  ok={n_ok} err={n_err}")
    if _word is not None:
        _word.Quit()
    print(f"DONE ok={n_ok} err={n_err} -> {out_path}")

if __name__ == "__main__":
    main()
