from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\user\Desktop\know-graph-lab")
ASSETS = ROOT / "docs" / "hcu-manual-assets"
OUTPUT = ROOT / "docs" / "玄奘佛學研究網站建置操作手冊.docx"
COVER = ROOT / "public" / "Hsuan_Chuang_Studies" / "covers" / "v45.jpg"

# compact_reference_guide preset
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION = "7A5A00"
RISK = "9B1C1C"
MUTED = "5F6B76"
WHITE = "FFFFFF"
GOLD_FILL = "FFF6D8"
RED_FILL = "FCE8E8"
GREEN_FILL = "E7F4EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B7C5D3", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = widths_dxa[min(i, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
    if table.rows:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = tr_pr.find(qn("w:tblHeader"))
        if tbl_header is None:
            tbl_header = OxmlElement("w:tblHeader")
            tr_pr.append(tbl_header)
        tbl_header.set(qn("w:val"), "true")


def set_font(run, size=None, bold=None, color=None, italic=None, east_asia="Microsoft JhengHei"):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_border_bottom(paragraph, color="D3DCE5", size="8", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)


def add_page_field(paragraph):
    begin_run = paragraph.add_run()
    set_font(begin_run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(fld_begin)

    instr_run = paragraph.add_run()
    set_font(instr_run, size=9, color=MUTED)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)

    sep_run = paragraph.add_run()
    set_font(sep_run, size=9, color=MUTED)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_sep)

    result_run = paragraph.add_run("1")
    set_font(result_run, size=9, color=MUTED)

    end_run = paragraph.add_run()
    set_font(end_run, size=9, color=MUTED)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    if "Caption Manual" not in styles:
        cap = styles.add_style("Caption Manual", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Caption Manual"]
    cap.font.name = "Calibri"
    cap.font.size = Pt(9)
    cap.font.color.rgb = RGBColor.from_string(MUTED)
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = True

    list_bullet = styles["List Bullet"]
    list_bullet.font.name = "Calibri"
    list_bullet.font.size = Pt(11)
    list_bullet._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    list_bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    list_bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    list_bullet.paragraph_format.left_indent = Inches(0.375)
    list_bullet.paragraph_format.first_line_indent = Inches(-0.188)
    list_bullet.paragraph_format.space_after = Pt(4)
    list_bullet.paragraph_format.line_spacing = 1.25


def add_numbering_definition(doc, num_type="decimal"):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if num_type == "bullet" else "decimal")
    lvl.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "●" if num_type == "bullet" else "%1.")
    lvl.append(text)
    restart = OxmlElement("w:lvlRestart")
    restart.set(qn("w:val"), "1")
    lvl.append(restart)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    r_pr.append(fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id, bold_prefix=None):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    n_id = OxmlElement("w:numId")
    n_id.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, n_id])
    p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet_item(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)
    return p


def add_callout(doc, title, body, kind="info"):
    colors = {
        "info": (LIGHT_BLUE, INK),
        "success": (GREEN_FILL, DARK_BLUE),
        "warning": (GOLD_FILL, CAUTION),
        "danger": (RED_FILL, RISK),
    }
    fill, color = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(title + "｜")
    set_font(r, bold=True, color=color)
    r2 = p.add_run(body)
    set_font(r2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    hdr = table.rows[0]
    for i, heading in enumerate(headers):
        set_cell_shading(hdr.cells[i], LIGHT_BLUE)
        hdr.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr.cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(heading)
        set_font(r, bold=True, color=INK)
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(str(value))
            set_font(r, size=10.5)
    set_table_geometry(table, widths_dxa)
    return table


def add_picture(doc, filename, caption, width=6.3):
    path = ASSETS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_together = True
    run = p.add_run()
    picture = run.add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    c = doc.add_paragraph(style="Caption Manual")
    c.add_run(caption)


def add_step(doc, number, title, instructions, num_id, screenshot=None, caption=None):
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(f"步驟 {number}｜{title}")
    for item in instructions:
        add_bullet_item(doc, item)
    if screenshot:
        add_picture(doc, screenshot, caption or title)


def add_section_title(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    set_keep_with_next(p)
    return p


def add_subsection(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    set_keep_with_next(p)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_page_break(doc):
    doc.add_page_break()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(2)
    hp.paragraph_format.line_spacing = 1.0
    hr = hp.add_run("玄奘佛學研究網站建置操作手冊")
    set_font(hr, size=9, bold=True, color=MUTED)
    set_paragraph_border_bottom(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    add_page_field(fp)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("實作操作手冊")
    set_font(r, size=11, bold=True, color=CAUTION)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("玄奘佛學研究")
    set_font(r, size=30, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("建置於「臺灣佛教研究中心」網站後台")
    set_font(r, size=15, color=DARK_BLUE)

    if COVER.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        cover_picture = p.add_run().add_picture(str(COVER), height=Inches(2.8))
        cover_picture._inline.docPr.set("descr", "玄奘佛學研究第四十五期封面")
        cover_picture._inline.docPr.set("title", "玄奘佛學研究第四十五期封面")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("版本 1.0｜2026 年 7 月 21 日")
    set_font(r, size=10.5, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("依玄奘大學後台實際介面與官方操作手冊整理")
    set_font(r, size=9.5, italic=True, color=MUTED)
    add_page_break(doc)


def build():
    doc = Document()
    configure_styles(doc)
    configure_document(doc)
    bullet_id = None
    number_id = None
    add_cover(doc)

    add_section_title(doc, "先看結論")
    add_callout(
        doc,
        "可以建置",
        "後台支援圖片列表、文章內容、PDF／Word 附件、超連結、預覽與發布。玄奘佛學研究封面可以點進各期內容頁。",
        "success",
    )
    add_table(
        doc,
        ["示範網站功能", "後台做法", "可行性"],
        [
            ("期刊封面牆", "多篇＋圖片列表；每一期建立一篇文章，第一張圖片做封面", "可行"),
            ("點封面進入單期", "圖片列表預設連到該篇文章內容；不填導轉網址", "可行"),
            ("單期論文表格", "文章內容加入表格，論文標題或 PDF 欄設超連結", "可行"),
            ("PDF 下載", "檔案庫上傳／沿用既有 PDF，再插入網站檔案連結", "可行"),
            ("投稿指引等固定頁", "建立單篇節點，版型選文章內容", "可行"),
            ("Word 範本下載", "固定頁的附件區加入 .doc／.docx", "可行"),
            ("中英即時切換", "現有中心僅有中文語系；英文需另建 en-us 語系或英文頁", "需另規劃"),
        ],
        [2200, 5560, 1600],
    )
    add_body(doc, "最重要的目的地：內容要建在「臺灣佛教研究中心 ＞ 中文 ＞ 首頁」之下。後台另有一個「玄奘佛教研究」空白網站樹，除非網站管理員明確指定，否則不要把正式內容建到那裡。")
    add_callout(doc, "安全原則", "本手冊不保存帳號密碼。操作中不要按「刪除」；正式發布前先做一期試刊並請中心確認。", "warning")

    add_subsection(doc, "建置量與對應位置")
    add_table(
        doc,
        ["項目", "示範站數量", "在後台的建置方式"],
        [
            ("期刊封面", "45 張", "每期文章的第一張圖片"),
            ("期別內容頁", "45 篇", "玄奘佛學研究節點內的 45 篇文章"),
            ("論文資料", "305 篇", "各期文章中的表格與 PDF 連結"),
            ("固定說明頁", "4 頁", "編輯委員、投稿指引、審查流程、學術倫理"),
            ("下載檔", "PDF／Word", "檔案庫、附件或文章內超連結"),
        ],
        [1900, 1600, 5860],
    )

    add_section_title(doc, "建議網站結構")
    add_body(doc, "建議先完成中文正式版。以下結構可保留目前臺灣佛教研究中心的導覽，同時加入期刊專區。")
    tree = [
        "臺灣佛教研究中心",
        "└─ 中文",
        "   └─ 首頁",
        "      ├─ 玄奘佛學研究（多篇／圖片列表／文章內容）",
        "      │  ├─ 第四十五期（文章）",
        "      │  ├─ 第四十四期（文章）",
        "      │  └─ ……第一期（文章）",
        "      ├─ 編輯委員（單篇／文章內容）",
        "      ├─ 投稿指引（單篇／文章內容＋Word 附件）",
        "      ├─ 審查流程（單篇／文章內容）",
        "      └─ 學術倫理（單篇／文章內容）",
    ]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="C9D5E1", size="6")
    set_cell_shading(table.cell(0, 0), CALLOUT)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(tree):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_font(r, size=10.5, bold=(idx == 0), color=INK)
    add_body(doc, "若希望四個固定頁顯示成「玄奘佛學研究」的下拉子選單，可改成在該節點上按右鍵建立子節點；先由中心確認主選單樣式，再選擇此做法。")

    add_subsection(doc, "上稿前先整理檔案")
    for item in [
        "建立一個資料夾，每一期各放：封面圖、論文 PDF、作者與頁碼清單。",
        "封面建議使用 JPG 或 PNG；所有期數採一致比例，避免列表高低不齊。",
        "檔名建議使用半形英數與連字號，例如 45-cover.jpg、45-01-author-keyword.pdf。",
        "先檢查後台檔案庫；目前已看到多期玄奘佛學研究 PDF，能沿用就不要重複上傳。",
        "絕對不要貼入 C:\\ 或 file:/// 開頭的電腦本機路徑；公開網站讀不到這些檔案。",
    ]:
        add_bullet_item(doc, item)
    add_callout(doc, "已發現的風險", "公開頁若出現 file:///C:/Users/...，表示連結誤指向上稿者電腦。正確做法是先上傳到後台檔案庫，再插入網站上的檔案網址。", "danger")

    add_section_title(doc, "第一階段｜登入與找到正確位置")
    add_step(
        doc,
        1,
        "登入後台",
        [
            "開啟 https://www.hcu.edu.tw/backend/login.aspx?ReturnUrl=%2Fbackend%2F。",
            "輸入學校配發的帳號、密碼與畫面驗證碼，按「登入」。",
            "不要把密碼寫在手冊、共用文件或瀏覽器畫面截圖裡。",
        ],
        number_id,
    )
    add_step(
        doc,
        2,
        "進入節點管理",
        [
            "上方選單按「CMS」。",
            "左側按「節點管理」。",
            "依序展開「臺灣佛教研究中心 ＞ 中文 ＞ 首頁」。",
            "在「首頁」上按滑鼠右鍵，選「建立」。",
        ],
        number_id,
        "01-node-context-menu.png",
        "圖 1｜在「首頁」節點按右鍵，選「建立」。右鍵選單同時有編輯、複製、重新命名、刪除與內容。",
    )

    add_section_title(doc, "第二階段｜建立「玄奘佛學研究」節點")
    add_step(
        doc,
        3,
        "填寫新節點基本資料",
        [
            "名稱：輸入「玄奘佛學研究」。",
            "子標題：可留白，或填「Hsuan Chuang Journal of Buddhism Studies」。",
            "友善網址：建議填 hsuan-chuang-studies；只用小寫英文字母、數字與連字號。",
            "網址：一般內部內容頁先留白；只有外部導轉才填。",
            "SEO 視窗標題：可填「玄奘佛學研究｜臺灣佛教研究中心」。",
            "先不要送出，往下完成版型與顯示設定。",
        ],
        number_id,
        "03-template-settings.png",
        "圖 2｜建立／編輯節點視窗。必填欄位為「名稱」，節點 ID 由系統自動產生。",
    )
    add_step(
        doc,
        4,
        "設定版型",
        [
            "版型設定選「多篇」。",
            "列表版型選「圖片列表」。",
            "內容版型選「文章內容」。",
            "系統設定：啟用選「是」、顯示選「是」。",
            "檢查無誤後按「送出」。",
        ],
        number_id,
    )
    add_callout(doc, "為什麼這樣選", "圖片列表會使用每篇文章的第一張圖片與標題；文章內容則可放表格、PDF 連結、圖片與附件。不要選只能顯示單張圖的「圖片內容」。", "info")

    add_section_title(doc, "第三階段｜新增第一期測試文章")
    add_body(doc, "先以「第四十五期」做試刊。通過畫面、PDF 與手機檢查後，再批次建立其他期數。")
    add_step(
        doc,
        5,
        "進入內容列表",
        [
            "回到「CMS ＞ 節點管理」。",
            "在「玄奘佛學研究」節點按右鍵，選「內容」。",
            "也可從「CMS ＞ 內容管理」找到同一節點。",
            "內容列表上方按「＋新增」。不要按任何既有資料列的「刪除」。",
        ],
        number_id,
        "04-content-list.png",
        "圖 4｜多篇節點的內容列表。上方有「調整順序」「節點內容」「新增」，資料列可發布、修改或刪除。",
    )

    add_step(
        doc,
        6,
        "填寫期別文章欄位",
        [
            "標題：輸入「第四十五期」；其他期數使用相同命名格式。",
            "導轉網址：留白。留白時封面會進入本篇文章內容；只有要直接跳外部頁面或 PDF 才填。",
            "重要性：一般使用預設值；若要置頂，再依中心規則調整。",
            "瀏覽權限：選「公開」。",
            "顯示日期：期刊典藏頁通常可把「是否顯示日期」設為不顯示；若中心要顯示出版日，再填正確日期。",
            "生效時間、過期時間：若要立即長期公開，可依中心既有文章設定；不要誤填已過期日期。",
        ],
        number_id,
        "05-article-editor.png",
        "圖 5｜文章編輯頁上半部：標題、導轉網址、瀏覽權限、日期、SEO，以及圖片／連結／附件／內容／影片區塊。",
    )

    add_section_title(doc, "第四階段｜讓封面可以點進單期內容")
    add_step(
        doc,
        7,
        "加入封面圖片",
        [
            "在文章編輯頁展開「圖片」。",
            "按「增加」，從圖片庫選擇或上傳本期封面。",
            "把封面設為本篇文章的第一張圖片。圖片列表會以第一張圖作為封面。",
            "替代文字建議填「玄奘佛學研究第四十五期封面」，方便無障礙與搜尋。",
        ],
        number_id,
    )
    add_callout(
        doc,
        "封面點擊規則",
        "要進入本期內容頁：導轉網址留白。要跳到別的網頁：才填導轉網址。若填成 PDF，封面將直接開 PDF，不會先看到論文目錄。",
        "success",
    )
    add_subsection(doc, "建議先做的測試")
    for item in [
        "預覽圖片列表：封面比例、標題、排列是否正確。",
        "點封面：是否進入「第四十五期」文章內容。",
        "用手機寬度查看：封面與文字是否被裁切。",
        "若點擊無反應：檢查節點是否為「圖片列表」、文章是否已啟用／發布，以及模板是否被改動。",
    ]:
        add_bullet_item(doc, item)

    add_subsection(doc, "若中心要求封面直接開啟外部頁")
    add_body(doc, "可在「導轉網址」貼入完整 https:// 網址；若該頁是本站新頁，先發布新頁取得正式網址，再回來設定導轉。不要貼入後台編輯網址，也不要貼入 file:/// 或 C:\\ 路徑。")

    add_section_title(doc, "第五階段｜建立單期論文目錄與 PDF 連結")
    add_step(
        doc,
        8,
        "開啟文章內容編輯器",
        [
            "在文章編輯頁展開「內容」。",
            "按「增加」。",
            "使用工具列的「表格」建立 4 欄：篇名、作者、頁碼、PDF。",
            "每篇論文建立一列；篇名與作者請核對期刊原始資料。",
        ],
        number_id,
        "05b-rich-text-editor.png",
        "圖 6｜文章內容編輯器：可加入表格、影像、超連結、字型與對齊等內容。",
    )
    add_subsection(doc, "論文表格建議格式")
    add_table(
        doc,
        ["篇名", "作者", "頁碼", "PDF"],
        [
            ("論文標題示例", "作者姓名", "1–24", "下載 PDF"),
            ("Another Article Title", "Author Name", "25–48", "PDF"),
        ],
        [4400, 1900, 1200, 1860],
    )
    add_step(
        doc,
        9,
        "替 PDF 文字加入網站連結",
        [
            "先到檔案庫確認該 PDF 是否已存在；同一檔案不要重複上傳。",
            "在編輯器反白「下載 PDF」或論文標題。",
            "按工具列的「插入／編輯超連結」。",
            "貼上檔案庫提供的網站網址，必須是 http:// 或 https:// 開頭。",
            "建議勾選在新視窗開啟；按確定後，再逐一測試。",
        ],
        number_id,
    )
    add_callout(doc, "連結檢查", "正確連結應是學校網站網址；看到 C:\\、file:///、桌面或 Downloads 字樣，立即取消並重新從檔案庫取得網址。", "danger")

    add_section_title(doc, "第六階段｜使用圖片庫與檔案庫")
    add_subsection(doc, "圖片庫：封面、文章插圖")
    for item in [
        "上方選單按「檔案管理 ＞ 圖片庫管理」。",
        "將 JPG／JPEG／PNG／GIF 拖到「點選或拖曳圖片至這邊」，或點該區選檔。",
        "單檔上限 100MB；一般封面不應接近上限，建議先壓縮圖片。",
        "上傳後按「預覽」確認；除非非常確定且已獲授權，不要按「刪除」。",
    ]:
        add_bullet_item(doc, item)
    add_picture(doc, "07-picture-library.png", "圖 7｜圖片庫管理：可點選或拖曳圖片，上傳後可預覽。", width=6.0)

    add_subsection(doc, "檔案庫：PDF、Word、其他下載檔")
    for item in [
        "上方選單按「檔案管理 ＞ 檔案庫管理」。",
        "先用清單查找既有期刊 PDF；本次檢查已看到多期玄奘佛學研究檔案。",
        "找不到才上傳；拖入檔案區後，等上傳完成再離開頁面。",
        "按「預覽」確認能開啟，再把網站檔案網址加入文章超連結。",
        "Word 投稿格式也可放入檔案庫，或直接加入固定頁的「附件」區。",
    ]:
        add_bullet_item(doc, item)
    add_picture(doc, "06-file-library.png", "圖 8｜檔案庫管理：已有多筆檔案；上傳前先查重，避免重複與誤刪。")
    add_callout(doc, "不要清理舊檔", "即使檔名看起來重複，也可能已被公開文章引用。沒有完整引用清單與管理員核准，不要刪除既有圖片或檔案。", "warning")

    add_page_break(doc)
    add_section_title(doc, "第七階段｜建立四個固定說明頁")
    add_body(doc, "依照「第二階段」建立節點，但每一頁的版型設定改為「單篇＋文章內容」。建議先作為首頁下的同層節點；若中心確認要做下拉選單，再移到玄奘佛學研究下方。")
    add_table(
        doc,
        ["節點名稱", "建議友善網址", "內容與附件"],
        [
            ("編輯委員", "editorial-board", "委員名單、職稱、單位；必要時中英文並列"),
            ("投稿指引", "submission-guidelines", "投稿規範、格式說明、Word 範本附件"),
            ("審查流程", "review-process", "收稿、初審、外審、修訂、錄用流程"),
            ("學術倫理", "publication-ethics", "研究倫理、抄襲、利益衝突、撤稿規則"),
        ],
        [1900, 2300, 5160],
    )
    add_step(
        doc,
        10,
        "建立單篇節點",
        [
            "在「首頁」上按右鍵選「建立」。",
            "填入節點名稱與友善網址。",
            "版型選「單篇」，內容版型選「文章內容」。",
            "啟用選「是」、顯示選「是」，按「送出」。",
            "在新節點按右鍵選「內容」，加入正文。",
        ],
        number_id,
    )
    add_step(
        doc,
        11,
        "加入 Word 下載",
        [
            "進入「投稿指引」內容編輯頁，展開「附件」。",
            "按「增加」，選擇已上傳的 .doc 或 .docx；若無檔案，再上傳。",
            "附件顯示名稱請寫清楚，例如「玄奘佛學研究投稿格式範本（Word）」。",
            "預覽頁面並實際下載，確認檔案可開啟且版本正確。",
        ],
        number_id,
    )

    add_subsection(doc, "中英文內容怎麼處理")
    add_body(doc, "目前臺灣佛教研究中心後台只看到中文語系節點，因此示範站的即時「中／英」切換無法只靠一般文章欄位原樣複製。建議分兩階段：")
    for item in [
        "第一階段：先完成中文期刊專區，英文篇名可直接在同一論文表格中並列。",
        "第二階段：請網站管理員建立或確認 en-us 英文語系，再複製節點與英文內容。",
        "若只需要少量英文說明，可在中文固定頁下方加入英文段落，避免先改整站語系架構。",
    ]:
        add_bullet_item(doc, item)

    add_section_title(doc, "第八階段｜預覽、儲存與發布")
    add_step(
        doc,
        12,
        "先預覽",
        [
            "文章編輯頁按「預覽」。",
            "確認標題、封面、表格、作者、頁碼與所有連結。",
            "用電腦與手機寬度查看；封面不可變形、表格不可超出畫面。",
            "若預覽有問題，回編輯頁修正，不要直接發布。",
        ],
        number_id,
    )
    add_step(
        doc,
        13,
        "儲存草稿",
        [
            "按「儲存」回到內容列表。",
            "確認該筆資料標題與狀態正確。",
            "若要調整期數順序，按「調整順序」，把最新一期放前面。",
        ],
        number_id,
    )
    add_step(
        doc,
        14,
        "發布",
        [
            "請中心確認試刊內容後，在該筆資料列按「發布」。",
            "若發布視窗要求方向，正式中心的一般內容應只在本網站範圍發布；不確定時先詢問管理員，不要任意選向上發布。",
            "按「送出」後，另開未登入或無痕視窗查看公開頁。",
            "逐一測試封面、返回列表、PDF、Word 與固定頁。",
        ],
        number_id,
    )
    add_callout(doc, "發布不是儲存", "「儲存」只保留後台內容；公開網站是否更新，要以登出後的公開頁為準。若內容未出現，先檢查發布狀態、生效／過期時間、啟用與顯示設定。", "warning")

    add_page_break(doc)
    add_section_title(doc, "上線檢查清單")
    checks = [
        "建立位置是「臺灣佛教研究中心 ＞ 中文 ＞ 首頁」，不是另一個空白網站樹。",
        "玄奘佛學研究節點為：多篇／圖片列表／文章內容。",
        "節點與文章都已啟用、顯示，且生效時間正確、尚未過期。",
        "每一期都有標題，第一張圖片是該期封面。",
        "封面點擊後進入正確期別；返回列表仍可正常操作。",
        "論文表格的篇名、作者、頁碼與 PDF 一致。",
        "所有 PDF／Word 連結都是 http:// 或 https://，沒有 C:\\ 或 file:///。",
        "下載檔能開啟，沒有傳錯版本，也沒有需要登入才能讀取。",
        "手機畫面封面不變形、表格不被截斷。",
        "固定頁內容與附件正確，主選單排序符合中心要求。",
        "以登出或無痕視窗重新檢查，確認不是只有管理員看得到。",
        "完成後按右上「登出」。",
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [750, 8610])
    set_table_borders(table)
    set_cell_shading(table.cell(0, 0), LIGHT_BLUE)
    set_cell_shading(table.cell(0, 1), LIGHT_BLUE)
    for i, txt in enumerate(("完成", "檢查項目")):
        r = table.cell(0, i).paragraphs[0].add_run(txt)
        set_font(r, bold=True, color=INK)
    for item in checks:
        cells = table.add_row().cells
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run("☐")
        set_font(r0, size=13, color=DARK_BLUE)
        r1 = cells[1].paragraphs[0].add_run(item)
        set_font(r1, size=10.5)
        for cell in cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    set_table_geometry(table, [750, 8610])

    add_subsection(doc, "批次上稿建議")
    add_table(
        doc,
        ["批次", "範圍", "完成條件"],
        [
            ("試刊", "第 45 期", "中心確認版面、封面點擊、PDF、手機畫面"),
            ("第二批", "第 44–41 期", "檔名與表格格式固定，可重複使用"),
            ("後續批次", "每次 5–10 期", "每批發布後抽查全部封面與至少每期 2 個 PDF"),
            ("總驗收", "第 45–1 期", "45 期、305 篇資料完整，固定頁與附件可用"),
        ],
        [1400, 2100, 5860],
    )

    add_section_title(doc, "常見問題與排除方法")
    add_table(
        doc,
        ["現象", "優先檢查", "處理方式"],
        [
            ("封面不能點", "列表版型、文章狀態", "確認圖片列表、第一張圖片、啟用／顯示／發布；導轉網址若錯誤則清除"),
            ("點封面直接開 PDF", "導轉網址", "若要先看目錄，把導轉網址留白，PDF 放入文章內容"),
            ("PDF 在自己電腦可開，別人不行", "網址是否 file:/// 或 C:\\", "上傳檔案庫，改成學校網站 http(s) 連結"),
            ("公開頁沒有更新", "發布、日期、快取", "確認已發布且在生效期間；仍未更新再請管理員協助清快取"),
            ("主選單看不到", "節點顯示／啟用", "兩者都設為「是」，並檢查父節點與排序"),
            ("封面比例不一", "原始圖比例", "統一裁切比例後重新上傳；不要在編輯器硬拉寬高"),
            ("同一 PDF 出現多份", "檔案庫查重", "新上傳前搜尋；既有重複檔不要自行刪除"),
            ("中文有頁面，英文沒有", "語系節點", "先中英並列；完整英文站需管理員建立 en-us 架構"),
        ],
        [2200, 2300, 4860],
    )

    add_subsection(doc, "建議資料命名")
    add_table(
        doc,
        ["類型", "建議檔名", "說明"],
        [
            ("封面", "45-cover.jpg", "期數在前，所有期數同一規則"),
            ("論文 PDF", "45-01-author-keyword.pdf", "期數－篇序－作者／短題名"),
            ("投稿範本", "submission-template-2026.docx", "加入年份或版本，避免誤用舊檔"),
            ("倫理規範", "publication-ethics-2026.pdf", "若另有正式 PDF，可在頁面同時提供"),
        ],
        [1600, 3300, 4460],
    )

    add_page_break(doc)
    add_section_title(doc, "正式上稿前的最後確認")
    add_callout(doc, "建議決策", "先發布第 45 期作為可回復的試刊；中心確認後再按批次匯入。這樣可以在大量上稿前修正版型、檔名與連結規則。", "success")
    add_body(doc, "本次檢查只進行後台唯讀瀏覽、開啟表單與取消，沒有儲存、發布、上傳或刪除任何正式網站內容；檢查完成後已登出。")
    add_subsection(doc, "交付給中心確認的 6 個問題")
    for item in [
        "「玄奘佛學研究」要放在現有主選單哪一個位置？",
        "四個固定頁要做主選單同層，還是期刊下拉子選單？",
        "期刊封面要進單期目錄，還是直接開整期 PDF？",
        "論文 PDF 是否全部公開？是否有延後公開期？",
        "英文版先採中英並列，還是正式建立 en-us 語系？",
        "每批上稿由誰預覽、誰核准、誰按發布？",
    ]:
        add_bullet_item(doc, item)

    add_subsection(doc, "快速按鍵路徑")
    add_table(
        doc,
        ["要做的事", "按鍵路徑"],
        [
            ("建立期刊節點", "CMS ＞ 節點管理 ＞ 臺灣佛教研究中心 ＞ 中文 ＞ 首頁 ＞ 右鍵「建立」"),
            ("加入一期", "期刊節點右鍵「內容」＞「＋新增」"),
            ("加封面", "文章編輯 ＞ 圖片 ＞ 增加"),
            ("加論文目錄", "文章編輯 ＞ 內容 ＞ 增加 ＞ 表格"),
            ("上傳 PDF／Word", "檔案管理 ＞ 檔案庫管理 ＞ 點選或拖曳"),
            ("上傳封面", "檔案管理 ＞ 圖片庫管理 ＞ 點選或拖曳"),
        ],
        [2250, 7110],
    )

    doc.core_properties.title = "玄奘佛學研究網站建置操作手冊"
    doc.core_properties.subject = "玄奘大學臺灣佛教研究中心網站後台操作指引"
    doc.core_properties.author = "玄奘佛學研究網站建置專案"
    doc.core_properties.comments = "不含帳號密碼；以實際後台介面為準。"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
