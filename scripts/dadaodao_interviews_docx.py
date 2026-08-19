# -*- coding: utf-8 -*-
"""《當代的大愛道革命》口述訪談集 → 三冊設計排版 Word。

第一冊　弘誓學院與玄奘大學（學團常住、玄奘同仁與護持者）
第二冊　印順學派的學者與僧人
第三冊　社會運動與宗教對話

版式：封面頁／目次（Word TOC 欄位，開檔後按 F9 更新即得頁碼）／
每篇起新頁；三冊正文頁碼連續（第一冊 1、第二冊 189、第三冊 414 起）；
篇名 Heading 1、節 Heading 2；受訪資訊小字；
「筆者：」問句粗體、受訪者答句頭銜粗體；註釋標記 [n](#footnoteN)
轉上標小字（全註見論文附錄）；頁尾置中頁碼。

冊別分派寫在 VOLUMES，要調整直接改清單即可。
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC_DIR = ROOT / 'public/content/interviews'
OUT_DIR = ROOT / 'public/content/works'

KAI = 'DFKai-SB'
MING = 'PMingLiU'
ACCENT = RGBColor(0x27, 0x4C, 0x5E)
MUTED = RGBColor(0x66, 0x6B, 0x73)


def style_fonts(doc, name, east, ascii_font='Times New Roman', size=None, bold=None):
    try:
        style = doc.styles[name]
    except KeyError:
        return None
    style.font.name = ascii_font
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn('w:eastAsia'), east)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    return style

# 排序原則：昭慧（上下）、性廣三篇置於第一冊之首，其餘一律依訪談日期排序。
VOLUMES = [
    ('第一冊　弘誓學院與玄奘大學', 'mahaprajapati-interviews-vol1.docx', [
        '04.06 釋昭慧法師口述訪談紀錄',            # 2023.04.06
        '05.18 釋昭慧法師口述訪談紀錄(下)',         # 2023.05.18
        '04.17 釋性廣法師口述訪談紀錄',            # 2024.04.17
        '04.21 黃運喜教授口述訪談紀錄',            # 2023.04.21
        '12.05 釋明一法師口述訪談紀錄',            # 2023.12.05
        '01.16 釋圓貌法師口述訪談紀錄',            # 2024.01.16
        '01.16 釋心宇法師口述訪談紀錄',            # 2024.01.16
        '01.16 釋印悅法師口述訪談紀錄',            # 2024.01.16
        '01.17 釋心皓法師口述訪談紀錄',            # 2024.01.17
        '01.17 王彩虹居士口述訪談紀錄',            # 2024.01.17
        '02.14 陳悅萱老師口述訪談紀錄',            # 2024.02.14
        '03.26 張莉筠居士口述訪談紀錄',            # 2024.03.26
        '03.01 釋心謙法師口述訪談紀錄',            # 2025.03.01
        '03.15 釋心玄法師口述訪談紀錄',            # 2025.03.15
    ]),
    ('第二冊　印順學派的學者與僧人', 'mahaprajapati-interviews-vol2.docx', [
        '12.22 侯坤宏教授口述訪談紀錄',            # 2022.12.22
        '04.10 邱敏捷教授口述訪談紀錄',            # 2023.04.10
        '05.08 楊惠南教授口述訪談紀錄',            # 2023.05.08
        '06.05 闞正宗教授口述訪談紀錄',            # 2023.06.05
        '08.27 林建德教授口述訪談紀錄',            # 2023.08.27
        '09.13 釋清德法師口述訪談紀錄',            # 2023.09.13
        '01.23 何日生教授口述訪談紀錄',            # 2024.01.23
        '04.12 溫金柯居士口述訪談紀錄',            # 2024.04.12
        '04.18 釋宏印法師口述訪談紀錄',            # 2024.04.18
        '05.11 釋見岸法師口述訪談紀錄',            # 2024.05.11
        '06.17 釋長叡法師口述訪談紀錄',            # 2024.06.17
        '08.20 釋長慈法師口述訪談紀錄',            # 2024.08.20
        '09.19 釋寬謙法師口述訪談紀錄',            # 2024.09.19
    ]),
    ('第三冊　社會運動與宗教對話', 'mahaprajapati-interviews-vol3.docx', [
        '12.27 盧俊義牧師口述訪談紀錄',            # 2023.12.27
        '01.16 何宗勳先生口述訪談紀錄',            # 2024.01.16
        '02.02 艾琳達教授口述訪談紀錄',            # 2024.02.02
        '04.08 莊秀美女士口述訪談紀錄',            # 2024.04.08
        '05.02 黃美瑜女士游雅婷女士口述訪談紀錄',  # 2024.05.02
        '05.07 洪山川主教口述訪談紀錄',            # 2024.05.07
        '05.07 張章得先生口述訪談紀錄',            # 2024.05.07
        '05.23 詹錫奎先生口述訪談紀錄',            # 2024.05.23
        '05.28 朱增宏先生口述訪談紀錄',            # 2024.05.28
        '06.20 葉菊蘭女士口述訪談紀錄0213(最終版)', # 2024.06.20
        '09.03 林蓉芝居士口述訪談紀錄',            # 2024.09.03
    ]),
]

# 三冊以同一套正文頁碼連續編排。封面與目次仍各自獨立編頁。
BODY_STARTS = (1, 189, 414)

FN_RE = re.compile(r'\[(\d+)\]\(#footnote\d+\)')
MD_LINK_RE = re.compile(r'\[([^\]]*)\]\([^)]*\)')


def run_ea(par, text, font=MING, size=11.0, bold=False, sup=False, color=None):
    r = par.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font)
    r.font.size = Pt(size)
    r.font.bold = bold
    if sup:
        r.font.superscript = True
    if color:
        r.font.color.rgb = color
    return r


def configure_a5(section):
    """A5 print geometry: mirrored margins, modest gutter, readable text block."""
    section.page_width = Cm(14.8)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.45)
    section.gutter = Cm(0.35)
    section.header_distance = Cm(0.72)
    section.footer_distance = Cm(0.72)


def enable_mirror_margins(doc):
    settings = doc.settings.element
    if settings.find(qn('w:mirrorMargins')) is None:
        settings.append(OxmlElement('w:mirrorMargins'))


def set_page_numbering(section, fmt='decimal', start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn('w:pgNumType'))
    if pg_num is None:
        pg_num = OxmlElement('w:pgNumType')
        sect_pr.append(pg_num)
    pg_num.set(qn('w:fmt'), fmt)
    pg_num.set(qn('w:start'), str(start))


def add_page_field(section, font=MING, size=8.5):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = run_ea(p, '', font, size, color=MUTED)
    for tag, attrs, value in [
        ('w:fldChar', {'w:fldCharType': 'begin'}, None),
        ('w:instrText', {'xml:space': 'preserve'}, ' PAGE '),
        ('w:fldChar', {'w:fldCharType': 'end'}, None),
    ]:
        el = OxmlElement(tag)
        for key, attr_value in attrs.items():
            el.set(qn(key), attr_value)
        if value:
            el.text = value
        run._element.append(el)


def set_running_header(section, text):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run_ea(p, text, MING, 8.2, color=MUTED)


def add_text_with_notes(par, text, font=MING, size=11.0, bold=False):
    """[n](#footnoteN) → 上標小字 n；其餘 markdown 連結去殼。"""
    pos = 0
    for m in FN_RE.finditer(text):
        before = MD_LINK_RE.sub(r'\1', text[pos:m.start()])
        if before:
            run_ea(par, before, font, size, bold)
        run_ea(par, m.group(1), font, max(size - 3, 7), sup=True, color=RGBColor(0x25, 0x63, 0xEB))
        pos = m.end()
    rest = MD_LINK_RE.sub(r'\1', text[pos:])
    if rest:
        run_ea(par, rest, font, size, bold)


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(14)
    run_ea(p, '目　次', KAI, 16, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-1" \h \z \u')
    inner = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '（在 Word 中按 F9 更新目次即顯示各篇頁碼）'
    inner.append(t)
    fld.append(inner)
    p._p.append(fld)


def build_volume(vol_title, out_name, filenames, body_start=1):
    doc = Document()
    enable_mirror_margins(doc)
    cover_sec = doc.sections[0]
    configure_a5(cover_sec)
    cover_sec.header.is_linked_to_previous = False
    cover_sec.footer.is_linked_to_previous = False
    cover_sec.header.paragraphs[0].clear()
    cover_sec.footer.paragraphs[0].clear()

    style_fonts(doc, 'Heading 1', KAI, size=15, bold=True)
    style_fonts(doc, 'Heading 2', KAI, size=12.5, bold=True)
    style_fonts(doc, 'Heading 3', KAI, size=11.5, bold=True)
    normal = style_fonts(doc, 'Normal', MING, size=10.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.45

    # 封面
    for txt, size, bold, before in [
        ('人間佛教與印順學派', 19.5, True, 88),
        ('訪談集', 19.5, True, 5),
        (vol_title, 14.5, False, 28),
        ('張辰瑋　訪問・整理', 11.5, False, 54),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        run_ea(p, txt, KAI, size, bold, color=ACCENT if bold else None)

    toc_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_a5(toc_sec)
    toc_sec.header.is_linked_to_previous = False
    toc_sec.header.paragraphs[0].clear()
    set_page_numbering(toc_sec, 'lowerRoman', 1)
    add_page_field(toc_sec)

    add_toc(doc)

    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_a5(body_sec)
    set_page_numbering(body_sec, 'decimal', body_start)
    set_running_header(body_sec, f'《人間佛教與印順學派訪談集》　{vol_title}')
    add_page_field(body_sec)

    n_ok = 0
    for idx, fname in enumerate(filenames):
        path = SRC_DIR / f'{fname}.txt'
        if not path.exists():
            print(f'  ⚠ 缺檔跳過：{fname}')
            continue
        if n_ok:
            doc.add_page_break()
        n_ok += 1
        lines = path.read_text(encoding='utf-8').split('\n')
        first = True
        for raw in lines:
            t = raw.strip()
            if not t:
                continue
            if first:  # 篇名
                first = False
                h = doc.add_heading(level=1)
                h.paragraph_format.space_before = Pt(8)
                h.paragraph_format.space_after = Pt(11)
                h.paragraph_format.keep_with_next = True
                h.paragraph_format.widow_control = True
                add_text_with_notes(h, t, KAI, 15, bold=True)
                for run in h.runs:
                    run.font.color.rgb = ACCENT
                continue
            if re.match(r'^(受訪者|訪問者|訪問時間|訪問地點|訪談時間|訪談地點|採訪者|採訪時間|採訪地點)：', t):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                add_text_with_notes(p, t, MING, 8.5)
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                continue
            if re.match(r'^[一二三四五六七八九十]+、', t) or t in ('結語', '前言'):
                h = doc.add_heading(level=2)
                h.paragraph_format.space_before = Pt(11)
                h.paragraph_format.space_after = Pt(5)
                h.paragraph_format.keep_with_next = True
                add_text_with_notes(h, t, KAI, 12.5, bold=True)
                for run in h.runs:
                    run.font.color.rgb = ACCENT
                continue
            if re.match(r'^（[一二三四五六七八九十]+）', t):
                h = doc.add_heading(level=3)
                h.paragraph_format.keep_with_next = True
                add_text_with_notes(h, t, KAI, 11.5, bold=True)
                continue
            m = re.match(r'^筆者：(.*)$', t)
            if m:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.45
                p.paragraph_format.widow_control = True
                run_ea(p, '問：', KAI, 10.5, bold=True, color=ACCENT)
                add_text_with_notes(p, m.group(1), MING, 10.5, bold=True)
                continue
            m = re.match(r'^([一-鿿]{1,6}(?:法師|教授|主教|和尚|居士|博士|老師|牧師|女士|先生))：(.+)$', t)
            if m:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.45
                p.paragraph_format.widow_control = True
                run_ea(p, f'{m.group(1)}：', KAI, 10.5, bold=True, color=ACCENT)
                add_text_with_notes(p, m.group(2), MING, 10.5)
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.45
            p.paragraph_format.widow_control = True
            add_text_with_notes(p, t, MING, 10.5)

    out = OUT_DIR / out_name
    doc.save(str(out))
    print(f'{vol_title}：{n_ok} 篇 → {out.name} ({out.stat().st_size:,} bytes)')


if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    total = sum(len(v[2]) for v in VOLUMES)
    print(f'三冊共 {total} 篇')
    for vi, (vol_title, out_name, files) in enumerate(VOLUMES):
        build_volume(vol_title, out_name, files, BODY_STARTS[vi])
