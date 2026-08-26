#!/usr/bin/env python3
"""Lay out the two Latin volumes as JIS-B5 DOCX.

The page geometry, type ladder, table rhythm and palette are the Hebrew
reader's, imported rather than re-specified, so the three readers in this series
sit on a shelf as one set. What differs is what a Latin page has to carry: no
right-to-left runs, no pointing, but a vocabulary table whose first column is a
full set of principal parts rather than a single form, and a reading column that
alternates between verse-numbered scripture and the versicle-and-response of the
Mass.

Each lesson prints the same four things in the same order -- twenty words, two
memory units, the reading, and the reading's Chinese -- because a reader that
reorders itself between lessons cannot be used as a reference.

Nothing here is generated. Every string comes from the frozen data masters, and
where a master has a gap the page says so rather than leaving a silent blank:
a reading still awaiting its Chinese prints 〔中譯待補〕, which is a thing the
owner can see and act on.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Mm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
import latin_source_texts as L  # noqa: E402
import build_hebrew_full_reader as H  # noqa: E402

ROOT = Path(__file__).resolve().parents[1]
CACHE = ROOT / "output" / "source-cache" / "original-readers" / "latin-full"
OUT_DIR = ROOT / "output" / "original-readers"

VOCABULARY = ROOT / "data" / "originalReaders" / "vocabulary" / "latin-2000.json"
APPENDICES = ROOT / "data" / "originalReaders" / "vocabulary" / "latin-appendices.json"
SCRIPTURE = CACHE / "scripture-plan.json"
SIGAO = CACHE / "sigao-zh.json"
CHURCH = CACHE / "church-plan.json"
LITURGY = CACHE / "liturgy.json"
READINGS_ZH = CACHE / "readings-zh.json"
MEMORY = CACHE / "memory-units.json"

LITURGY_NOTE = "禮儀經文不作機器翻譯；中譯應採用教會通行本文，付印前由人補上"

FONT_LA = "Noto Serif"
LATIN_PT = 11.0
GLOSS_PT = 9.6

VOLUMES = {
    "上冊": {
        "subtitle": "武加大譯本",
        "blurb": "十篇禮儀短經，四十章完整武加大經文，中文並列思高譯本。",
        "file": "latin-original-reader-vol1.docx",
        "appendix": "upper",
    },
    "下冊": {
        "subtitle": "從教父到教廷",
        "blurb": "五十篇教父、中世紀與教廷文獻，終卷為常年期主日彌撒經文全文。",
        "file": "latin-original-reader-vol2.docx",
        "appendix": "lower",
    },
}

COLOPHON = [
    ("拉丁文本", "武加大譯本用 Clementine Vulgate（eBible.org latVUC 轉錄，公有領域）；"
                 "教父與中世紀文本取自 The Latin Library；教廷文獻取自本專案既有拉丁文檔；"
                 "彌撒經文取自 Collins《A Primer of Ecclesiastical Latin》讀本部分所印之現行彌撒常規。"),
    ("中文", "聖經章節用思高譯本（思高聖經學會）。其餘篇章凡標「自譯」者為本讀本研讀用譯文，"
             "非教會核准之禮儀譯本；中文彌撒經文另有《感恩祭典》，本書不取代之。"),
    ("詞彙", "上冊一千詞依 Collins《A Primer of Ecclesiastical Latin》原書順序；"
             "下冊一千詞依教父／中世紀與近現代教廷語料詞頻，與上冊互斥。"
             "詞形主要部分取自 Whitaker's WORDS。"),
    ("發音", "全書採羅馬式教會發音。古典重建音為另一軌，本書不混用。"),
    ("授權", "本書為私人研讀用途，非賣品。製作已取得口頭同意；"
             "所引各版本之著作權仍屬原權利人，不得再散布。"),
]


def load(path: Path, default=None):
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


# --------------------------------------------------------------------------
# page furniture
# --------------------------------------------------------------------------

def heading(document, text: str, size: float, *, color=None, space_before=10,
            space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    H.add_mixed_script_text(paragraph, text, H.FONT_ZH, size, bold=True,
                            color=color or H.ACCENT_DARK)
    return paragraph


def body(document, text: str, size=H.BODY_SIZE_PT, *, font=H.FONT_ZH, color=H.INK,
         italic=False, space_after=4, indent_mm=0.0):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = H.BODY_LINE_MULTIPLE
    if indent_mm:
        paragraph.paragraph_format.left_indent = Mm(indent_mm)
    H.add_mixed_script_text(paragraph, text, font, size, italic=italic, color=color)
    return paragraph


def page_break(document):
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def title_page(document, volume: str, spec: dict, counts: str):
    for _ in range(4):
        document.add_paragraph()
    heading(document, "教會拉丁文原文讀本", H.TITLE_SIZE_PT,
            align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    heading(document, f"{volume}　{spec['subtitle']}", H.H1_SIZE_PT,
            align=WD_ALIGN_PARAGRAPH.CENTER, color=H.ACCENT, space_after=18)
    para = body(document, spec["blurb"], H.BODY_SIZE_PT, color=H.MUTED)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para = body(document, counts, H.CAPTION_PT, color=H.MUTED)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_break(document)

    heading(document, "凡例", H.H1_SIZE_PT)
    for label, text in COLOPHON:
        heading(document, label, H.H3_SIZE_PT, space_before=8, space_after=2)
        body(document, text, H.TRANSLATION_PT, color=H.MUTED)
    page_break(document)


def vocabulary_table(document, rows: list[dict]):
    heading(document, "本課詞彙", H.H3_SIZE_PT, space_before=6, space_after=4)
    table = document.add_table(rows=1, cols=3)
    widths = [H.USABLE_WIDTH_MM * 0.46, H.USABLE_WIDTH_MM * 0.14, H.USABLE_WIDTH_MM * 0.40]
    H.set_table_geometry(table, widths)
    H.set_borders(table)
    header = table.rows[0]
    for cell, label in zip(header.cells, ("拉丁文", "詞類", "繁體中文")):
        H.shade(cell, H.PALE)
        H.set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        H.add_mixed_script_text(paragraph, label, H.FONT_ZH, H.LABEL_PT, bold=True,
                                color=H.ACCENT_DARK)
    H.set_repeat_header(header)
    for entry in rows:
        row = table.add_row()
        # Keep a vocabulary row whole.  A row that splits across a page break
        # leaves what looks like an empty first row at the top of the next page,
        # which reads as a missing word rather than as a continuation.
        H.prevent_row_split(row)
        cells = row.cells
        H.set_cell_margins(cells[0])
        H.add_mixed_script_text(cells[0].paragraphs[0], entry.get("forms") or entry["headword"],
                                FONT_LA, H.TABLE_SIZE_PT)
        H.set_cell_margins(cells[1])
        H.add_mixed_script_text(cells[1].paragraphs[0], short_pos(entry), H.FONT_ZH,
                                H.LABEL_PT, color=H.MUTED)
        H.set_cell_margins(cells[2])
        H.add_mixed_script_text(cells[2].paragraphs[0], entry.get("glossZh") or "〔待補〕",
                                H.FONT_ZH, H.TABLE_SIZE_PT)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


POS_ZH = {"N": "名", "V": "動", "ADJ": "形", "ADV": "副", "PREP": "介", "CONJ": "連",
          "PRON": "代", "NUM": "數", "INTERJ": "嘆",
          "NOUN": "名", "VERB": "動", "PROPN": "名", "ADP": "介",
          "CCONJ": "連", "SCONJ": "連", "DET": "限", "AUX": "動", "INTJ": "嘆"}

GRAM_HINTS = (
    ("prep", "介"), ("conj", "連"), ("adv", "副"), ("pron", "代"),
    ("num", "數"), ("interj", "嘆"), ("indecl", "不變"),
)


def short_pos(entry: dict) -> str:
    """Say what part of speech this is, reading the dictionary line if need be.

    Collins does not label his nouns and verbs: the gender abbreviation at the
    end of a noun entry and the four principal parts of a verb entry *are* the
    labels.  Taking the label only from an explicit field leaves the column
    empty for most of the book, which is what the first print run did.
    """
    for key in ("gram", "pos"):
        value = (entry.get(key) or "").strip()
        if value in POS_ZH:
            return POS_ZH[value]
    gram = (entry.get("gram") or "").lower()
    for needle, label in GRAM_HINTS:
        if needle in gram:
            return label
    forms = (entry.get("forms") or "").strip()
    parts = [p.strip() for p in forms.split(",")]
    if re.search(r"(^|[, ])(m|f|n|c)\.$", forms):
        return "名"
    if re.search(r"-(a|ae), -(um|a)$|-is, -e$|, -a, -um$", forms):
        return "形"
    if len(parts) >= 4 or re.search(r"(are|ēre|ere|īre|ire)$", parts[1] if len(parts) > 1 else ""):
        return "動"
    if len(parts) == 2 and parts[1]:
        return "名"
    return ""


def memory_block(document, units: list[dict]):
    if not units:
        return
    heading(document, "記憶單元", H.H3_SIZE_PT, space_before=6, space_after=3)
    for unit in units:
        body(document, unit["text"], LATIN_PT, font=FONT_LA, space_after=1, indent_mm=4)
        zh = unit.get("zh") or ""
        if zh and zh != "reading-has-chinese":
            body(document, zh, H.TRANSLATION_PT, color=H.MUTED, space_after=3, indent_mm=4)
        else:
            body(document, f"〔{unit.get('ref', '')}〕", H.CAPTION_PT, color=H.MUTED,
                 space_after=3, indent_mm=4)


def reading_block(document, title: str, pairs: list[tuple[str, str]], note: str = ""):
    heading(document, f"讀本　{title}", H.H3_SIZE_PT, space_before=8, space_after=3)
    if note:
        body(document, note, H.CAPTION_PT, color=H.MUTED, space_after=4)
    for latin, chinese in pairs:
        body(document, latin, LATIN_PT, font=FONT_LA, space_after=1)
        body(document, chinese or "〔中譯待補〕", H.TRANSLATION_PT, color=H.MUTED,
             space_after=5)


# --------------------------------------------------------------------------
# content assembly
# --------------------------------------------------------------------------

def upper_readings() -> dict[int, dict]:
    plan = load(SCRIPTURE)
    chinese = load(SIGAO, {"chapters": []})
    translated = load(READINGS_ZH, {"units": {}})
    liturgy = {row["id"]: row for row in load(LITURGY, {"formulas": []})["formulas"]}
    zh_by_chapter = {(c["book"], c["latinChapter"]): c for c in chinese["chapters"]}
    verses = L.vulgate_chapters()

    out: dict[int, dict] = {}
    for row in plan["chapters"]:
        if row["kind"] == "liturgy":
            source = liturgy.get(row["id"], {})
            unit = translated["units"].get(f"formula:{row['id']}", {})
            zh_lines = [z for segment in unit.get("segments", []) for z in segment["zh"]]
            pairs = list(zip(source.get("lines", []), zh_lines + [""] * len(source.get("lines", []))))
            # The liturgical Chinese is deliberately absent, not merely late:
            # a machine rendering of a formula the congregation knows by heart
            # is an error the label 自譯 does not cover.
            note = "　".join(x for x in (row.get("note"), LITURGY_NOTE) if x)
            out[row["lesson"]] = {"title": f"{row['title']}　{row['latinTitle']}",
                                  "pairs": pairs, "note": note}
            continue
        chapter_zh = zh_by_chapter.get((row["book"], row["chapter"]))
        zh_by_verse = {v["verse"]: v["text"] for v in chapter_zh["verses"]} if chapter_zh else {}
        pairs = [(f"{number}　{text}", zh_by_verse.get(number, ""))
                 for number, text in sorted(verses[(row["book"], row["chapter"])].items())]
        note = row.get("note") or ""
        if chapter_zh and chapter_zh.get("alignmentNote"):
            note = (note + "　" if note else "") + chapter_zh["alignmentNote"]
        out[row["lesson"]] = {"title": row["title"], "pairs": pairs, "note": note}
    return out


SECTION_NUMBER = re.compile(r"^\s*(\d{1,3})[.、]")


def chinese_by_section(path: str) -> dict[int, str]:
    """Index a published translation by the section numbers it prints.

    Pairing the two sides by paragraph index is what this replaced, and it was
    wrong every time: Sacrosanctum Concilium has 362 Latin paragraphs against
    11 Chinese ones, so paragraph five of each is five different places in the
    document.  Where both sides number their sections, the number is the join.
    """
    latin = ROOT / path
    chinese = latin.with_name(latin.name.replace("-latin.txt", "-chinese.txt"))
    if not chinese.exists():
        return {}
    raw = chinese.read_text(encoding="utf-8", errors="replace")
    body = re.sub(r"^#.*$", "", raw, flags=re.M)
    sections: dict[int, list[str]] = {}
    current = 0
    for line in body.splitlines():
        match = SECTION_NUMBER.match(line)
        if match:
            current = int(match.group(1))
        if current and line.strip():
            sections.setdefault(current, []).append(line.strip())
    return {number: " ".join(rows) for number, rows in sections.items()}


def lower_readings() -> dict[int, dict]:
    plan = load(CHURCH)
    translated = load(READINGS_ZH, {"units": {}})
    out: dict[int, dict] = {}
    for row in plan["readings"]:
        key = f"reading:{row['sourceRef']}"
        unit = translated["units"].get(key)
        if unit:
            pairs = [(segment["latin"][0], segment["zh"][0])
                     for segment in unit["segments"]]
            note = f"{row['excerptRule']}；{unit['translationNote']}"
        else:
            # Cut with the same rule the plan measured: whole divisions of the
            # work, never part of one.  Re-splitting on blank lines here instead
            # would print sixteen thousand words of Vincent of Lerins, because
            # several Latin Library files contain no blank line at all.
            import build_latin_church_plan as plan_module
            latin_text = (ROOT / row["sourcePath"]).read_text(encoding="utf-8", errors="replace")
            if row.get("section"):
                latin_text = plan_module.section(latin_text, tuple(row["section"]))
            if row["extent"] == "excerpt":
                latin_text, _, _ = plan_module.complete_unit(latin_text)
            paragraphs = [re.sub(r"\s+", " ", part).strip()
                          for part in latin_text.split(chr(10) * 2) if part.strip()]
            chinese = (chinese_by_section(row["sourcePath"])
                       if row["chineseParallel"] == "repo-aligned-by-number" else {})
            pairs = []
            for paragraph in paragraphs:
                match = SECTION_NUMBER.match(paragraph)
                zh = chinese.get(int(match.group(1)), "") if match else ""
                pairs.append((paragraph, zh))
            note = row["excerptRule"]
            if row["chineseSource"] in {"denzinger-excerpts", "placeholder",
                                        "full-translation-unnumbered"}:
                note += f"；既有中文檔為 {row['chineseSource']}，無法逐段並排，中譯另行自譯"

        out[row["lesson"]] = {
            "title": f"{row['title']}　{row['latinTitle']}", "pairs": pairs, "note": note,
        }
    return out


def appendix_section(document, tables: dict):
    page_break(document)
    heading(document, "附錄", H.H1_SIZE_PT)
    for table in tables.values():
        entries = table["entries"]
        heading(document, f"{table['title']}（{len(entries)} 條）", H.H2_SIZE_PT,
                space_before=10, space_after=4)
        grouped: dict[str, list[dict]] = {}
        for entry in entries:
            grouped.setdefault(entry.get("group", ""), []).append(entry)
        for group, rows in grouped.items():
            if group:
                heading(document, group, H.H3_SIZE_PT, space_before=6, space_after=2)
            for row in rows[:200]:
                latin = row.get("forms") or row.get("headword", "")
                zh = row.get("zh") or row.get("glossZh") or row.get("glossEn") or ""
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(1)
                H.add_mixed_script_text(paragraph, latin + "　", FONT_LA, H.TABLE_SIZE_PT)
                H.add_mixed_script_text(paragraph, zh, H.FONT_ZH, H.TABLE_SIZE_PT,
                                        color=H.MUTED)
            if len(rows) > 200:
                body(document, f"（本組尚有 {len(rows) - 200} 條，見資料檔）",
                     H.CAPTION_PT, color=H.MUTED)


def relabel(document, volume: str, spec: dict) -> None:
    """Put this book's name in the running head.

    The layout is imported from the Hebrew reader, and so is its running header;
    left alone, every page of the Latin volumes says 聖經希伯來文原文讀本.
    """
    running = f"教會拉丁文原文讀本　{volume}　{spec['subtitle']}"
    for section in document.sections:
        for part in (section.header, section.first_page_header):
            for paragraph in part.paragraphs:
                for run in paragraph.runs:
                    if "希伯來" in run.text:
                        run.text = running
    document.core_properties.title = f"教會拉丁文原文讀本：{volume}　{spec['subtitle']}"
    document.core_properties.subject = spec["blurb"]


def build(volume: str) -> Path:
    spec = VOLUMES[volume]
    vocabulary = load(VOCABULARY)["entries"]
    memory = load(MEMORY, {"上冊": [], "下冊": []})
    appendices = load(APPENDICES, {})
    readings = upper_readings() if volume == "上冊" else lower_readings()

    per_lesson: dict[int, list[dict]] = {}
    for entry in vocabulary:
        if entry["volume"] == volume:
            per_lesson.setdefault(entry["lesson"], []).append(entry)
    memory_by_lesson: dict[int, list[dict]] = {}
    for unit in memory.get(volume, []):
        memory_by_lesson.setdefault(unit["lesson"], []).append(unit)

    document = Document()
    H.configure(document)
    relabel(document, volume, spec)
    words = sum(len(per_lesson.get(n, [])) for n in range(1, 51))
    reading_words = sum(len(L.words(latin)) for row in readings.values()
                        for latin, _ in row["pairs"])
    title_page(document, volume, spec,
               f"五十課．{words} 詞．讀本 {reading_words:,} 詞")

    for lesson in range(1, 51):
        reading = readings.get(lesson, {"title": "", "pairs": [], "note": ""})
        heading(document, f"第 {lesson} 課　　{reading['title']}", H.H1_SIZE_PT,
                space_before=0, space_after=6)
        vocabulary_table(document, per_lesson.get(lesson, []))
        memory_block(document, memory_by_lesson.get(lesson, []))
        if reading["pairs"]:
            reading_block(document, reading["title"], reading["pairs"], reading["note"])
        page_break(document)

    appendix_section(document, appendices.get(spec["appendix"], {}))

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / spec["file"]
    document.save(path)
    return path


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--volume", choices=("上冊", "下冊", "both"), default="both")
    args = ap.parse_args()
    targets = ["上冊", "下冊"] if args.volume == "both" else [args.volume]
    for volume in targets:
        path = build(volume)
        print(f"{volume} -> {path.relative_to(ROOT)}  {path.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
