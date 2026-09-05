"""把 /works 學位論文計畫書的 markdown 轉成可直接送件的 .docx。

用法：
    python scripts/build_proposal_docx.py public/content/works/hcu-phd-proposal.md

版式（比照系上送件慣例）：
    A4、頁邊 2.54cm；內文 12pt 新細明體 + Times New Roman、行距 1.5、首行縮排兩字
    #    → 主標 18pt 置中粗體      ##   → 章節標 14pt 粗體
    ###  → 子標 13pt 粗體          #### → 小標 12pt 粗體
    *斜體行* → 置中斜體（英文題名）  **粗體** → 行內粗體
    ---  → 分隔線（略過不印）
    「參考書目」／「徵引書目」之後的段落自動改成懸掛縮排（書目體例）

腳註（國史館體例要求隨頁附註）：
    正文寫 `文字。[^1]`，檔案末尾另起 `[^1]: 註文`。轉檔時變成 Word 真正的
    footnote，號碼由 Word 自動編、排在當頁下緣。實作見 docx_footnotes.py。
    🚨 註號位置照國史館體例「置於標點符號之後」，所以 markdown 也寫在句號後面。

表格：
    連續的 `| a | b |` 行轉成 Word 表格（吃掉 `|---|` 那一行）。首列作表頭。
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx_footnotes import Footnotes  # noqa: E402

EN_FONT = "Times New Roman"
# 送件體例以使用者的《張辰瑋，論文研究計畫》為準（2026-09 量測）：
#   標楷體、行距 1.5、左右邊界 3.17cm、內文**不縮排**、全篇沒有項目符號；
#   封面 18pt（校系／姓名／日期）、15pt（英文機構四行）、20pt（中英題名，不粗體）；
#   章標「一、」14pt 粗體，節標「（一）」粗體但不放大。
CJK_FONT = "標楷體"
STYLE = {"cjk": "標楷體", "side": Cm(3.17), "indent": None}

# 這幾節各自另起一頁（送件文件的基本體例）
PAGE_BREAK_BEFORE = ("Abstract", "目錄", "圖表目錄", "前言", "參考書目", "徵引書目", "附錄")


def add_run(par, text, *, bold=False, italic=False, size=12):
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = EN_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    return run


INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def add_inline(par, text, *, size=12, bold=False, italic=False):
    """處理 **粗體** 與 *斜體*（西文期刊名、書名要斜體，國史館體例亦然）。
    🚨 斜體一定要一起處理：只認粗體的話，*Journal of Buddhist Ethics* 會把星號
       原樣印出來——表格欄位裡尤其看不出來，因為那裡不會換行。"""
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            add_run(par, chunk[2:-2], bold=True, italic=italic, size=size)
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            add_run(par, chunk[1:-1], bold=bold, italic=True, size=size)
        else:
            add_run(par, chunk, bold=bold, italic=italic, size=size)


def new_par(doc, *, align=None, space_after=6, first_indent=None, hanging=None):
    par = doc.add_paragraph()
    fmt = par.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(space_after)
    if align is not None:
        par.alignment = align
    if first_indent is not None:
        fmt.first_line_indent = first_indent
    if hanging is not None:
        fmt.left_indent = hanging
        fmt.first_line_indent = -hanging
    return par


FN_DEF = re.compile(r"^\[\^([^\]]+)\]:\s*(.+)$")
FN_REF = re.compile(r"\[\^([^\]]+)\]")


def collect_footnotes(lines):
    """把檔尾的 `[^1]: 註文` 收成字典，並回傳去掉這些行的正文。"""
    notes, body = {}, []
    for ln in lines:
        m = FN_DEF.match(ln.strip())
        if m:
            notes[m.group(1)] = m.group(2).strip()
        else:
            body.append(ln)
    return notes, body


def add_table(doc, rows):
    """rows 是每列的欄位串列；首列作表頭。"""
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            par = t.cell(i, j).paragraphs[0]
            par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            par.paragraph_format.space_after = Pt(2)
            if i == 0:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(par, cell, bold=(i == 0))
    return t


def build(md_path: Path, out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_height, section.page_width = Cm(29.7), Cm(21.0)
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = STYLE["side"]

    normal = doc.styles["Normal"]
    normal.font.name = EN_FONT
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), STYLE["cjk"])

    notes, lines = collect_footnotes(md_path.read_text(encoding="utf8").splitlines())
    fn = Footnotes(doc)

    in_biblio = False
    on_cover = True  # 「摘要」之前都算封面：一律置中、不縮排
    # 沒有「摘要」的短件（如指導規劃）用這行標記封面到哪裡結束
    break_next = False
    seen_title = False
    i = 0
    while i < len(lines):
        raw = lines[i]
        i += 1
        line = raw.strip()
        if not line or re.fullmatch(r"-{3,}", line):
            continue

        if line == "<!-- 封面結束 -->":
            on_cover, break_next = False, True
            continue

        if line.startswith("|") and line.endswith("|"):
            block = []
            j = i - 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    block.append(cells)
                j += 1
            i = j
            if block:
                add_table(doc, block)
            continue

        head = re.match(r"^(#{1,6})\s+(.*)$", line)
        if head:
            level, text = len(head.group(1)), head.group(2)
            page_break = False
            if level == 1:
                seen_title = True
            if text.strip() == "摘要":
                on_cover = False
                page_break = True             # 封面獨立一頁
            elif break_next:
                page_break, break_next = True, False
            elif any(text.startswith(k) for k in PAGE_BREAK_BEFORE):
                page_break = True
            if "參考書目" in text or "徵引書目" in text:
                in_biblio = True
            elif level <= 2 and in_biblio and "附錄" in text:
                in_biblio = False
            if on_cover:
                size, bold = 20, False          # 中英題名
            else:
                size = {2: 15 if text in ("摘要", "Abstract") else 14, 3: 12}.get(level, 12)
                bold = True
            align = WD_ALIGN_PARAGRAPH.CENTER if (level == 1 or on_cover) else None
            par = new_par(doc, align=align, space_after=10)
            par.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
            # 直接設在標題段上，而不是插一個只帶分頁符的空段——後者會留下空行、
            # 在標題剛好落在頁首時還會多出一整頁空白。
            par.paragraph_format.page_break_before = page_break
            add_inline(par, text, size=size, bold=bold)
            continue

        # 整行斜體（英文題名）
        whole_italic = re.fullmatch(r"\*([^*].*[^*])\*", line)
        if whole_italic:
            par = new_par(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_run(par, whole_italic.group(1), italic=True, size=20 if on_cover else 12)
            continue

        if on_cover:
            par = new_par(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
            # 英文機構那四行 15pt，其餘（校系、姓名、指導教授、日期）18pt。
            # 判準是「在中文題名之前的純 ASCII 行」——姓名的英譯在題名之後，是 18pt。
            add_run(par, line, size=15 if (not seen_title and line.isascii()) else 18)
            continue
        if in_biblio:
            par = new_par(doc, space_after=4, hanging=Cm(0.85))
        else:
            par = new_par(doc, first_indent=STYLE["indent"])
        # 註號要變成真正的 footnote，所以文字必須按 [^n] 切開逐段加
        for k, chunk in enumerate(FN_REF.split(line)):
            if k % 2 == 0:
                if chunk:
                    add_inline(par, chunk)
            elif chunk in notes:
                fn.add(par, notes[chunk])
            else:
                raise SystemExit(f"🚨 正文引用了 [^{chunk}] 但檔尾沒有對應的註文")

    unused = set(notes) - {n for n in notes if f"[^{n}]" in "".join(lines)}
    if unused:
        print(f"⚠ 有註文沒被引用：{sorted(unused)}")
    fn.save()                      # 一定要在 doc.save() 之前
    doc.save(out_path)
    print(f"{out_path}  ({out_path.stat().st_size // 1024} KB)　腳註 {len(fn.items)} 條")


if __name__ == "__main__":
    src = Path(sys.argv[1])
    build(src, src.with_suffix(".docx"))
