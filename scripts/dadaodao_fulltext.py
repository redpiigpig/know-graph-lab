#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
《當代的大愛道革命》研究資料全文轉錄。
逐檔抽文字（docx/xlsx/有文字層 PDF，免 API）或 OCR（掃描 PDF／JPG／PNG）。
引擎：Gemini（4 key 輪流）為主，全部 429／配額用罄 → Sonnet（claude-sonnet-4-6，OAuth）救援。
全文存 R2 `dadaodao-fulltext/<相對路徑>.txt`；冪等（已存在跳過），可重跑接續。

  python -X utf8 scripts/dadaodao_fulltext.py [--limit N] [--only ext1,ext2] [--ocr-only] [--no-ocr]

階段建議：先 `--no-ocr`（秒抽 docx/xlsx/文字 PDF），再不帶旗標跑完掃描 OCR（過夜）。
"""
import argparse
import base64
import io
import json
import os
import sys
import time
from pathlib import Path

import boto3
import requests

ROOT_DIR = Path(__file__).resolve().parents[1]
ENV = {}
for _l in (ROOT_DIR / ".env").read_text(encoding="utf-8").splitlines():
    if _l and not _l.startswith("#") and "=" in _l:
        _k, _v = _l.split("=", 1)
        ENV[_k.strip()] = _v.strip().strip('"').strip("'")

SRC_ROOT = Path(r"G:\我的雲端硬碟\公事\國北教\碩士論文\論文資料")
R2_BUCKET = ENV["R2_BUCKET"]
TEXT_PREFIX = "dadaodao-fulltext"
GEMINI_MODEL = "gemini-2.5-flash"
SONNET_MODEL = "claude-sonnet-4-6"
OCR_PROMPT = (
    "請完整且原樣轉錄這份文件中的所有文字內容，保留原本的段落、標題與註腳結構。"
    "只輸出文字本身，不要翻譯、不要摘要、不要加任何說明或標記。"
)
MIN_TEXT_PER_PAGE = 60  # PDF 平均每頁字元數低於此 → 視為掃描檔，需 OCR

# ── clients ───────────────────────────────────────────────────────────────
s3 = boto3.client(
    "s3", region_name="auto", endpoint_url=ENV["R2_ENDPOINT"],
    aws_access_key_id=ENV["R2_ACCESS_KEY"], aws_secret_access_key=ENV["R2_SECRET_KEY"],
)

def _gemini_keys():
    keys, seen = [], set()
    for base in ("Gemini_API_Key", "GEMINI_API_KEY", "GOOGLE_API_KEY"):
        for n in [""] + [f"_{i}" for i in range(1, 11)]:
            v = ENV.get(f"{base}{n}")
            if v:
                for piece in v.split(","):
                    k = piece.strip()
                    if k and k not in seen:
                        seen.add(k); keys.append(k)
    return keys

GEMINI_KEYS = _gemini_keys()
_gem_idx = 0  # round-robin start

_CC_OAUTH_REFRESH_URL = "https://console.anthropic.com/v1/oauth/token"
_CC_OAUTH_CLIENT_ID = "9d1c250a-e61b-44d9-88ed-5944d1962f5e"

def _cred_path():
    return Path(os.environ.get("USERPROFILE", os.environ.get("HOME", ""))) / ".claude" / ".credentials.json"

def _refresh_oauth():
    cred = _cred_path()
    if not cred.exists():
        return None
    data = json.loads(cred.read_text(encoding="utf-8"))
    oauth = data.get("claudeAiOauth", {})
    rt = oauth.get("refreshToken")
    if not rt:
        return None
    r = requests.post(_CC_OAUTH_REFRESH_URL, headers={"Content-Type": "application/json", "User-Agent": "anthropic"},
                      json={"grant_type": "refresh_token", "refresh_token": rt, "client_id": _CC_OAUTH_CLIENT_ID}, timeout=30)
    if not r.ok:
        return None
    b = r.json()
    oauth["accessToken"] = b["access_token"]
    oauth["refreshToken"] = b.get("refresh_token") or rt
    if b.get("expires_in"):
        oauth["expiresAt"] = int(time.time() * 1000) + int(b["expires_in"]) * 1000
    data["claudeAiOauth"] = oauth
    cred.write_text(json.dumps(data, indent=2), encoding="utf-8")
    return oauth["accessToken"]

_anthropic_client = None
def anthropic_client():
    global _anthropic_client
    if _anthropic_client is not None:
        return _anthropic_client
    import anthropic
    api_key = ENV.get("ANTHROPIC_API_KEY") or os.environ.get("ANTHROPIC_API_KEY")
    if api_key:
        _anthropic_client = anthropic.Anthropic(api_key=api_key, timeout=600.0, max_retries=2)
        return _anthropic_client
    cred = _cred_path()
    creds = json.loads(cred.read_text(encoding="utf-8"))
    oauth = creds.get("claudeAiOauth", {})
    token = oauth.get("accessToken", "")
    if oauth.get("expiresAt", 0) <= int(time.time() * 1000) + 300_000:
        token = _refresh_oauth() or token
    _anthropic_client = anthropic.Anthropic(auth_token=token, timeout=600.0, max_retries=2)
    return _anthropic_client

# ── helpers ───────────────────────────────────────────────────────────────
def r2_exists(key):
    try:
        s3.head_object(Bucket=R2_BUCKET, Key=key)
        return True
    except Exception:
        return False

def r2_put_text(key, text):
    s3.put_object(Bucket=R2_BUCKET, Key=key, Body=text.encode("utf-8"), ContentType="text/plain; charset=utf-8")

class RateLimited(Exception):
    pass

def _is_quota(err):
    s = str(err).lower()
    return any(t in s for t in ("429", "quota", "rate limit", "resource_exhausted", "exhausted", "503", "overloaded"))

# ── extractors (no API) ─────────────────────────────────────────────────────
def extract_docx(path):
    import docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for tbl in d.tables:
        for row in tbl.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(x for x in parts if x is not None).strip()

def extract_xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"# {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                out.append("\t".join(cells))
    return "\n".join(out).strip()

def pdf_text_layer(path):
    """Return extracted text if the PDF has a real text layer, else None."""
    import fitz
    doc = fitz.open(str(path))
    n = doc.page_count
    txt = []
    for pg in doc:
        txt.append(pg.get_text())
    doc.close()
    full = "\n".join(txt).strip()
    if n and len(full) >= MIN_TEXT_PER_PAGE * n:
        return full
    return None

# ── OCR engines ─────────────────────────────────────────────────────────────
def gemini_ocr(path, mime):
    """Try each Gemini key once; raise RateLimited if all are quota-limited."""
    global _gem_idx
    from google import genai
    from google.genai import types
    last = None
    for off in range(len(GEMINI_KEYS)):
        idx = (_gem_idx + off) % len(GEMINI_KEYS)
        try:
            client = genai.Client(api_key=GEMINI_KEYS[idx])
            data = path.read_bytes()
            resp = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=[types.Part.from_bytes(data=data, mime_type=mime), OCR_PROMPT],
            )
            _gem_idx = (idx + 1) % len(GEMINI_KEYS)
            return (resp.text or "").strip()
        except Exception as e:
            last = e
            if _is_quota(e):
                continue
            raise
    raise RateLimited(f"all gemini keys limited: {last}")

def sonnet_ocr(path, mime):
    client = anthropic_client()
    b64 = base64.standard_b64encode(path.read_bytes()).decode()
    block_type = "document" if mime == "application/pdf" else "image"
    src = {"type": "base64", "media_type": mime, "data": b64}
    msg = client.messages.create(
        model=SONNET_MODEL, max_tokens=64000,
        messages=[{"role": "user", "content": [
            {"type": block_type, "source": src},
            {"type": "text", "text": OCR_PROMPT},
        ]}],
    )
    return "".join(b.text for b in msg.content if getattr(b, "type", "") == "text").strip()

class BothLimited(Exception):
    pass

def _sonnet_is_429(e):
    return e.__class__.__name__ == "RateLimitError" or "429" in str(e) or "rate_limit" in str(e).lower()

PACE = 1.5  # seconds between Gemini calls (set from --pace)

def ocr_file(path, mime, pace=None):
    pace = PACE if pace is None else pace
    """Gemini→Sonnet；兩者皆 429 時退避重試，仍不行則 raise BothLimited。"""
    for attempt in range(4):
        try:
            txt = gemini_ocr(path, mime)
            time.sleep(pace)
            return txt, "gemini"
        except RateLimited:
            pass
        try:
            return sonnet_ocr(path, mime), "sonnet"
        except Exception as e:
            if not _sonnet_is_429(e):
                raise
        back = 20 * (attempt + 1)
        print(f"    …both engines limited, backoff {back}s (attempt {attempt+1}/4)", flush=True)
        time.sleep(back)
    raise BothLimited("gemini + sonnet both rate-limited after retries")

# ── per-file dispatch ───────────────────────────────────────────────────────
IMG_MIME = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png"}

def process(path, no_ocr, ocr_only):
    ext = path.suffix.lower()
    if ext == ".docx":
        return (None if ocr_only else extract_docx(path)), "docx"
    if ext == ".xlsx":
        return (None if ocr_only else extract_xlsx(path)), "xlsx"
    if ext == ".pdf":
        if not ocr_only:
            t = pdf_text_layer(path)
            if t:
                return t, "pdf-text"
        if no_ocr:
            return None, "pdf-scan(skip)"
        txt, eng = ocr_file(path, "application/pdf")
        return txt, f"pdf-ocr:{eng}"
    if ext in IMG_MIME:
        if no_ocr:
            return None, "img(skip)"
        txt, eng = ocr_file(path, IMG_MIME[ext])
        return txt, f"img-ocr:{eng}"
    return None, f"unsupported:{ext}"

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--limit", type=int, default=0)
    ap.add_argument("--only", default="", help="comma ext filter e.g. docx,pdf")
    ap.add_argument("--no-ocr", action="store_true", help="只抽免 API 文字，跳過掃描 OCR")
    ap.add_argument("--ocr-only", action="store_true", help="只跑需 OCR 的掃描檔")
    ap.add_argument("--pace", type=float, default=1.5, help="Gemini 呼叫間隔秒數")
    args = ap.parse_args()
    global PACE
    PACE = args.pace
    only = {e.strip().lstrip(".").lower() for e in args.only.split(",") if e.strip()}

    files = sorted(p for p in SRC_ROOT.rglob("*") if p.is_file())
    # 便宜的先做：docx/xlsx → 文字 PDF → 圖片 → 其餘
    order = {".docx": 0, ".xlsx": 0, ".doc": 1, ".pdf": 2, ".jpg": 3, ".jpeg": 3, ".png": 3}
    files.sort(key=lambda p: (order.get(p.suffix.lower(), 9), p.stat().st_size))

    done = skip = fail = 0
    strikes = 0  # 連續「兩引擎皆配額用罄」次數；達 2 即停（過夜可重跑接續）
    for p in files:
        if only and p.suffix.lower().lstrip(".") not in only:
            continue
        rel = str(p.relative_to(SRC_ROOT)).replace("\\", "/")
        key = f"{TEXT_PREFIX}/{rel}.txt"
        if r2_exists(key):
            skip += 1
            continue
        try:
            text, method = process(p, args.no_ocr, args.ocr_only)
            strikes = 0
            if text is None:
                continue  # skipped by mode (e.g. --no-ocr on scan)
            if not text.strip():
                print(f"  ∅ empty: {rel} [{method}]", flush=True)
                fail += 1
                continue
            r2_put_text(key, text)
            done += 1
            print(f"  ✓ {rel}  [{method}] {len(text)} chars", flush=True)
            if args.limit and done >= args.limit:
                break
        except BothLimited:
            strikes += 1
            print(f"  ⏸ quota strike {strikes}/2 — {rel}", flush=True)
            if strikes >= 2:
                print("\n2-strike 配額停機；稍後重跑可接續（已完成的會跳過）。", flush=True)
                break
        except Exception as e:
            fail += 1
            print(f"  ✗ {rel}: {type(e).__name__}: {str(e)[:160]}", flush=True)
    print(f"\nDONE — transcribed {done}, skipped(exist) {skip}, failed {fail}", flush=True)

if __name__ == "__main__":
    main()
