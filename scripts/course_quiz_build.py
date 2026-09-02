# -*- coding: utf-8 -*-
"""小考產生器：題庫 → 線上考卷 HTML ＋ quizzes.json ＋ 紙本考卷 docx。

三種產出：
  1. 線上版　public/content/works/{slug}/quizzes/wr2-chNN.html
             並更新 {slug}-quizzes.json（bookId=WR2，與 WR1 的考卷分流）
  2. 紙本考卷（學生用，無答案）　Drive .../小考/第NN章小考.docx
  3. 教師用解答卷（含答案與解析）　Drive .../小考/第NN章小考（解答）.docx

「小考卷另外出」＝考卷是獨立檔案，不印在講義裡；線上版則附可切換顯示的參考答案。

用法：
  python scripts/course_quiz_build.py            # 全部十六章
  python scripts/course_quiz_build.py 1 2 3
"""
import html
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
ROOT = Path(__file__).resolve().parent.parent
PUB = ROOT / 'public'
DRIVE_ROOT = Path(r'G:\我的雲端硬碟\資料\知識圖工作室\教學')

# 一個課程一組設定；預設世界宗教文化導論，`--course=ch` 切到基督宗教概論。
COURSES = {
    'wr2': dict(slug='world-religions-intro', book_id='WR2', prefix='wr2',
                chapters='chapters-wr2', title='世界宗教文化導論',
                folder='115-1_世界宗教文化導論'),
    'ch': dict(slug='christianity-intro', book_id='CH1', prefix='ch1',
               chapters='chapters', title='基督宗教概論',
               folder='115-2_基督宗教概論'),
    'sl': dict(slug='sinographic-literature', book_id='SL1', prefix='sl1',
               chapters='chapters', title='宗教系國文講義',
               folder='宗教系國文講義'),
}


def load_course(key):
    """把課程設定攤成模組層常數；產生器各函式直接讀這些名字。"""
    global ALL, DRIVE, SLUG, BOOK_ID, PREFIX, CHAPTERS, COURSE_TITLE
    c = COURSES[key]
    SLUG, BOOK_ID, PREFIX = c['slug'], c['book_id'], c['prefix']
    CHAPTERS, COURSE_TITLE = c['chapters'], c['title']
    DRIVE = DRIVE_ROOT / c['folder'] / '小考'
    if key == 'wr2':
        from course_quiz_data import QUIZZES
        from course_quiz_data2 import QUIZZES2
        ALL = {**QUIZZES, **QUIZZES2}
    elif key == 'ch':
        from course_quiz_data_ch import QUIZZES_CH
        ALL = QUIZZES_CH
    else:
        from course_quiz_data_sl import QUIZZES_SL
        ALL = QUIZZES_SL


KLASS = '玄奘大學宗教與文化學系‧二年制在職專班1年A班'
KAI, MING, HEI = 'DFKai-SB', 'PMingLiU', 'Microsoft JhengHei'
GRAY = RGBColor(0x60, 0x60, 0x60)
NAVY = RGBColor(0x1F, 0x39, 0x64)

CN = '一二三四五六七八九十十一十二十三十四十五十六'
OPT = 'ABCD'


def chapter_title(n):
    """章名一律取自講義章節檔的 <h2>（去掉「第N章　」前綴），避免題庫與講義走鐘。"""
    f = PUB / 'content/works' / SLUG / CHAPTERS / f'ch{n:02d}.html'
    h = re.sub(r'<[^>]+>', '', re.search(r'<h2>(.*?)</h2>',
               f.read_text(encoding='utf-8')).group(1)).strip()
    return re.sub(r'^第[一二三四五六七八九十]+章[　\s]*', '', h)


def cn(n):
    # CN 前十個字是一~十，之後每兩個字一組（十一、十二…）
    return CN[n - 1] if n <= 10 else CN[10:][(n - 11) * 2:(n - 11) * 2 + 2]


# 出題時容易把正解都寫在同一個位置。這裡以確定性的輪轉把正解均勻攤到 A–D，
# 輪轉保持選項的相對順序，所以讀起來仍然自然；同一題每次產生的結果都一樣。
TARGETS = [0, 2, 1, 3, 2, 0, 3, 1, 1, 3, 0, 2, 3, 1, 2, 0]


def balanced(n, i, opts, ans):
    target = TARGETS[(n * 7 + i * 3) % len(TARGETS)] % 4
    shift = (ans - target) % 4
    return opts[shift:] + opts[:shift], target


# ── 線上版 HTML ─────────────────────────────────────────────────────────────
def build_html(n, q):
    e = html.escape
    out = [f'<p class="quiz-meta">{COURSE_TITLE}　第{cn(n)}章小考　'
           f'{e(chapter_title(n))}　共 10 題，每題 10 分</p>',
           '<h4>選擇題</h4>', '<ol>']
    for i, (stem, opts, ans, _) in enumerate(q['items']):
        opts, _ans = balanced(n, i, opts, ans)
        out.append(f'<li>{e(stem)}')
        out.append('<ul class="q-options">')
        for j, o in enumerate(opts):
            out.append(f'<li>({OPT[j]}) {e(o)}</li>')
        out.append('</ul></li>')
    out.append('</ol>')
    out.append('<div class="quiz-answers">')
    out.append('<h3>參考答案與解析</h3>')
    out.append('<ol>')
    for i, (_, opts, ans, why) in enumerate(q['items']):
        _, a = balanced(n, i, opts, ans)
        out.append(f'<li><strong>({OPT[a]})</strong>　{e(why)}</li>')
    out.append('</ol></div>')
    return '\n'.join(out) + '\n'


# ── 紙本考卷 ────────────────────────────────────────────────────────────────
def run_ea(par, text, font=MING, size=11.0, bold=False, color=None):
    r = par.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font)
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    return r


def page_numbers(doc, label):
    par = doc.sections[0].footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ea(par, f'{label}　－ ', MING, 9, color=GRAY)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    par._p.append(fld)
    run_ea(par, ' －', MING, 9, color=GRAY)


def build_docx(n, q, with_answers):
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Cm(2.0)
    s.left_margin = s.right_margin = Cm(2.3)
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'
    st.element.rPr.rFonts.set(qn('w:eastAsia'), MING)
    st.font.size = Pt(11)

    # 卷頭
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run_ea(p, KLASS, HEI, 10, color=GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run_ea(p, f'{COURSE_TITLE}　第{cn(n)}章小考', KAI, 18, bold=True, color=NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run_ea(p, chapter_title(n), KAI, 12, color=GRAY)

    if with_answers:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        run_ea(p, '【教師用‧含參考答案與解析】', HEI, 11, bold=True, color=NAVY)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        run_ea(p, '姓名：＿＿＿＿＿＿＿＿　　學號：＿＿＿＿＿＿＿＿　　'
                  '日期：＿＿＿＿＿＿　　得分：＿＿＿＿', MING, 11)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run_ea(p, '選擇題　共 10 題，每題 10 分。請選出最適當的一個答案。', HEI, 11, bold=True)

    for i, (stem, opts, ans, why) in enumerate(q['items'], 1):
        opts, ans = balanced(n, i - 1, opts, ans)
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(0.95); pf.first_line_indent = Cm(-0.95)
        pf.line_spacing = 1.4; pf.space_before = Pt(8); pf.space_after = Pt(3)
        run_ea(p, f'（　　）{i}. ' if not with_answers else f'（{OPT[ans]}）{i}. ',
               MING, 11, bold=with_answers)
        run_ea(p, stem, MING, 11)
        for j, o in enumerate(opts):
            po = doc.add_paragraph()
            pf = po.paragraph_format
            pf.left_indent = Cm(1.9); pf.first_line_indent = Cm(-0.75)
            pf.line_spacing = 1.35; pf.space_after = Pt(1)
            bold = with_answers and j == ans
            run_ea(po, f'({OPT[j]}) ', MING, 10.5, bold=bold)
            run_ea(po, o, MING, 10.5, bold=bold)
        if with_answers:
            pw = doc.add_paragraph()
            pf = pw.paragraph_format
            pf.left_indent = Cm(1.9); pf.line_spacing = 1.3
            pf.space_before = Pt(3); pf.space_after = Pt(2)
            run_ea(pw, '解析：', HEI, 9.5, bold=True, color=GRAY)
            run_ea(pw, why, MING, 9.5, color=GRAY)

    if not with_answers:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        run_ea(p, '※ 本卷不計入學期成績，用來確認自己讀懂了沒有。'
                  '答錯的題目請回頭看講義該章對應的一節。', MING, 10, color=GRAY)

    page_numbers(doc, f'{COURSE_TITLE}　第{cn(n)}章小考'
                      + ('（解答）' if with_answers else ''))
    DRIVE.mkdir(parents=True, exist_ok=True)
    name = f'第{n:02d}章小考' + ('（解答）' if with_answers else '') + '.docx'
    out = DRIVE / name
    doc.save(out)
    return out


# ── 主流程 ──────────────────────────────────────────────────────────────────
def main(nums):
    qdir = PUB / 'content/works' / SLUG / 'quizzes'
    qdir.mkdir(parents=True, exist_ok=True)
    meta_path = PUB / 'content/works' / f'{SLUG}-quizzes.json'
    meta = (json.loads(meta_path.read_text(encoding='utf-8'))
            if meta_path.exists() else {'quizzes': []})
    entries = [q for q in meta['quizzes'] if q.get('bookId') != BOOK_ID]

    for n in sorted(ALL):
        q = ALL[n]
        assert len(q['items']) == 10, f'第{n}章題數不是 10'
        for stem, opts, ans, why in q['items']:
            assert len(opts) == 4 and 0 <= ans < 4, f'第{n}章選項或答案有誤：{stem[:20]}'
        entries.append({
            'id': f'{PREFIX}-ch{n:02d}',
            'title': f'第{cn(n)}章',
            'range': chapter_title(n),
            'file': f'/content/works/{SLUG}/quizzes/{PREFIX}-ch{n:02d}.html',
            'bookId': BOOK_ID,
        })
        if n in nums:
            (qdir / f'{PREFIX}-ch{n:02d}.html').write_text(build_html(n, q), encoding='utf-8')
            build_docx(n, q, with_answers=False)
            build_docx(n, q, with_answers=True)
            print(f'✔ 第{n:02d}章　線上 HTML ＋ 考卷 ＋ 解答卷')

    meta['quizzes'] = entries
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2) + '\n',
                         encoding='utf-8')
    print(f'✔ quizzes.json：{BOOK_ID} '
          f'{len([q for q in entries if q.get("bookId") == BOOK_ID])} 張'
          f'（本檔共 {len(entries)} 張）')


if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    args = sys.argv[1:]
    course = next((a.split('=')[1] for a in args if a.startswith('--course=')), 'wr2')
    load_course(course)
    nums = set(int(a) for a in args if not a.startswith('--')) or set(ALL)
    main(nums)
