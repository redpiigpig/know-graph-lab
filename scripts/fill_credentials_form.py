# -*- coding: utf-8 -*-
"""填玄奘大學研究生助學金「證件黏貼表」，並把證件影本貼進表格。

資料一律抄自同一批已填好的 01 申請書與身分資料夾，不另行輸入，免得兩份表不一致。

🚨 這份表帶身分證字號。產出檔留在 Drive 的獎學金資料夾，不進 repo（repo 是公開的）。
🚨 表格有合併儲存格：python-docx 讀合併列時同一個 cell 會重複出現，
   直接照索引寫會把同一格寫好幾次。用 id() 去重才對得上。

  python -X utf8 scripts/fill_credentials_form.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE = Path("G:/我的雲端硬碟/玄奘/博一/獎學金/玄奘助學金")
ID = Path("G:/我的雲端硬碟/資料/身分資料")
SRC = BASE / "03-研究生助學金證件表(docx).docx"
OUT = BASE / "03-研究生助學金證件表（已填）.docx"

FIELDS = {                      # 標籤 → 要填的值
    "姓  名": "張辰瑋",
    "身分證字號": "A129181512",
    "出生年月日": "民國84年10月19日",
    "年  級": "一年級",
    "學  號": "DB1153002",
    "手  機": "09099095",
    "電子郵件": "redpiigpig@gmail.com",
}
IMAGES = {                      # 黏貼欄標籤 → 影本檔
    "身分證影本正面": ID / "身分證正面.png",
    "身分證影本背面": ID / "身分證反面.png",
    "匯款往來郵局封面影本": ID / "辰瑋郵局存摺.jpg",
}
CJK, EN = "標楷體", "Times New Roman"


def write(cell, text, *, size=12):
    cell.text = ""
    par = cell.paragraphs[0]
    run = par.add_run(text)
    run.font.size = Pt(size)
    run.font.name = EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)


def put_image(cell, path, width=Cm(7.2)):
    cell.text = ""
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(str(path), width=width)


def main():
    doc = Document(SRC)

    # 🚨 標題的「玄奘大學   學年度」被 Word 切成好幾個 run，逐 run 取代永遠對不上。
    #    要把整段文字併起來改完，再塞回第一個 run（其餘清空），才留得住原字型。
    for p in doc.paragraphs:
        if "學年度研究生助學金證件" in p.text and p.runs:
            fixed = p.text.replace("玄奘大學", "玄奘大學 115").replace("   學年度", " 學年度")
            p.runs[0].text = fixed
            for r in p.runs[1:]:
                r.text = ""

    t0 = doc.tables[0]
    for row in t0.rows:
        seen = []
        for k, cell in enumerate(row.cells):
            if id(cell._tc) in seen:            # 合併格會重複出現
                continue
            seen.append(id(cell._tc))
            label = cell.text.strip()
            if label in FIELDS and k + 1 < len(row.cells):
                write(row.cells[k + 1], FIELDS[label])
            elif label.startswith("系  所"):
                write(row.cells[k + 1], "宗教與文化學系\n■博士班　□碩士班")

    t1 = doc.tables[1]
    for row in t1.rows:
        for cell in row.cells:
            label = cell.text.strip()
            for key, path in IMAGES.items():
                if label.startswith(key) and path.exists():
                    put_image(cell, path)
            if label.startswith("學生證影本"):
                write(cell, f"{label}\n（115 學年度學生證尚未領取，註冊後補貼）", size=10)

    doc.save(OUT)
    print(f"→ {OUT.name}")


if __name__ == "__main__":
    main()
