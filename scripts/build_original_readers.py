"""Build three visually matched JIS B5 original-language reader samples."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "data" / "originalReaders" / "samples" / "psalm-shepherd.json"
OUTPUT_DIR = ROOT / "output" / "original-readers"

PAGE_WIDTH_MM = 182
PAGE_HEIGHT_MM = 257
MARGIN_INSIDE_MM = 24
MARGIN_OUTSIDE_MM = 17
MARGIN_TOP_MM = 18
MARGIN_BOTTOM_MM = 20
USABLE_WIDTH_MM = PAGE_WIDTH_MM - MARGIN_INSIDE_MM - MARGIN_OUTSIDE_MM

FONT_ZH = "Noto Serif TC"
FONT_UI = "Noto Sans TC"
FONT_LATIN = "Noto Serif"
INK = "25221E"
PAPER = "FFFDF7"
MUTED = "6F685F"
RULE = "D8D0C4"

HEBREW_WORD = re.compile(r"[\u05D0-\u05EA]")
HEBREW_VOWEL = re.compile(r"[\u05B0-\u05BB\u05C7]")

VOCABULARY_INTRO = {
    "hbo": "詞形欄照錄正文；lemma 欄才是查字典的形式。希伯來文的 pointed lemma 與無標點索引形必須分開儲存，紙本只讓無標點形式退居索引用途。",
    "grc": "詞形欄照錄正文；lemma 欄保留帶多調號的詞典形。供搜尋使用的無調號／正規化索引鍵另行儲存，不得取代正文或 lemma。",
    "la": "詞形欄照錄來源正文；lemma 欄使用詞典形，教學長音只屬 lemma／讀音層。搜尋索引鍵另行儲存，不得把後加長音寫回正文。",
}

VOCABULARY_QA = {
    "hbo": "排版檢查：正文詞形、詞典形、索引形三者不得因正規化而互相覆寫。",
    "grc": "排版檢查：正文多調號詞形、帶調號 lemma 與搜尋正規化鍵不得互相覆寫。",
    "la": "排版檢查：來源正文詞形、詞典 lemma、教學長音與搜尋索引鍵不得互相覆寫。",
}

LATIN_PSALM_SOURCE_URL = "https://ebible.org/study/content/texts/latVUC/PS22.html"


def assert_fully_pointed_hebrew(verse: str) -> None:
    """Reject any Hebrew orthographic word that lacks a Masoretic vowel sign."""
    for word in re.split(r"[\s\u05BE]+", verse):
        if HEBREW_WORD.search(word) and not HEBREW_VOWEL.search(word):
            raise AssertionError(f"Hebrew word missing niqqud: {word}")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_mm: float) -> None:
    cell.width = Mm(width_mm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(round(width_mm / 25.4 * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_fixed_table_layout(table, widths_mm: list[float]) -> None:
    """Lock table width/grid so LibreOffice does not redistribute columns."""
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(round(sum(widths_mm) / 25.4 * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_mm in widths_mm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(round(width_mm / 25.4 * 1440)))
        grid.append(grid_col)
    for col_idx, width_mm in enumerate(widths_mm):
        table.columns[col_idx].width = Mm(width_mm)
        for cell in table.columns[col_idx].cells:
            set_cell_width(cell, width_mm)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def set_table_borders(table, color=RULE, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_run_font(run, font: str, size: float, *, bold=False, color=INK, italic=False) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font)


def set_paragraph_rtl(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    for run in paragraph.runs:
        r_pr = run._r.get_or_add_rPr()
        rtl = r_pr.find(qn("w:rtl"))
        if rtl is None:
            rtl = OxmlElement("w:rtl")
            r_pr.append(rtl)
        rtl.set(qn("w:val"), "1")


def set_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def set_keep_lines(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_lines = OxmlElement("w:keepLines")
    p_pr.append(keep_lines)


def set_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    set_run_font(run, FONT_UI, 8, color=MUTED)


def configure_document(document: Document, volume: dict[str, Any]) -> None:
    section = document.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = Mm(MARGIN_TOP_MM)
    section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    section.left_margin = Mm(MARGIN_INSIDE_MM)
    section.right_margin = Mm(MARGIN_OUTSIDE_MM)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(9)
    section.different_first_page_header_footer = True

    settings = document.settings._element
    mirror = settings.find(qn("w:mirrorMargins"))
    if mirror is None:
        settings.append(OxmlElement("w:mirrorMargins"))

    normal = document.styles["Normal"]
    normal.font.name = FONT_ZH
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in (
        ("Title", 26, INK),
        ("Heading 1", 16, volume["accent"]),
        ("Heading 2", 13, INK),
        ("Heading 3", 11, MUTED),
    ):
        style = document.styles[style_name]
        style.font.name = FONT_ZH
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_ZH)
        style.paragraph_format.space_before = Pt(9 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{volume['title']}  ·  {volume['sampleTitle']}")
    set_run_font(run, FONT_UI, 7.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("私人研讀樣張   ·   ")
    set_run_font(run, FONT_UI, 8, color=MUTED)
    set_page_field(p)

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PRIVATE STUDY EDITION  ·  2026")
    set_run_font(run, FONT_UI, 7.5, color=MUTED)

    props = document.core_properties
    props.title = f"{volume['title']}｜{volume['sampleTitle']}"
    props.subject = "JIS B5 original-language reader sample"
    props.author = "Know Graph Lab"
    props.keywords = "Biblical languages, B5 reader, private study"


def add_rule(document: Document, color: str, height_pt=2) -> None:
    table = document.add_table(rows=1, cols=1)
    set_fixed_table_layout(table, [USABLE_WIDTH_MM])
    shade_cell(table.cell(0, 0), color)
    table.rows[0].height = Pt(height_pt)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(0)
    set_cell_margins(table.cell(0, 0), top=0, bottom=0, start=0, end=0)
    remove_table_borders(table)


def add_label(document: Document, text: str, accent: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text.upper())
    set_run_font(run, FONT_UI, 7.5, bold=True, color=accent)
    run.font.character_spacing = Pt(1.2)


def add_body(document: Document, text: str, *, size=10.5, color=INK, bold=False, italic=False, align=None) -> None:
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, FONT_ZH, size, bold=bold, color=color, italic=italic)
    return p


def page_break(document: Document) -> None:
    document.add_page_break()


def add_cover(document: Document, volume: dict[str, Any]) -> None:
    accent = volume["accent"]
    table = document.add_table(rows=1, cols=1)
    set_fixed_table_layout(table, [USABLE_WIDTH_MM])
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=420, bottom=420, start=360, end=360)
    shade_cell(cell, "17231F")
    remove_table_borders(table)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ORIGINAL-LANGUAGE READER")
    set_run_font(run, FONT_UI, 8, bold=True, color="E8BC67")
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(volume["title"])
    set_run_font(run, FONT_ZH, 25, bold=True, color="FFF7E8")
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(volume["subtitle"])
    set_run_font(run, FONT_LATIN, 10, color="D6D2C8")

    document.add_paragraph().paragraph_format.space_after = Pt(32)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(volume["sampleOriginalTitle"])
    set_run_font(run, volume["font"], 24, color=accent)
    if volume["direction"] == "rtl":
        set_paragraph_rtl(p)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(volume["sampleTitle"])
    set_run_font(run, FONT_ZH, 17, bold=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("三冊共用版式樣張  ·  校訂原文／繁中對照／詞形／音訊索引")
    set_run_font(run, FONT_UI, 8.5, color=MUTED)

    document.add_paragraph().paragraph_format.space_after = Pt(32)
    add_rule(document, accent, 22)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("JIS B5  182 × 257 mm")
    set_run_font(run, FONT_UI, 8.5, bold=True, color=MUTED)


def add_how_to_page(document: Document, volume: dict[str, Any]) -> None:
    add_label(document, "Reader architecture", volume["accent"])
    document.add_heading("這一本如何閱讀", level=1)
    add_body(document, "每一篇先讀完整原文，再用繁中對照確認句意；低頻詞、詞典形與形態分析集中在頁邊／詞表，避免中文遮斷原文視線。線上版沿用相同段落 ID，可直接跳到慢速、自然速或吟誦音軌。", size=10.5)

    cards = [
        ("01", "先讀原文", "保留校訂本的重音、母音點、標點與詩行。"),
        ("02", "再查詞形", "正文詞形、詞典形、字根／詞幹與中文義分欄。"),
        ("03", "最後跟讀", "音訊時間碼對齊段落與單字；裝置 TTS 不等於校訂音。"),
    ]
    table = document.add_table(rows=1, cols=3)
    set_fixed_table_layout(table, [USABLE_WIDTH_MM / 3] * 3)
    remove_table_borders(table)
    for idx, (number, title, text) in enumerate(cards):
        cell = table.cell(0, idx)
        set_cell_margins(cell, top=180, bottom=180, start=150, end=150)
        shade_cell(cell, "F2EEE5")
        p = cell.paragraphs[0]
        r = p.add_run(number)
        set_run_font(r, FONT_UI, 8, bold=True, color=volume["accent"])
        p = cell.add_paragraph()
        r = p.add_run(title)
        set_run_font(r, FONT_ZH, 11, bold=True)
        p = cell.add_paragraph()
        r = p.add_run(text)
        set_run_font(r, FONT_ZH, 8.5, color=MUTED)

    document.add_heading("本冊的正文規則", level=2)
    add_body(document, volume["textPolicy"], size=9.5, color=MUTED)
    document.add_heading("詞彙課程排序", level=2)
    add_body(document, volume["curriculum"], size=9.5, color=MUTED)

    document.add_heading("樣張來源", level=2)
    add_body(document, volume["sourceEdition"], size=9, color=MUTED)
    if volume["languageCode"] == "la":
        add_body(
            document,
            "來源版本：Vulgata Clementina / latVUC；Psalmus 22 = MT Psalm 23",
            size=7.8,
            color=volume["accent"],
        )
        source_link = add_body(
            document,
            f"完整 URL：{LATIN_PSALM_SOURCE_URL}",
            size=7.8,
            color=volume["accent"],
        )
        source_link.paragraph_format.space_after = Pt(6)
    else:
        source_url = urlparse(volume["sourceUrl"])
        compact_source = f"來源連結：{source_url.netloc}{source_url.path}"
        add_body(document, compact_source, size=7.8, color=volume["accent"])

    table = document.add_table(rows=1, cols=4)
    set_fixed_table_layout(table, [USABLE_WIDTH_MM / 4] * 4)
    remove_table_borders(table)
    specs = [("成品", "JIS B5"), ("內／外", "24 / 17 mm"), ("上／下", "18 / 20 mm"), ("裝訂", "鏡像邊界")]
    for i, (label, value) in enumerate(specs):
        cell = table.cell(0, i)
        set_cell_margins(cell, top=120, bottom=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + "\n")
        set_run_font(r, FONT_UI, 7, color=MUTED)
        r = p.add_run(value)
        set_run_font(r, FONT_UI, 8.5, bold=True)


def add_verse_pair(document: Document, volume: dict[str, Any], translation: str, verse: str, number: int) -> None:
    table = document.add_table(rows=2, cols=2)
    set_fixed_table_layout(table, [10, USABLE_WIDTH_MM - 10])
    remove_table_borders(table)
    for cell in table.row_cells(0) + table.row_cells(1):
        set_cell_margins(cell, top=80, bottom=80, start=80, end=80)
    shade_cell(table.cell(0, 0), volume["accent"])
    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(number))
    set_run_font(r, FONT_UI, 8, bold=True, color="FFFFFF")

    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if volume["direction"] == "rtl" else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.45 if volume["direction"] == "rtl" else 1.25
    set_keep_lines(p)
    r = p.add_run(verse)
    set_run_font(r, volume["font"], 15.5 if volume["languageCode"] == "hbo" else 12.2, color=INK)
    if volume["direction"] == "rtl":
        set_paragraph_rtl(p)

    p = table.cell(1, 1).paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(translation)
    set_run_font(r, FONT_ZH, 9, color=MUTED)


def add_reading_page(document: Document, volume: dict[str, Any], translation: list[str], start: int, end: int) -> None:
    add_label(document, volume["reference"], volume["accent"])
    document.add_heading(volume["sampleTitle"], level=1)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(volume["sampleOriginalTitle"])
    set_run_font(r, volume["font"], 13, color=volume["accent"])
    if volume["direction"] == "rtl":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_rtl(p)
    for idx in range(start, end):
        add_verse_pair(document, volume, translation[idx], volume["verses"][idx], idx + 1)


def add_vocabulary_page(document: Document, volume: dict[str, Any]) -> None:
    add_label(document, "Vocabulary & morphology", volume["accent"])
    document.add_heading("讀文詞彙與詞形", level=1)
    add_body(
        document,
        VOCABULARY_INTRO[volume["languageCode"]],
        size=9.5,
        color=MUTED,
    )

    headers = ["正文詞形", "詞典形", "字根", "中文義", "形態／句法"]
    widths = [30, 28, 14, 29, 40]
    table = document.add_table(rows=1, cols=5)
    set_fixed_table_layout(table, widths)
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_margins(cell, top=110, bottom=110, start=80, end=80)
        shade_cell(cell, volume["accent"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, FONT_UI, 7.5, bold=True, color="FFFFFF")
    set_repeat_table_header(table.rows[0])

    for item in volume["vocabulary"]:
        cells = table.add_row().cells
        for i, value in enumerate(item):
            set_cell_margins(cells[i], top=90, bottom=90, start=80, end=80)
            if len(table.rows) % 2 == 1:
                shade_cell(cells[i], "F7F3EB")
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if volume["direction"] == "rtl" and i < 3 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            font = volume["font"] if i < 3 else FONT_ZH
            entry_size = 8.4 if volume["languageCode"] == "grc" and i < 3 else (9 if i < 3 else 7.8)
            set_run_font(r, font, entry_size, color=INK if i != 4 else MUTED)
            if volume["direction"] == "rtl" and i < 3:
                set_paragraph_rtl(p)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(VOCABULARY_QA[volume["languageCode"]])
    set_run_font(r, FONT_UI, 8, bold=True, color=volume["accent"])


def add_close_reading_page(document: Document, volume: dict[str, Any], translation: list[str]) -> None:
    add_label(document, "Close reading", volume["accent"])
    document.add_heading("第一節：逐層細讀", level=1)

    table = document.add_table(rows=1, cols=1)
    set_fixed_table_layout(table, [USABLE_WIDTH_MM])
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=220, bottom=200, start=220, end=220)
    shade_cell(cell, "F2EEE5")
    remove_table_borders(table)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if volume["direction"] == "rtl" else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(volume["verses"][0])
    set_run_font(r, volume["font"], 15 if volume["languageCode"] == "hbo" else 12.5)
    if volume["direction"] == "rtl":
        set_paragraph_rtl(p)
    p = cell.add_paragraph()
    r = p.add_run(translation[0])
    set_run_font(r, FONT_ZH, 9.5, color=MUTED)

    document.add_heading("觀察", level=2)
    for index, note in enumerate(volume["grammar"], 1):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Mm(5)
        p.paragraph_format.first_line_indent = Mm(-5)
        r = p.add_run(f"{index:02d}  ")
        set_run_font(r, FONT_UI, 8, bold=True, color=volume["accent"])
        r = p.add_run(note)
        set_run_font(r, FONT_ZH, 9.5, color=INK)

    document.add_heading("翻譯決策", level=2)
    add_body(document, "先標出主語、謂語與代名詞指涉，再處理詩歌意象。不要先背中文句子後反推原文；讀本的目的，是讓原文結構成為理解的第一層。", size=9.5, color=MUTED)

    add_rule(document, volume["accent"], 2)
    add_body(document, "練習：不看中文，圈出本節的主語、動詞與第一人稱形式；再用一句中文寫出句法骨架。", size=9, bold=True, color=volume["accent"])


def add_audio_practice_page(document: Document, volume: dict[str, Any]) -> None:
    add_label(document, "Aligned audio & practice", volume["accent"])
    document.add_heading("線上音訊與紙本練習", level=1)
    add_body(document, "每個紙本段落都有固定 segment ID；正式錄音以毫秒時間碼對齊段落，完成第二輪校訂後再加 token cue。使用者可切換慢速、自然速與吟誦，並重複 A–B 區間。", size=9.8)

    widths = [25, 50, 66]
    table = document.add_table(rows=4, cols=3)
    set_fixed_table_layout(table, widths)
    set_table_borders(table)
    rows = [
        ("音軌", "用途", "品質門檻"),
        ("慢速", "辨認母音／詞尾／連音", "不得改動音位；停頓按句法"),
        ("自然速", "整段跟讀與理解", "母語／受訓讀者；雙人校聽"),
        ("吟誦", "禮儀或傳統讀法", "標示傳統、角色、旋律版本"),
    ]
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            set_cell_margins(cell, top=110, bottom=110, start=100, end=100)
            if row_idx == 0:
                shade_cell(cell, volume["accent"])
            p = cell.paragraphs[0]
            r = p.add_run(value)
            set_run_font(r, FONT_UI if row_idx == 0 else FONT_ZH, 8, bold=row_idx == 0, color="FFFFFF" if row_idx == 0 else INK)

    document.add_heading("本篇練習", level=2)
    exercises = [
        "朗讀兩遍：第一遍只求準確，第二遍按詩行語義分組。",
        "遮住中文，為六節各寫一個不超過八字的中文標題。",
        "從詞表選四字，寫出正文詞形、lemma 與形態差異。",
        "聽自然速音軌，標出與自己預期不同的重音或停頓。",
        "隔日重讀，不查詞表完成 80% 句意後才標記『已讀』。",
    ]
    for i, text in enumerate(exercises, 1):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Mm(7)
        p.paragraph_format.first_line_indent = Mm(-7)
        r = p.add_run(f"□ {i}. ")
        set_run_font(r, FONT_UI, 9, bold=True, color=volume["accent"])
        r = p.add_run(text)
        set_run_font(r, FONT_ZH, 9.2)

    if volume["languageCode"] == "hbo":
        add_body(document, "注意：希伯來文線上版不啟用現代以色列語裝置 TTS；沒有校訂音軌時，寧可顯示『待匯入』。", size=8.7, bold=True, color=volume["accent"])
    else:
        add_body(document, "裝置語音只供介面測試與定位，畫面必須明示『非校訂發音』，不可計入正式音訊完成率。", size=8.7, bold=True, color=volume["accent"])


def add_series_page(document: Document, volume: dict[str, Any]) -> None:
    add_label(document, "Three-volume system", volume["accent"])
    document.add_heading("三冊共同系統", level=1)
    add_body(document, "紙本不只是輸出檔，線上也不只是播放器。兩者共用選文 ID、段落 ID、詞位 ID、版本資訊與校訂狀態；任何修訂只建立一次，再同步到 B5、PDF 與私人網站。", size=10)

    series = [
        ("希伯來文", "15 章 Tanakh＋20 禱文＋完整 Haggadah＋20 拉比選文", "50 × 20 BBH 詞彙課"),
        ("希臘文", "15 章 NT＋15 章 LXX／次經＋金口若望禮儀＋20 希臘教父", "Mounce 基礎＋Graded Reader"),
        ("拉丁文", "15 章 Vulgata＋20 禱文／信經＋完整 Ordo Missae＋20 拉丁教父", "聖經／禮儀／教父分級詞彙"),
    ]
    table = document.add_table(rows=3, cols=1)
    set_fixed_table_layout(table, [USABLE_WIDTH_MM])
    remove_table_borders(table)
    for idx, (title, contents, vocab) in enumerate(series):
        cell = table.cell(idx, 0)
        set_cell_margins(cell, top=120, bottom=120, start=220, end=220)
        shade_cell(cell, "F2EEE5" if idx != 1 else "E9E4D9")
        p = cell.paragraphs[0]
        r = p.add_run(title + "  ")
        set_run_font(r, FONT_ZH, 11, bold=True, color=volume["accent"])
        r = p.add_run(contents)
        set_run_font(r, FONT_ZH, 8.8)
        p = cell.add_paragraph()
        r = p.add_run("詞彙路徑｜" + vocab)
        set_run_font(r, FONT_UI, 8, color=MUTED)

    document.add_heading("完成判準", level=2)
    checks = [
        "正文版本、段落範圍與授權來源可追溯",
        "希伯來文 niqqud／希臘文多調號／拉丁文句讀未被正規化破壞",
        "詞形與 lemma 經人工抽查，中文對照不遮蔽原文",
        "音訊有發音方案、讀者、速度、checksum 與 cue 時間碼",
        "DOCX 與 PDF 每頁渲染檢查；線上路由登入限制與 noindex 驗證",
    ]
    for text in checks:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Mm(5)
        p.paragraph_format.first_line_indent = Mm(-5)
        r = p.add_run("✓ ")
        set_run_font(r, FONT_UI, 9, bold=True, color=volume["accent"])
        r = p.add_run(text)
        set_run_font(r, FONT_ZH, 8.6)

    add_rule(document, volume["accent"], 2)
    add_body(document, "本檔是版式與資料結構樣張；未匯入的詞表、教父全文與正式錄音不標示為完成。", size=8, bold=True, color=MUTED)


def build_volume(volume: dict[str, Any], translation: list[str]) -> Path:
    if volume["languageCode"] == "hbo":
        for verse in volume["verses"]:
            assert_fully_pointed_hebrew(verse)

    document = Document()
    configure_document(document, volume)
    add_cover(document, volume)
    page_break(document)
    add_how_to_page(document, volume)
    page_break(document)
    add_reading_page(document, volume, translation, 0, 3)
    page_break(document)
    add_reading_page(document, volume, translation, 3, 6)
    page_break(document)
    add_vocabulary_page(document, volume)
    page_break(document)
    add_close_reading_page(document, volume, translation)
    page_break(document)
    add_audio_practice_page(document, volume)
    page_break(document)
    add_series_page(document, volume)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / f"{volume['slug']}-original-reader-sample.docx"
    document.save(path)
    return path


def main() -> None:
    payload = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    paths = [build_volume(volume, payload["translation"]) for volume in payload["volumes"]]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
