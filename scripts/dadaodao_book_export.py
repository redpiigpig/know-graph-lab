# -*- coding: utf-8 -*-
"""《當代的大愛道革命》全書初稿 → 網頁版 md + Word（docx，真頁尾註）。

來源：repo 根目錄《當代的大愛道革命_全書初稿.md》
輸出：public/content/works/mahaprajapati-revolution-book.md（網頁渲染用，去工作筆記，保留章末尾註清單）
      public/content/works/mahaprajapati-revolution-book.docx（下載用）

Word 流程：md（去工作筆記、去「尾註（第X章）」標題）→ pandoc（[^n] 轉真正的
Word 頁尾註，自動連續編號）→ python-docx 後處理：封面頁、章起新頁、
標題／引文標楷體、正文新細明體與首行縮排、註腳小字。
"""
import re
import sys
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def add_page_numbers(doc):
    """頁尾置中頁碼（PAGE 欄位），方便確認頁數。"""
    for sec in doc.sections:
        p = sec.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        for tag, attrs, text in [
            ('w:fldChar', {'w:fldCharType': 'begin'}, None),
            ('w:instrText', {'xml:space': 'preserve'}, ' PAGE '),
            ('w:fldChar', {'w:fldCharType': 'end'}, None),
        ]:
            e = OxmlElement(tag)
            for k, v in attrs.items():
                e.set(qn(k), v)
            if text:
                e.text = text
            run._element.append(e)

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / '當代的大愛道革命_全書初稿.md'
OUT_MD = ROOT / 'public/content/works/mahaprajapati-revolution-book.md'
OUT_DOCX = ROOT / 'public/content/works/mahaprajapati-revolution-book.docx'

KAI = 'DFKai-SB'      # 標楷體
MING = 'PMingLiU'     # 新細明體（正文）


def style_fonts(doc, name, east, ascii_font='Times New Roman', size=None, bold=None):
    try:
        st = doc.styles[name]
    except KeyError:
        return None
    st.font.name = ascii_font
    rf = st.element.get_or_add_rPr().get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), east)
    if size is not None:
        st.font.size = Pt(size)
    if bold is not None:
        st.font.bold = bold
    return st


def build():
    md = SRC.read_text(encoding='utf-8')
    cut = md.find('## 【全書初稿工作筆記】')
    if cut >= 0:
        md = md[:cut].rstrip() + '\n'
    OUT_MD.write_text(md, encoding='utf-8')  # 網頁版：保留章末尾註清單

    # Word 版：尾註標題拿掉（pandoc 會把 [^n] 定義收成頁尾註，標題會變空節）
    word_md = re.sub(r'(?m)^###\s*尾註（[^）]*）\s*\n', '', md)
    pypandoc.convert_text(
        word_md, 'docx', format='markdown+footnotes',
        outputfile=str(OUT_DOCX),
        extra_args=['--wrap=none'],
    )

    doc = Document(str(OUT_DOCX))

    # 樣式：正文明體、標題與引文標楷體、註腳小字
    for n in ('Normal', 'Body Text', 'First Paragraph'):
        style_fonts(doc, n, MING, size=11)
    style_fonts(doc, 'Block Text', KAI, size=11)          # 引文＝標楷體
    style_fonts(doc, 'Footnote Text', MING, size=9)       # 頁尾註小字
    style_fonts(doc, 'Heading 1', KAI, size=18, bold=True)
    style_fonts(doc, 'Heading 2', KAI, size=14, bold=True)
    style_fonts(doc, 'Heading 3', KAI, size=12, bold=True)
    style_fonts(doc, 'Title', KAI, size=26, bold=True)

    # 段落版式：正文首行縮排、引文縮排；章（Heading 1）起新頁＋置中
    for n in ('Body Text', 'First Paragraph'):
        st = doc.styles[n] if n in [s.name for s in doc.styles] else None
        if st is not None:
            st.paragraph_format.first_line_indent = Cm(0.85)
            st.paragraph_format.line_spacing = 1.7
    try:
        bq = doc.styles['Block Text']
        bq.paragraph_format.left_indent = bq.paragraph_format.right_indent = Cm(1.0)
        bq.paragraph_format.line_spacing = 1.6
    except KeyError:
        pass
    h1s = [p for p in doc.paragraphs if p.style.name == 'Heading 1']
    for i, p in enumerate(h1s):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i > 0:
            p.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)

    # 封面頁＋目次頁（插到文首）
    cover = []
    for txt, size, bold, before in [
        ('當代的大愛道革命', 26, True, 220),
        ('昭慧法師與性廣法師的人間佛教思想與實踐', 14, False, 18),
        ('張辰瑋', 14, False, 60),
        ('（全書初稿）', 10, False, 10),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        r = p.add_run(txt)
        r.font.name = 'Times New Roman'
        r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), KAI)
        r.font.size = Pt(size)
        r.font.bold = bold
        cover.append(p)
    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)
    cover.append(pb)
    # 目次（TOC 欄位；scripts/dadaodao_pages_sync.py 會用 Word 更新出頁碼）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('目　次')
    r.font.name = 'Times New Roman'
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), KAI)
    r.font.size = Pt(18)
    r.font.bold = True
    cover.append(p)
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-2" \h \z \u')
    r_el = OxmlElement('w:r')
    t_el = OxmlElement('w:t')
    t_el.text = '（目次欄位：由頁數同步腳本或 Word F9 更新）'
    r_el.append(t_el)
    fld.append(r_el)
    p._p.append(fld)
    cover.append(p)
    pb2 = doc.add_paragraph()
    pb2.add_run().add_break(WD_BREAK.PAGE)
    cover.append(pb2)
    body = doc.element.body
    for i, p in enumerate(cover):
        body.remove(p._element)
        body.insert(i, p._element)

    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = sec.right_margin = Cm(2.8)
    add_page_numbers(doc)
    doc.save(str(OUT_DOCX))

    n_fn = len(re.findall(r'(?m)^\[\^[^\]]+\]:', word_md))
    print(f'md   → {OUT_MD} ({OUT_MD.stat().st_size:,} bytes)')
    print(f'docx → {OUT_DOCX} ({OUT_DOCX.stat().st_size:,} bytes)，頁尾註 {n_fn} 條')


if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    build()
