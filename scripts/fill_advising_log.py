# -*- coding: utf-8 -*-
"""產生整個補助期間的「研究生論文指導紀錄表」——**每次討論一個獨立的 Word 檔**。

表上註 2 寫「每次討論前請列印本表」，所以要一次一個檔案分開帶去印，
不是一份 21 頁的合訂本。檔名照日期排序，一眼看得出哪一張是哪一次。

《玄奘大學研究生助學金實施辦法》第六條第二款：每月與指導教授討論二次以上，
次月五日前彙整送研發處。表上註 2 又寫「每次討論前請列印本表，討論後師生簽名」，
所以這是**一次一張**的表，不是一學期一張。雙週三 = 每月兩次，剛好對得上。

只填不會變的欄位（姓名學號、指導教授、院系、論文題目、日期時間、第幾次），
討論內容與簽名留白給當天填。

🚨 表格有合併儲存格：python-docx 讀合併列時同一個 tc 會在 row.cells 重複出現，
   照索引寫會把同一格寫好幾次、或把值寫到隔壁去。一律用 id(tc) 去重。
🚨 **多段落的儲存格不可以用 write() 覆蓋。** write() 會先把整格清空，
   而「學生與指導教授互動情形」那一格有八段共六個項目（定期討論／同時討論人數／
   兩題五等第／主動提問／上次待修正），整格覆蓋等於把第 2–6 項刪掉，
   印出來完全正常、只是少了四題。同一列的「□博士班 □碩士班 □碩專班」也是三個選項。
   這種格子一律用 edit_par() 就地改該行的字，其餘段落不動。
🚨 產出檔留 Drive，不進 repo（含學號等個資，repo 是公開的）。

  python -X utf8 scripts/fill_advising_log.py
"""
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

BASE = Path("G:/我的雲端硬碟/玄奘/博一上/獎學金/玄奘助學金")
SRC = BASE / "04-研究生論文指導記錄表(docx) (1).docx"
OUTDIR = BASE / "04-研究生論文指導記錄表（已填）"

FIRST = date(2026, 9, 16)     # 115 學年度第一學期開學後第一個週三
LAST = date(2027, 6, 30)      # 助學金補助期間到 2027 年 6 月
EVERY = 14                    # 雙週

NAME = "張辰瑋（DB1153002）"
ADVISOR = "釋昭慧"
COLLEGE = "社會科學院／宗教與文化學系"
TITLE = "從彼岸向此岸的轉向：台灣佛教與基督教公共性之宗教史比較研究（1920–2020）"
PLACE = "玄奘大學宗教與文化學系研究室"
TIME = "14 時 00 分至　15 時 00 分"
CJK, EN = "標楷體", "Times New Roman"


def dates():
    out, d = [], FIRST
    while d <= LAST:
        out.append(d)
        d += timedelta(days=EVERY)
    return out


def write(cell, text, *, size=12):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.font.name = EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)


def edit_par(par, old, new):
    """就地把段落裡的 old 換成 new，其餘段落與格式不動。

    跨 run 的文字要先併起來再換：Word 常把一行切成好幾個 run，
    逐 run 比對永遠對不上。併完塞回第一個 run、其餘清空，格式沿用第一個 run 的。
    """
    if old not in par.text or not par.runs:
        return False
    par.runs[0].text = par.text.replace(old, new)
    for r in par.runs[1:]:
        r.text = ""
    return True


def uniq(row):
    """合併格在 row.cells 會重複出現，回傳 (欄索引, cell) 的去重序列。"""
    seen, out = [], []
    for i, c in enumerate(row.cells):
        if id(c._tc) in seen:
            continue
        seen.append(id(c._tc))
        out.append((i, c))
    return out


def fill(table, d, n):
    roc = d.year - 1911
    stamp = f"{roc} 年 {d.month} 月 {d.day} 日（星期三）"
    pairs = {
        "研究生姓名": NAME, "指導教授姓名": ADVISOR, "院      系": COLLEGE,
        "討論日期": stamp, "討論時間": TIME, "討論地點": PLACE, "論文題目": TITLE,
    }
    for row in table.rows:
        cells = uniq(row)
        for k, (idx, cell) in enumerate(cells):
            label = cell.text.strip().replace("\n", "")
            for key, val in pairs.items():
                if label.startswith(key.replace(" ", "")) or label.startswith(key):
                    if k + 1 < len(cells):
                        write(cells[k + 1][1], val, size=11 if key == "論文題目" else 12)
            if label.startswith("院"):
                # 學制那格是「□博士班 □碩士班 □碩專班」三選一，只勾不覆蓋
                for par in cells[-1][1].paragraphs:
                    edit_par(par, "□博士班", "■博士班")
            if label.startswith("1.本次為本學期第"):
                for par in cell.paragraphs:
                    if edit_par(par, "第    次討論", f"第 {n} 次討論"):
                        edit_par(par, "是否為定期討論？□是", "是否為定期討論？■是")


def main():
    OUTDIR.mkdir(parents=True, exist_ok=True)
    ds = dates()
    per_sem = {}
    for d in ds:
        # 🚨 第一學期是 9 月到**隔年 1 月**，不是到 12 月為止；用「年份換了就換學期」
        #    會把 1 月那兩次算成第二學期的第 1、2 次。
        sem = 1 if (d.month >= 8 or d.month == 1) else 2
        per_sem[sem] = per_sem.get(sem, 0) + 1
        doc = Document(SRC)
        fill(doc.tables[0], d, per_sem[sem])
        name = (f"指導紀錄_{d:%Y-%m-%d}_115-{sem}第{per_sem[sem]:02d}次.docx")
        doc.save(OUTDIR / name)
    print(f"{len(ds)} 個檔（{ds[0]} … {ds[-1]}）→ {OUTDIR.name}/")


if __name__ == "__main__":
    main()
