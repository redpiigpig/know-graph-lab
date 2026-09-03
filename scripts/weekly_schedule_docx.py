# -*- coding: utf-8 -*-
"""辰瑋 115-1 週課表——單頁 Word。

一頁裝得下是硬條件，因此橫式、窄邊界、節次只印到第 13 節（家教最晚到 20:30）。
教課／修課／家教三色以底色區分，單雙週在方塊裡標「單」「雙」。

用法：python scripts/weekly_schedule_docx.py
輸出：G:\\我的雲端硬碟\\玄奘\\博一\\115-1 週課表.docx
"""
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(r'G:\我的雲端硬碟\玄奘\博一') / '115-1 週課表.docx'

KAI, MING, HEI = 'DFKai-SB', 'PMingLiU', 'Microsoft JhengHei'
INK = RGBColor(0x18, 0x21, 0x2A)
GRAY = RGBColor(0x6C, 0x76, 0x81)
NAVY = RGBColor(0x1F, 0x5E, 0x78)

PERIODS = [
    (1, '08:30–09:20'), (2, '09:25–10:15'), (3, '10:25–11:15'), (4, '11:20–12:10'),
    (5, '12:10–13:10'), (6, '13:10–14:00'), (7, '14:10–15:00'), (8, '15:10–16:00'),
    (9, '16:10–17:00'), (10, '17:10–18:00'), (11, '18:30–19:15'), (12, '19:15–20:00'),
    (13, '20:10–20:55'),
]

# kind: teach 教課／study 修課／tutor 家教
# day: 1=一 … 7=日；p0/p1 起訖節次；cyc 單／雙週
EVENTS = [
    dict(kind='study', day=1, p0=2, p1=4, name='宗教研究基本問題與研究方法',
         who='根瑟馬庫斯', room='妙然 208', code='BBJ001', tag='博一‧專必‧3 學分'),
    dict(kind='tutor', day=1, p0=11, p1=13, name='國中家教', who='', room='',
         code='', tag='19:00–20:30'),
    dict(kind='teach', day=3, p0=1, p1=2, name='世界宗教文化導論',
         who='我', room='妙然 401', code='BBE275', tag='宗教系 1A‧專必‧2 學分'),
    dict(kind='study', day=3, p0=3, p1=4, name='唯識思想專題研討',
         who='釋昭慧、丹增南卓', room='妙然 206', code='BBJ029', tag='博一‧專選‧2 學分'),
    dict(kind='teach', day=3, p0=8, p1=9, name='基督宗教概論',
         who='我', room='妙然 401', code='BBE150', tag='宗教系 2A‧專必‧2 學分'),
    dict(kind='tutor', day=3, p0=10, p1=10, name='小五家教', who='', room='',
         code='', tag='17:00–18:00'),
    dict(kind='tutor', day=3, p0=11, p1=13, name='國中家教', who='', room='',
         code='', tag='19:00–20:30'),
    dict(kind='tutor', day=5, p0=11, p1=13, name='國中家教', who='', room='',
         code='', tag='19:00–20:30'),
    dict(kind='study', day=6, p0=1, p1=4, cyc='單', name='宗教學理論與方法（一）',
         who='根瑟馬庫斯', room='妙然 308', code='JJA051', tag='碩專 1A‧專必‧2 學分'),
    dict(kind='teach', day=6, p0=6, p1=9, cyc='單', name='國文',
         who='我', room='妙然 206', code='PPA066', tag='二專 1A‧通必‧2 學分'),
    dict(kind='teach', day=7, p0=6, p1=9, cyc='雙', name='世界宗教文化導論',
         who='我', room='妙然 401', code='PPA001', tag='二專 1A‧專必‧2 學分'),
]

FILL = {'teach': 'DFEDF3', 'study': 'E6EFDD', 'tutor': 'F6E7DE', 'rest': 'F2F2F2'}
FG = {'teach': RGBColor(0x17, 0x51, 0x6A), 'study': RGBColor(0x3F, 0x5F, 0x30),
      'tutor': RGBColor(0x8A, 0x45, 0x2A)}


def run(par, text, size, font=HEI, bold=False, color=INK, space=0):
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(space)
    par.paragraph_format.line_spacing = 1.0
    r = par.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    return r


def shade(cell, hexcolor):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def build():
    doc = Document()
    s = doc.sections[0]
    s.orientation = WD_ORIENT.LANDSCAPE
    s.page_width, s.page_height = Cm(29.7), Cm(21.0)
    s.top_margin = s.bottom_margin = Cm(1.0)
    s.left_margin = s.right_margin = Cm(1.1)

    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'
    st.element.rPr.rFonts.set(qn('w:eastAsia'), HEI)
    st.font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, '玄奘大學宗教與文化學系　115 學年度第 1 學期　張辰瑋　週課表',
        15, KAI, bold=True, color=NAVY, space=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, '教課 4 門 8 學分　修課 3 門 7 學分　家教 4 場／週　'
           '教室皆在妙然樓', 9, HEI, color=GRAY, space=4)

    days = ['一', '二', '三', '四', '五', '六', '日']
    tbl = doc.add_table(rows=len(PERIODS) + 1, cols=8)
    tbl.style = 'Table Grid'
    tbl.autofit = False
    widths = [Cm(2.5)] + [Cm(3.5)] * 7
    for ci, w in enumerate(widths):
        for row in tbl.rows:
            row.cells[ci].width = w

    hdr = tbl.rows[0]
    shade(hdr.cells[0], '1F5E78')
    run(hdr.cells[0].paragraphs[0], '節次／時間', 9, HEI, bold=True,
        color=RGBColor(0xFF, 0xFF, 0xFF))
    hdr.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, d in enumerate(days):
        c = hdr.cells[i + 1]
        shade(c, '1F5E78')
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(c.paragraphs[0], f'週{d}', 10, HEI, bold=True,
            color=RGBColor(0xFF, 0xFF, 0xFF))

    for ri, (no, clock) in enumerate(PERIODS, start=1):
        row = tbl.rows[ri]
        row.height = Cm(0.86)
        c = row.cells[0]
        c.vertical_anchor = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(c.paragraphs[0], f'{no}　{clock}', 8, HEI, color=GRAY)
        if no == 5:
            shade(c, FILL['rest'])
            for ci in range(1, 8):
                shade(row.cells[ci], FILL['rest'])
            run(row.cells[4].paragraphs[0], '午休 12:10–13:10', 8, HEI, color=GRAY)
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for e in EVENTS:
        col = e['day']
        top = tbl.rows[e['p0']].cells[col]
        if e['p1'] > e['p0']:
            top = top.merge(tbl.rows[e['p1']].cells[col])
        shade(top, FILL[e['kind']])
        top.vertical_anchor = WD_ALIGN_VERTICAL.CENTER
        for par in list(top.paragraphs)[1:]:
            par._element.getparent().remove(par._element)
        par = top.paragraphs[0]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        head = (f'［{e["cyc"]}週］' if e.get('cyc') else '') + e['name']
        run(par, head, 9.5, HEI, bold=True, color=FG[e['kind']], space=1)
        if e.get('room'):
            p2 = top.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(p2, f'{e["room"]}　{e["code"]}', 8, HEI, color=GRAY, space=1)
        if e.get('tag'):
            p3 = top.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(p3, e['tag'], 7.5, HEI, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run(p, '■ 教課（藍）　■ 修課（綠）　■ 家教（橙）　'
           '單／雙週依學期週次：第一週 9/7–9/13。'
           '單週六＝9/12、9/26、10/10、10/24、11/7、11/21、12/5、12/19、1/2；'
           '雙週日＝9/20、10/4、10/18、11/1、11/15、11/29、12/13、12/27、1/10。',
        8, HEI, color=GRAY, space=1)
    p = doc.add_paragraph()
    run(p, '待確認：10/10 為國慶日，單週六的國文是否調課；'
           'BBJ001 開課系統為第 2–4 節（09:25–12:10、3 學分）。',
        8, HEI, color=RGBColor(0x8A, 0x45, 0x2A))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    print('✔', build())
