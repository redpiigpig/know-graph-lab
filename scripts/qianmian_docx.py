# -*- coding: utf-8 -*-
"""千面上帝：把 28 章 markdown 組成七卷 Word，交到 Drive。

每卷四章。註釋走 docx_footnotes 的真頁下註（Word 自己編號、排在當頁下緣），
不是章末尾註，也不是把註文塞在段落後面。

版面沿用本 repo 既有的中文書稿慣例：A4、12pt 新細明體＋Times New Roman、
行距 1.5、首行縮排兩字；章與卷首各自另起一頁。

用法：
    python scripts/qianmian_docx.py              # 出七卷到 Drive
    python scripts/qianmian_docx.py --volume 3   # 只出第三卷
    python scripts/qianmian_docx.py --out c:/tmp # 改輸出位置（試印用）
"""
import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx_footnotes import Footnotes  # noqa: E402

sys.stdout.reconfigure(encoding="utf-8")
ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "output" / "qianmian" / "sources"
CH = ROOT / "output" / "qianmian" / "chapters"
DRIVE = Path(r"G:\我的雲端硬碟\資料\知識圖工作室\寫作計畫\書籍寫作\千面上帝")

EN_FONT, CJK_FONT = "Times New Roman", "新細明體"
INDENT = Pt(24)                      # 12pt × 2 字
FN_DEF = re.compile(r"^\[\^(\d+)\]:\s*(.+)$")
FN_REF = re.compile(r"\[\^(\d+)\]")


def run(par, text, *, size=12, bold=False, color=None):
    r = par.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = EN_FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if color:
        r.font.color.rgb = color
    return r


def para(doc, *, align=None, after=6, indent=None, before=0):
    p = doc.add_paragraph()
    f = p.paragraph_format
    f.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    f.space_after, f.space_before = Pt(after), Pt(before)
    if align is not None:
        p.alignment = align
    if indent is not None:
        f.first_line_indent = indent
    return p


def read_chapter(no):
    """回傳 (標題, 副標, [(節名, [段落])], {註號: 註文})。"""
    lines = (CH / f"ch{no:02d}.md").read_text(encoding="utf-8").splitlines()
    notes, body = {}, []
    for ln in lines:
        m = FN_DEF.match(ln.strip())
        if m:
            notes[int(m.group(1))] = m.group(2).strip()
        else:
            body.append(ln)

    title = sub = ""
    blocks, cur = [], ("導言", [])
    for ln in body:
        t = ln.strip()
        if not t or t == "---" or t.startswith("<!--"):
            continue
        if t.startswith("# "):
            title = t[2:].strip()
        elif t.startswith("### "):
            sub = t[4:].strip()
        elif t.startswith("## "):
            if cur[1]:
                blocks.append(cur)
            cur = (t[3:].strip(), [])
        else:
            cur[1].append(t)
    if cur[1]:
        blocks.append(cur)
    return title, sub, blocks, notes


def add_body(doc, fn, text, notes):
    """一個段落：遇到 [^n] 就插真註腳，其餘照排。"""
    p = para(doc, indent=INDENT)
    pos = 0
    for m in FN_REF.finditer(text):
        if m.start() > pos:
            run(p, text[pos:m.start()])
        body = notes.get(int(m.group(1)))
        if body:
            fn.add(p, body)
        pos = m.end()
    if pos < len(text):
        run(p, text[pos:])
    return p


def build_volume(vol_no, vol_name, chapters, out_dir):
    doc = Document()
    s = doc.sections[0]
    s.page_height, s.page_width = Cm(29.7), Cm(21.0)
    s.top_margin = s.bottom_margin = Cm(2.54)
    s.left_margin = s.right_margin = Cm(3.17)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = EN_FONT, Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)

    fn = Footnotes(doc)
    grey = RGBColor(0x60, 0x60, 0x60)

    # ── 卷首 ──────────────────────────────────────────────
    for _ in range(6):
        para(doc, after=0)
    run(para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=18), "千面上帝", size=30, bold=True)
    run(para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=6), f"第{vol_no}卷", size=16)
    run(para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=48), vol_name, size=22, bold=True)
    run(para(doc, align=WD_ALIGN_PARAGRAPH.CENTER), "宗教史的故事接力賽", size=12, color=grey)

    doc.add_page_break()
    run(para(doc, after=18), "本卷目次", size=16, bold=True)
    for no in chapters:
        title, sub, _, _ = read_chapter(no)
        p = para(doc, after=10)
        run(p, title, size=13, bold=True)
        run(p, "　" + sub, size=11, color=grey)

    # ── 正文 ──────────────────────────────────────────────
    for no in chapters:
        title, sub, blocks, notes = read_chapter(no)
        doc.add_page_break()
        run(para(doc, after=4, before=24), title, size=18, bold=True)
        run(para(doc, after=24), sub, size=12, color=grey)
        for name, paras in blocks:
            if name not in ("導言",):
                run(para(doc, after=8, before=18), name, size=14, bold=True)
            for text in paras:
                add_body(doc, fn, text, notes)

    fn.save()                       # 一定要在 doc.save() 之前
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"千面上帝　第{vol_no}卷　{vol_name}.docx"
    doc.save(str(path))
    return path, len(fn.items)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--volume", type=int, help="只出某一卷（1-7）")
    ap.add_argument("--out", default=str(DRIVE), help="輸出資料夾")
    a = ap.parse_args()

    # 卷名直接讀來源包，不要在這裡再抄一次目錄
    vols = {}
    for f in sorted(SRC.glob("ch*.json")):
        d = json.loads(f.read_text(encoding="utf-8"))
        vols.setdefault(d["volume"], []).append(d["no"])

    out_dir = Path(a.out)
    for i, (vol, chs) in enumerate(vols.items(), 1):
        if a.volume and i != a.volume:
            continue
        missing = [n for n in chs if not (CH / f"ch{n:02d}.md").exists()]
        if missing:
            print(f"  ✗ 第{i}卷 {vol}：第 {missing} 章還沒寫完，跳過")
            continue
        name = vol.split("：", 1)[-1]
        path, n = build_volume(i, name, chs, out_dir)
        print(f"  ✓ 第{i}卷 {name}：第 {chs[0]}–{chs[-1]} 章、{n} 個頁下註 → {path}")


if __name__ == "__main__":
    main()
