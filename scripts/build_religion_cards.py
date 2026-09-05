#!/usr/bin/env python3
"""可裁切的印刷宗教卡：世界宗教十六張。

版面、卡框、雙面鏡像全部沿用單字卡那一套（`build_flashcards`），同一把裁刀通吃，
所以這裡不重算任何尺寸——只換卡面。

跟撲克牌（`build_playing_cards`）的差別是**正反面的意義**：撲克牌背面統一才玩得了，
這副是教具，正面問、背面答，跟單字卡同一個路數。

課堂用法決定了兩件設計：

1. **正面不能洩漏答案。** 十六張要讓學生先分成四堆（泛神論／多神論／一神論／
   實用神論）再翻背面對，所以正面一律深灰框，四色只出現在背面。
   講義章次同理印在背面——第十一章就等於實用神論，印在正面就破梗了。
2. **沒有符號的傳統就留白。** 原住民傳統、儒教、無神論沒有自己採用的符號，
   硬配一張鼓或一個原子只是把刻板印象印上去。留白是誠實的，也是可以講的一課。
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))

import build_flashcards as base  # noqa: E402
import build_playing_cards as cards  # noqa: E402  （借它的 OpenMoji 載入與查圖）

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data/teachingCards/world-religions.json"
OUTPUT_DIR = ROOT / "output/teaching-cards"

FRONT_COLOR = "4A4A4A"      # 正面一律深灰：四色分類是答案，不能印在題目上
ICON_MM = 24
TRANSLIT_PT = 8.5
VIEW_PT = 10
FIELD_PT = 8
LABEL_PT = 11
NAME_BACK_PT = 10.5
FOOT_PT = 7.5

# 「臺灣原住民族傳統宗教」十一個字與「道教」兩個字要並存，字級按長度縮
NAME_STEPS = ((3, 22), (4, 19), (6, 17), (8, 14), (10, 12), (99, 10.5))


def name_size(text: str) -> float:
    for limit, size in NAME_STEPS:
        if len(text) <= limit:
            return size
    return NAME_STEPS[-1][1]


def fill_front(cell, card: dict, icon: Path | None, place) -> None:
    cell = base.framed(cell, FRONT_COLOR, place)
    if icon:
        picture = cell.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.paragraph_format.space_after = Pt(6)
        picture.add_run().add_picture(str(icon), width=Mm(ICON_MM))
    else:
        base.blank(cell, 16)
    name = cell.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(4)
    base.write(name, card["name"], base.FONT_ZH, name_size(card["name"]))
    translit = cell.add_paragraph()
    translit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    translit.paragraph_format.space_after = Pt(0)
    base.write(translit, card["translit"], cards.FONT_LATIN, TRANSLIT_PT, color=base.MUTED)


def fill_back(cell, card: dict, category: dict, place) -> None:
    cell = base.framed(cell, category["color"], place)
    # 背面也要印名字：翻開的牌散在桌上時，只有「泛神論」認不出是哪一張
    name = cell.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(2)
    base.write(name, card["name"], base.FONT_ZH, NAME_BACK_PT)

    label = cell.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.space_after = Pt(5)
    base.write(label, card["category"], base.FONT_ZH, LABEL_PT, color=category["color"])

    view = cell.add_paragraph()
    view.alignment = WD_ALIGN_PARAGRAPH.CENTER
    view.paragraph_format.space_after = Pt(7)
    base.write(view, card["view"], base.FONT_ZH, VIEW_PT)

    for field, value in (("起源", card["origin"]), ("經典", card["scripture"]),
                         ("人數", card["adherents"])):
        line = cell.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.paragraph_format.space_after = Pt(3)
        base.write(line, f"{field}｜{value}", base.FONT_UI, FIELD_PT)

    foot = cell.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.paragraph_format.space_after = Pt(0)
    base.write(foot, f"講義第 {card['chapter']} 章", base.FONT_UI, FOOT_PT, color=base.MUTED)


def add_cover(document: Document, deck: dict, total: int, with_icon: int) -> None:
    colors = "、".join(f"{name}{info['colorName']}" for name, info in deck["categories"].items())
    lines = [
        (deck["title"], 30, base.INK, 14),
        (f"{deck['subtitle']}　{total} 張・{-(-total // 8)} 組雙面", 13, base.MUTED, 22),
        ("列印：A4 橫式，雙面列印選「沿長邊翻頁」，縮放 100%（不要選「符合頁面大小」）。",
         11, base.INK, 6),
        (f"裁切：不印裁切線。成品每張 {base.CARD_W_MM:.2f}×{base.CARD_H_MM:.0f} mm，"
         f"卡框比裁切線內縮 {base.FRAME_INSET_MM:.0f} mm，裁歪一兩毫米只會讓白邊不等寬。",
         11, base.INK, 6),
        (deck["note"], 11, base.INK, 6),
        (f"背面框色即分類：{colors}。四堆分完再翻面，顏色一致就是分對了。", 11, base.INK, 6),
        (f"插圖：{with_icon} 張有符號，其餘留白——原住民傳統、儒教、無神論沒有自己採用的符號，"
         "硬配一張圖只會把刻板印象印上去。哪些傳統用符號代表自己、哪些不用，本身就是一課。",
         11, base.INK, 22),
        ("圖片來源：Noto Emoji 舊版（Apache 2.0），經 Iconify 取得。整副統一用同一套繪者的畫法，"
         "免得有的圖帶底色方框、有的留白，擺在一起像貼錯。", 9.5, base.MUTED, 4),
    ]
    lines.extend((text, 9.5, base.MUTED, 4) for text in deck["sources"])
    for text, size, color, after in lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.space_before = Pt(0)
        base.write(paragraph, text, base.FONT_UI, size, color=color, bold=size >= 30)
    base.page_break(document)
    blank = document.add_paragraph()
    blank.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.write(blank, "（此頁留白，供雙面列印對齊）", base.FONT_UI, 9.5, color=base.MUTED)


def build(deck: dict, icons: dict[str, Path | None], output: Path) -> Path:
    document = Document()
    base.configure(document)
    add_cover(document, deck, len(deck["cards"]), sum(1 for v in icons.values() if v))

    per_page = base.COLS * base.ROWS
    for start in range(0, len(deck["cards"]), per_page):
        page = deck["cards"][start : start + per_page]
        for side in ("front", "back"):
            base.page_break(document)
            table = base.new_grid(document)
            for offset, card in enumerate(page):
                row, column = divmod(offset, base.COLS)
                target_column = column if side == "front" else base.COLS - 1 - column
                place = (target_column, row, 1000 + start * 2 + offset * 2 + (side == "back"))
                cell = table.cell(row, target_column)
                if side == "front":
                    fill_front(cell, card, icons[card["name"]], place)
                else:
                    fill_back(cell, card, deck["categories"][card["category"]], place)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


def main() -> None:
    parser = argparse.ArgumentParser(description="產生宗教卡（A4 橫式、每頁 8 張、雙面）")
    parser.add_argument("--output", default="world-religion-cards.docx")
    args = parser.parse_args()

    deck = json.loads(DATA.read_text("utf-8"))
    index = cards.openmoji()
    icons = {card["name"]: (cards.icon_path(index, card["icon"]) if card["icon"] else None)
             for card in deck["cards"]}

    unknown = {card["category"] for card in deck["cards"]} - set(deck["categories"])
    if unknown:
        raise SystemExit(f"卡片用了沒定義的分類：{sorted(unknown)}")

    path = build(deck, icons, OUTPUT_DIR / args.output)
    tally = {name: sum(1 for c in deck["cards"] if c["category"] == name)
             for name in deck["categories"]}
    print(f"  {deck['title']}：{len(deck['cards'])} 張，"
          f"正反共 {-(-len(deck['cards']) // 8) * 2} 頁（每頁 8 張，"
          f"{base.CARD_W_MM:.2f}×{base.CARD_H_MM:.0f} mm）")
    print("  " + "　".join(f"{k} {v} 張" for k, v in tally.items()))
    print(f"  有符號 {sum(1 for v in icons.values() if v)}，留白 {sum(1 for v in icons.values() if not v)}")
    print(path)


if __name__ == "__main__":
    main()
