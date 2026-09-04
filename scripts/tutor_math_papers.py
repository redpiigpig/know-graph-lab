# -*- coding: utf-8 -*-
"""家教（私中入學數學）試卷產生器。

三種產出，全部落在 Drive，不進版控：
  1. B5 題目卷（學生用，無答案）　.../家教_私中入學數學/複習卷/複習卷（N）XXX.docx
  2. B5 詳解答案卷（教師用，逐題列式）　.../複習卷（N）XXX　詳解.docx
  3. A4 一頁備考說明（單元清單＋教材建議）　.../私中數學備考說明.docx

版面刻意比照曙光的實況：全卷 35 分鐘、選擇題為主（正式考試劃答案卡），
末四題非選是家教自己加的，用來看解題過程，正式考試不考。

用法：python scripts/tutor_math_papers.py
"""
import hashlib
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DRIVE = Path(r'G:\我的雲端硬碟\資料\知識圖工作室\教學\家教_私中入學數學')
KAI, MING, HEI = 'DFKai-SB', 'PMingLiU', 'Microsoft JhengHei'
GRAY = RGBColor(0x5A, 0x5A, 0x5A)
NAVY = RGBColor(0x1F, 0x39, 0x64)
OPT = 'ABCD'

# ── 題庫 ────────────────────────────────────────────────────────────────────
# 每題 (題幹, [四選項], 正解索引, 詳解)。這裡寫的是「原始順序」的正解索引，
# 實際印出來的選項順序由 shuffled() 依內容雜湊決定 —— 出題時很容易把正解都放在
# 同一個位置（初版 16 題有 9 題落在 B），孩子不必讀題就能猜。改成吃內容的雜湊後：
#   ‧ 同一題每次產生的順序都一樣，題目卷與詳解卷才對得起來
#   ‧ 正解會散在 A–D，看不出規律

PAPER1 = dict(
    no='一', name='小五範圍總複習',
    scope='整數四則與併式‧因數倍數‧分數加減乘‧小數乘除‧面積‧體積與容積‧'
          '比率與百分率‧平均‧時間‧單位換算',
    mc=[
        ('計算 25 × 4 － 36 ÷ 6 × 3 ＝ ？',
         ['82', '88', '64', '246'], 0,
         '先乘除後加減，乘除同級由左而右：36 ÷ 6 ＝ 6，6 × 3 ＝ 18；'
         '25 × 4 ＝ 100。100 － 18 ＝ 82。'),
        ('有紅色色紙 36 張、藍色色紙 48 張，要平分給若干人，兩種色紙都不剩，'
         '最多可以分給幾個人？',
         ['6 人', '12 人', '16 人', '144 人'], 1,
         '「都不剩、最多」＝求最大公因數。36 ＝ 2×2×3×3，48 ＝ 2×2×2×2×3，'
         '公同的是 2×2×3 ＝ 12。答 12 人（每人紅 3 張、藍 4 張）。'),
        ('3/4 ＋ 5/6 ＝ ？',
         ['8/10', '1又7/12', '15/24', '2又1/12'], 1,
         '4 和 6 的最小公倍數是 12，通分：3/4 ＝ 9/12，5/6 ＝ 10/12。'
         '9/12 ＋ 10/12 ＝ 19/12 ＝ 1又7/12。'),
        ('2.4 × 0.35 ＝ ？',
         ['0.084', '0.84', '8.4', '84'], 1,
         '先當整數算：24 × 35 ＝ 840。兩個乘數小數點後共 1 ＋ 2 ＝ 3 位，'
         '所以是 0.840 ＝ 0.84。'),
        ('一個梯形的上底 6 公分、下底 10 公分、高 8 公分，面積是多少平方公分？',
         ['48', '64', '80', '128'], 1,
         '梯形面積 ＝（上底 ＋ 下底）× 高 ÷ 2 ＝（6 ＋ 10）× 8 ÷ 2 '
         '＝ 16 × 8 ÷ 2 ＝ 64 平方公分。'),
        ('一個三角形的面積是 42 平方公分，底邊長 12 公分，這個底邊上的高是幾公分？',
         ['3.5 公分', '7 公分', '14 公分', '28 公分'], 1,
         '三角形面積 ＝ 底 × 高 ÷ 2，所以 高 ＝ 面積 × 2 ÷ 底 '
         '＝ 42 × 2 ÷ 12 ＝ 84 ÷ 12 ＝ 7 公分。'),
        ('一個長方體水箱，內部長 20 公分、寬 15 公分、高 10 公分，'
         '裝滿水時是多少公升？',
         ['0.3 公升', '3 公升', '30 公升', '300 公升'], 1,
         '容積 ＝ 20 × 15 × 10 ＝ 3000 立方公分。1 公升 ＝ 1000 立方公分，'
         '所以是 3 公升。'),
        ('一件外套原價 1200 元，週年慶打七五折，要付多少元？',
         ['840 元', '900 元', '960 元', '1125 元'], 1,
         '七五折就是原價的 75％。1200 × 0.75 ＝ 900 元。'
         '（提醒：台灣的「七五折」＝ 付 75％，不是折掉 75％。）'),
        ('全班 40 人，其中男生 24 人。男生占全班的百分之幾？',
         ['24％', '40％', '60％', '64％'], 2,
         '24 ÷ 40 ＝ 0.6 ＝ 60％。'),
        ('下列哪一個數同時是 3 的倍數，也是 4 的倍數？',
         ['114', '126', '132', '146'], 2,
         '3 的倍數看各位數字和、4 的倍數看末兩位。132：1＋3＋2 ＝ 6 是 3 的倍數，'
         '末兩位 32 ÷ 4 ＝ 8 整除。其餘三個各只符合一個條件或都不符合。'),
        ('一條繩子長 4/5 公尺，用掉了其中的 3/4，用掉幾公尺？',
         ['1/5 公尺', '3/5 公尺', '3/20 公尺', '16/15 公尺'], 1,
         '「某數的幾分之幾」用乘法：4/5 × 3/4 ＝ 12/20 ＝ 3/5 公尺。'),
        ('8.4 ÷ 0.7 ＝ ？',
         ['1.2', '12', '120', '0.12'], 1,
         '除數與被除數同時乘 10，把除數變成整數：84 ÷ 7 ＝ 12。'),
        ('一部電影 14:35 開演，片長 1 小時 55 分，幾點幾分結束？',
         ['15:30', '16:20', '16:30', '16:40'], 2,
         '14:35 ＋ 1 小時 ＝ 15:35；再加 55 分，35 ＋ 55 ＝ 90 分 ＝ 1 小時 30 分，'
         '所以是 16:30。'),
        ('小美五次小考的分數是 82、90、76、88、94 分，平均幾分？',
         ['84 分', '85 分', '86 分', '88 分'], 2,
         '總分 ＝ 82 ＋ 90 ＋ 76 ＋ 88 ＋ 94 ＝ 430，430 ÷ 5 ＝ 86 分。'),
        ('下列哪一種圖形有無限多條對稱軸？',
         ['正方形', '長方形', '圓形', '正三角形'], 2,
         '正方形 4 條、長方形 2 條、正三角形 3 條；圓形通過圓心的每一條直線'
         '都是對稱軸，有無限多條。'),
        ('一塊農地面積 0.8 公頃，等於多少平方公尺？',
         ['80 平方公尺', '800 平方公尺', '8000 平方公尺', '80000 平方公尺'], 2,
         '1 公頃 ＝ 10000 平方公尺，0.8 × 10000 ＝ 8000 平方公尺。'),
    ],
    fr=[
        ('有一張長 60 公分、寬 45 公分的長方形色紙，要剪成大小相同的正方形'
         '且不能剩下。正方形的邊長最長是幾公分？這時共可剪出幾張？',
         '邊長要能同時整除 60 和 45，而且要最長 → 求最大公因數。\n'
         '60 ＝ 2×2×3×5，45 ＝ 3×3×5，最大公因數 ＝ 3×5 ＝ 15。\n'
         '邊長最長 15 公分。\n'
         '一排可剪 60 ÷ 15 ＝ 4 張，共有 45 ÷ 15 ＝ 3 排，4 × 3 ＝ 12 張。\n'
         '答：邊長 15 公分，共 12 張。'),
        ('一桶油，第一次用掉全桶的 1/4，第二次用掉「剩下的」2/3。'
         '請問還剩下全桶的幾分之幾？',
         '陷阱在「剩下的 2/3」不是「全桶的 2/3」。\n'
         '第一次用後還剩 1 － 1/4 ＝ 3/4。\n'
         '第二次用掉 3/4 × 2/3 ＝ 6/12 ＝ 1/2（全桶的一半）。\n'
         '還剩 3/4 － 1/2 ＝ 3/4 － 2/4 ＝ 1/4。\n'
         '（也可以直接算：剩下的是 3/4 的 1/3，3/4 × 1/3 ＝ 1/4。）\n'
         '答：還剩全桶的 1/4。'),
        ('一塊梯形花圃，上底 8 公尺、下底 14 公尺、高 6 公尺。'
         '若每 1 平方公尺鋪草皮要 250 元，鋪滿整塊花圃共要多少元？',
         '面積 ＝（8 ＋ 14）× 6 ÷ 2 ＝ 22 × 6 ÷ 2 ＝ 132 ÷ 2 ＝ 66 平方公尺。\n'
         '費用 ＝ 66 × 250 ＝ 16500 元。\n'
         '答：66 平方公尺，共 16500 元。'),
        ('一件衣服標價 2000 元。店家先打八折，結帳時再折抵 100 元。'
         '實際付了多少元？實付金額是標價的百分之幾？',
         '打八折：2000 × 0.8 ＝ 1600 元。\n'
         '再折 100：1600 － 100 ＝ 1500 元。\n'
         '1500 ÷ 2000 ＝ 0.75 ＝ 75％。\n'
         '答：實付 1500 元，是標價的 75％。'),
    ])

PAPER2 = dict(
    no='二', name='小六範圍總複習',
    scope='分數除法‧比與比值‧比例尺‧圓周長與圓面積‧柱體體積與表面積‧速率‧'
          '成數與折扣‧正比反比‧統計與機率‧簡易方程式‧放大縮小',
    mc=[
        ('3/4 ÷ 2/5 ＝ ？',
         ['3/10', '6/20', '1又7/8', '2又2/3'], 2,
         '除以一個分數 ＝ 乘它的倒數：3/4 × 5/2 ＝ 15/8 ＝ 1又7/8。'),
        ('甲數與乙數的比是 3 : 5。若甲數是 24，乙數是多少？',
         ['14.4', '30', '40', '45'], 2,
         '3 : 5 ＝ 24 : ?。24 ÷ 3 ＝ 8（一份是 8），乙 ＝ 5 × 8 ＝ 40。'),
        ('一個圓的半徑是 7 公分，圓周長是多少公分？（圓周率取 3.14）',
         ['21.98 公分', '43.96 公分', '153.86 公分', '87.92 公分'], 1,
         '圓周長 ＝ 2 × 圓周率 × 半徑 ＝ 2 × 3.14 × 7 ＝ 43.96 公分。'
         '（153.86 是它的面積，別選錯。）'),
        ('一個圓的直徑是 20 公分，面積是多少平方公分？（圓周率取 3.14）',
         ['62.8', '314', '1256', '400'], 1,
         '半徑 ＝ 20 ÷ 2 ＝ 10。面積 ＝ 3.14 × 10 × 10 ＝ 314 平方公分。'
         '（面積公式吃的是半徑，不是直徑。）'),
        ('一輛車以時速 60 公里行駛，2 小時 30 分可以走多少公里？',
         ['120 公里', '130 公里', '150 公里', '180 公里'], 2,
         '2 小時 30 分 ＝ 2.5 小時。距離 ＝ 速率 × 時間 ＝ 60 × 2.5 ＝ 150 公里。'
         '（30 分要換成 0.5 小時，不是 0.3。）'),
        ('地圖的比例尺是 1 : 50000。圖上量得兩地相距 4 公分，實際相距多少公里？',
         ['0.2 公里', '2 公里', '20 公里', '200 公里'], 1,
         '實際 ＝ 4 × 50000 ＝ 200000 公分 ＝ 2000 公尺 ＝ 2 公里。'),
        ('一個圓柱，底面半徑 5 公分、高 10 公分，體積是多少立方公分？'
         '（圓周率取 3.14）',
         ['157', '314', '785', '1570'], 2,
         '柱體體積 ＝ 底面積 × 高。底面積 ＝ 3.14 × 5 × 5 ＝ 78.5，'
         '78.5 × 10 ＝ 785 立方公分。'),
        ('某商品的進貨成本是 800 元，老闆想賺兩成，售價要訂多少元？',
         ['820 元', '960 元', '1000 元', '1600 元'], 1,
         '兩成 ＝ 20％。售價 ＝ 800 ×（1 ＋ 0.2）＝ 800 × 1.2 ＝ 960 元。'),
        ('下列哪一組是「成反比」的關係？',
         ['買同一種筆的枝數與總價', '正方形的邊長與周長',
          '長方形面積固定時的長與寬', '走路速率固定時的時間與距離'], 2,
         '兩量相乘是定值就成反比。長 × 寬 ＝ 面積（固定），長變大寬就變小。'
         '其餘三組都是一個變大、另一個跟著等倍變大，屬於正比。'),
        ('某校對 600 名學生做調查，圓形圖上「最喜歡籃球」占 35％，有幾人？',
         ['165 人', '180 人', '210 人', '350 人'], 2,
         '600 × 35％ ＝ 600 × 0.35 ＝ 210 人。'),
        ('甲數是乙數的 1.25 倍。若甲數是 45，乙數是多少？',
         ['30', '36', '40', '56.25'], 1,
         '甲 ＝ 乙 × 1.25，所以 乙 ＝ 45 ÷ 1.25 ＝ 36。'
         '（比較量除以倍數才是基準量；45 × 1.25 ＝ 56.25 是做反了。）'),
        ('解方程式 3x ＋ 7 ＝ 25，x ＝ ？',
         ['4', '6', '8', '32/3'], 1,
         '兩邊同減 7：3x ＝ 18；兩邊同除以 3：x ＝ 6。'),
        ('觀察這一列數：1、4、9、16、25、…… 第 10 個數是多少？',
         ['64', '81', '100', '121'], 2,
         '第 n 個數就是 n × n（1×1、2×2、3×3……）。第 10 個 ＝ 10 × 10 ＝ 100。'),
        ('袋子裡有紅球 3 顆、白球 5 顆、藍球 2 顆。任意摸出一顆，'
         '摸到白球的機率是多少？',
         ['1/5', '3/10', '1/2', '5/8'], 2,
         '球共 3 ＋ 5 ＋ 2 ＝ 10 顆，白球 5 顆，機率 ＝ 5/10 ＝ 1/2。'),
        ('一個正方體的邊長是 6 公分，表面積是多少平方公分？',
         ['36', '144', '216', '432'], 2,
         '正方體有 6 個相同的正方形面。一面 ＝ 6 × 6 ＝ 36，'
         '表面積 ＝ 36 × 6 ＝ 216 平方公分。（216 立方公分是它的體積，單位別混。）'),
        ('把一個圖形的每一邊都放大為原來的 3 倍，面積會變成原來的幾倍？',
         ['3 倍', '6 倍', '9 倍', '27 倍'], 2,
         '長和寬各變 3 倍，面積 ＝ 3 × 3 ＝ 9 倍。'
         '（27 倍是體積會變的倍數，考題很愛拿這個混淆。）'),
    ],
    fr=[
        ('甲、乙兩地相距 240 公里。A 車從甲地出發，時速 70 公里；'
         'B 車同時從乙地出發，時速 50 公里，兩車相向而行。'
         '幾小時後相遇？相遇的地點距離甲地多少公里？',
         '相向而行時，兩車「合起來」每小時拉近 70 ＋ 50 ＝ 120 公里。\n'
         '相遇時間 ＝ 240 ÷ 120 ＝ 2 小時。\n'
         '相遇點距甲地 ＝ A 車走的距離 ＝ 70 × 2 ＝ 140 公里。\n'
         '（檢查：B 車走 50 × 2 ＝ 100，140 ＋ 100 ＝ 240，對得起來。）\n'
         '答：2 小時後相遇，距甲地 140 公里。'),
        ('一個正方形，邊長 10 公分，在裡面畫一個最大的圓。'
         '求這個圓的面積，以及正方形內、圓以外部分的面積。（圓周率取 3.14）',
         '最大的圓，直徑就等於正方形的邊長 10 公分，所以半徑 ＝ 5 公分。\n'
         '圓面積 ＝ 3.14 × 5 × 5 ＝ 78.5 平方公分。\n'
         '正方形面積 ＝ 10 × 10 ＝ 100 平方公分。\n'
         '圓以外 ＝ 100 － 78.5 ＝ 21.5 平方公分。\n'
         '答：圓 78.5 平方公分，圓以外 21.5 平方公分。'),
        ('某商品成本 1500 元，標價比成本高四成。後來以標價的八折賣出，'
         '請問售價多少元？這筆買賣是賺還是賠？金額多少？',
         '標價 ＝ 1500 ×（1 ＋ 0.4）＝ 1500 × 1.4 ＝ 2100 元。\n'
         '售價 ＝ 2100 × 0.8 ＝ 1680 元。\n'
         '1680 － 1500 ＝ 180，售價高於成本。\n'
         '答：售價 1680 元，賺 180 元。\n'
         '（常見錯誤：以為「加四成再打八折」等於加兩成。其實 1.4 × 0.8 ＝ 1.12，'
         '只賺一成二。）'),
        ('把 84 顆糖果依照 2 : 3 : 2 的比分給甲、乙、丙三人，各得幾顆？',
         '總份數 ＝ 2 ＋ 3 ＋ 2 ＝ 7 份。\n'
         '一份 ＝ 84 ÷ 7 ＝ 12 顆。\n'
         '甲 ＝ 2 × 12 ＝ 24 顆，乙 ＝ 3 × 12 ＝ 36 顆，丙 ＝ 2 × 12 ＝ 24 顆。\n'
         '（檢查：24 ＋ 36 ＋ 24 ＝ 84，對得起來。）\n'
         '答：甲 24 顆、乙 36 顆、丙 24 顆。'),
    ])


# ── docx 工具 ───────────────────────────────────────────────────────────────
def run_ea(par, text, font=MING, size=11.0, bold=False, color=None):
    r = par.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font)
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    return r


def _h(s):
    return hashlib.md5(s.encode('utf-8')).hexdigest()


def targets(spec):
    """每卷 16 題的正解位置：A、B、C、D 各出現 4 次，順序由卷別雜湊打散。

    純粹靠選項文字的雜湊排序不夠 —— 實測會出現某一卷 8 題正解都在 A、
    一題 D 都沒有。這裡先把「四個字母各四次」的名單洗開當目標位置，
    再把正解搬過去，分布才真的平均。
    """
    n = len(spec['mc'])
    slots = (list(range(4)) * (n // 4 + 1))[:n]
    order = sorted(range(n), key=lambda k: _h(f'{spec["no"]}-slot-{k}'))
    return [slots[k] for k in order]


def shuffled(spec, i, opts, ans):
    """把正解排到 targets() 指定的位置，其餘三個選項依雜湊填滿剩下的格子。"""
    tgt = targets(spec)[i - 1]
    others = sorted((o for j, o in enumerate(opts) if j != ans),
                    key=lambda t: _h(f'{spec["no"]}-{i}||{t}'))
    out = []
    for k in range(4):
        out.append(opts[ans] if k == tgt else others.pop(0))
    return out, tgt


def new_doc(w_cm, h_cm, margin_cm, base_pt=11):
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Cm(w_cm), Cm(h_cm)
    s.top_margin = s.bottom_margin = Cm(margin_cm)
    s.left_margin = s.right_margin = Cm(margin_cm)
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'
    st.element.rPr.rFonts.set(qn('w:eastAsia'), MING)
    st.font.size = Pt(base_pt)
    st.paragraph_format.space_after = Pt(0)
    return doc


def footer(doc, label):
    par = doc.sections[0].footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ea(par, f'{label}　－ ', MING, 8.5, color=GRAY)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    par._p.append(fld)
    run_ea(par, ' －', MING, 8.5, color=GRAY)


def para(doc, indent=0.0, hang=0.0, before=0, after=0, spacing=1.0, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if indent:
        pf.left_indent = Cm(indent)
    if hang:
        pf.first_line_indent = Cm(-hang)
    pf.space_before, pf.space_after = Pt(before), Pt(after)
    pf.line_spacing = spacing
    if align:
        p.alignment = align
    return p


# ── B5 試卷 ────────────────────────────────────────────────────────────────
def build_paper(spec, with_answers):
    doc = new_doc(18.2, 25.7, 1.7, base_pt=11)      # JIS B5
    title = f'私中入學數學　複習卷（{spec["no"]}）'

    p = para(doc, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    run_ea(p, title, KAI, 17, bold=True, color=NAVY)
    p = para(doc, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    run_ea(p, spec['name'], KAI, 12, color=GRAY)

    if with_answers:
        p = para(doc, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
        run_ea(p, '【教師用‧詳解答案卷】', HEI, 11, bold=True, color=NAVY)
        p = para(doc, after=6, spacing=1.2)
        run_ea(p, '範圍：', HEI, 9, bold=True, color=GRAY)
        run_ea(p, spec['scope'], MING, 9, color=GRAY)
        # 答案總表：改卷時先看這一列，不必翻整份
        p = para(doc, after=10, spacing=1.3)
        run_ea(p, '選擇題答案　', HEI, 10, bold=True)
        run_ea(p, '　'.join(
            f'{i}.{OPT[shuffled(spec, i, q[1], q[2])[1]]}'
            for i, q in enumerate(spec['mc'], 1)), MING, 10)
    else:
        p = para(doc, after=6)
        run_ea(p, '姓名：＿＿＿＿＿＿＿＿　　日期：＿＿＿＿＿＿＿＿　　'
                  '得分：＿＿＿＿＿＿', MING, 11)
        p = para(doc, after=10, spacing=1.25)
        run_ea(p, '作答時間 35 分鐘（比照曙光國中部數理科一節的長度）。'
                  '第一部分 16 題、每題 5 分；第二部分 4 題、每題 5 分，'
                  '要寫出算式。分數請寫成「3/4」、「1又7/12」的形式。',
               MING, 9.5, color=GRAY)

    p = para(doc, before=2, after=6)
    run_ea(p, '第一部分　選擇題（每題 5 分，共 80 分）', HEI, 11.5, bold=True)

    for i, (stem, opts, ans, why) in enumerate(spec['mc'], 1):
        opts, ans = shuffled(spec, i, opts, ans)
        p = para(doc, indent=1.05, hang=1.05, before=6, after=2, spacing=1.3)
        run_ea(p, f'（{OPT[ans]}）{i}. ' if with_answers else f'（　　）{i}. ',
               MING, 11, bold=with_answers)
        run_ea(p, stem, MING, 11)
        po = para(doc, indent=1.05, after=1, spacing=1.25)
        for j, o in enumerate(opts):
            bold = with_answers and j == ans
            run_ea(po, f'({OPT[j]}) ', MING, 10.5, bold=bold)
            run_ea(po, o + ('　　' if j < 3 else ''), MING, 10.5, bold=bold)
        if with_answers:
            pw = para(doc, indent=1.05, before=2, after=2, spacing=1.25)
            run_ea(pw, '詳解：', HEI, 9.5, bold=True, color=GRAY)
            run_ea(pw, why, MING, 9.5, color=GRAY)

    p = para(doc, before=14, after=6)
    run_ea(p, '第二部分　計算與應用（每題 5 分，共 20 分，須列出算式）',
           HEI, 11.5, bold=True)
    if not with_answers:
        p = para(doc, after=8, spacing=1.2)
        run_ea(p, '※ 正式入學考全部劃答案卡、不考這一部分；'
                  '這裡出四題是要看你怎麼想，過程比答案重要。', MING, 9, color=GRAY)

    for i, (stem, sol) in enumerate(spec['fr'], 1):
        p = para(doc, indent=0.85, hang=0.85, before=8, after=3, spacing=1.35)
        run_ea(p, f'{i}. ', MING, 11, bold=True)
        run_ea(p, stem, MING, 11)
        if with_answers:
            for line in sol.split('\n'):
                pw = para(doc, indent=1.5, after=1, spacing=1.3)
                run_ea(pw, line, MING, 10, color=GRAY)
        else:
            for _ in range(5):
                para(doc, after=0, spacing=1.9)

    if not with_answers:
        p = para(doc, before=16, spacing=1.25)
        run_ea(p, '※ 寫完先自己檢查一遍再交。不會的題目直接跳過，'
                  '不要卡在同一題上——正式考試 35 分鐘要寫完整張，'
                  '平均一題只有一分多鐘。', MING, 9.5, color=GRAY)

    footer(doc, title + ('　詳解' if with_answers else ''))
    out = DRIVE / '複習卷' / (
        f'複習卷（{spec["no"]}）{spec["name"]}' +
        ('　詳解' if with_answers else '') + '.docx')
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


# ── A4 一頁備考說明 ─────────────────────────────────────────────────────────
GRADE5 = [
    ('數與計算', '整數四則與併式（先乘除後加減、括號）、多位小數乘除、概數'),
    ('因數倍數', '質數、質因數分解、最大公因數與最小公倍數的應用題'),
    ('分數', '擴分約分、通分、異分母加減、分數乘法、假分數帶分數互換'),
    ('量與實測', '面積（平行四邊形、三角形、梯形）、長方體體積與容積、'
                 '時間計算、公頃等大單位'),
    ('關係統計', '比率與百分率、折扣、平均數、長條圖與折線圖、線對稱'),
]
GRADE6 = [
    ('數與計算', '分數除法、小數與分數混合計算、四則的整體策略'),
    ('比與比例', '比值、最簡整數比、比例式、比例尺、按比分配'),
    ('圓與立體', '圓周長、圓面積、扇形、角柱圓柱體積與表面積、複合圖形'),
    ('關係應用', '速率（相遇追及）、基準量與比較量、成數折扣、正比反比、'
                 '放大縮小（面積變平方倍）'),
    ('代數統計', '用符號表示數、簡易方程式、圓形圖、規律問題、機率初步'),
]
MATERIALS = [
    ('① 把課內補滿（最優先）',
     '學校課本＋習作，加一本同版本自修（南一新超群／康軒新挑戰／翰林小無敵，'
     '看孩子學校用哪一版）。考題再怎麼變，八成分數還是課內範圍。'),
    ('② 跨冊總複習',
     '一本國小數學總複習（上述三家都有出，六上開始用剛好）：把六冊打散重編成'
     '單元，複習時不必翻六本課本。'),
    ('③ 私中考古題（已下載，在同一個 Drive 資料夾）',
     '縣市學力檢測 15 份（官方、附答案，當診斷與單元卷）；延和國中 13 份'
     '（數理合卷，格式最像曙光）；協同中學 7 份（私中自命題風格、較難、無答案）。'),
    ('④ 進階題感（行有餘力）',
     '書店的「私中入學模擬試題」題本各家都有，挑選只有一個硬條件：附詳解。'
     '補洞可搭配免費的均一教育平台。'),
]
TIMELINE = [
    ('現在～小五結束', '主攻五年級單元，因數倍數、分數、面積三塊練到不用想；'
                       '每週一份計時小卷。'),
    ('小六上學期', '邊上新課邊補，圓、比例、速率是六年級最吃分的三塊；'
                   '開始寫整份考古題。'),
    ('小六下 1–2 月', '全範圍總複習，每週一份 35 分鐘完整模擬，練配速與劃卡。'),
    ('小六下 3 月', '報名曙光的小六適性學力大會考，用真實考場走一遍，'
                    '回來針對失分處收尾。'),
]


def cell_para(cell, first=False, indent=0.0, hang=0.0, after=1, spacing=1.15):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    pf = p.paragraph_format
    if indent:
        pf.left_indent = Cm(indent)
    if hang:
        pf.first_line_indent = Cm(-hang)
    pf.space_before, pf.space_after = Pt(0), Pt(after)
    pf.line_spacing = spacing
    return p


def build_guide():
    doc = new_doc(21.0, 29.7, 1.5, base_pt=9)      # A4，壓成一頁
    p = para(doc, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    run_ea(p, '準備私中數學：要上哪些單元、要備哪些教材', KAI, 15,
           bold=True, color=NAVY)
    p = para(doc, after=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    run_ea(p, '目標校：新竹市私立曙光女子高級中學國中部', KAI, 10, color=GRAY)

    p = para(doc, after=3)
    run_ea(p, '一、先認清這場考試的形狀', HEI, 11, bold=True, color=NAVY)
    for line in [
        '每年兩關：3 月中週六「小六適性學力大會考」（對外模擬考，可先報名'
        '試水溫）、4 月中週日「新生入學測驗暨抽籤」（正式）。',
        '四科只分兩節：數理科 08:40–09:50（09:15 換科）、語文科 10:20–11:30'
        '——數學實際只有約 35 分鐘。全部選擇題、2B 鉛筆劃答案卡，'
        '題本與答案卡一律收回，所以曙光自己沒有考古題流出。',
        '結論：這場考試考的是熟練與速度，不是難題。練習一定要計時，'
        '並且練跳題與檢查劃卡。',
    ]:
        q = para(doc, indent=0.6, hang=0.6, after=1, spacing=1.2)
        run_ea(q, '‧ ', MING, 9)
        run_ea(q, line, MING, 9)

    p = para(doc, before=7, after=3)
    run_ea(p, '二、要複習的單元清單', HEI, 11, bold=True, color=NAVY)
    t = doc.add_table(rows=2, cols=2)
    t.style = 'Table Grid'
    for col, (label, items) in enumerate((('五年級', GRADE5), ('六年級', GRADE6))):
        h = cell_para(t.cell(0, col), first=True, after=0)
        run_ea(h, f'【{label}】', HEI, 9.5, bold=True, color=NAVY)
        c = t.cell(1, col)
        for k, (topic, detail) in enumerate(items):
            r = cell_para(c, first=(k == 0), indent=1.35, hang=1.35)
            run_ea(r, f'{topic}　', HEI, 8.5, bold=True)
            run_ea(r, detail, MING, 8.5)

    q = para(doc, before=5, indent=0.6, hang=0.6, after=1, spacing=1.2)
    run_ea(q, '最容易失分的五處：', HEI, 9, bold=True, color=NAVY)
    run_ea(q, '① 分數除法與「剩下的幾分之幾」這類兩層敘述　'
              '② 折扣成數（八折是乘 0.8，加四成是乘 1.4）　'
              '③ 圓的公式吃半徑不吃直徑　'
              '④ 單位換算（公升與立方公分、公頃、時速裡的幾分鐘）　'
              '⑤ 基準量與比較量倒過來除。', MING, 9)

    p = para(doc, before=7, after=3)
    run_ea(p, '三、教材怎麼備', HEI, 11, bold=True, color=NAVY)
    for label, detail in MATERIALS:
        q = para(doc, indent=0.6, hang=0.6, after=2, spacing=1.2)
        run_ea(q, f'{label}　', HEI, 9, bold=True)
        run_ea(q, detail, MING, 9)

    p = para(doc, before=7, after=3)
    run_ea(p, '四、時程建議', HEI, 11, bold=True, color=NAVY)
    for when, what in TIMELINE:
        q = para(doc, indent=2.6, hang=2.6, after=1, spacing=1.2)
        run_ea(q, f'{when}　', HEI, 9, bold=True)
        run_ea(q, what, MING, 9)

    out = DRIVE / '私中數學備考說明.docx'
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def main():
    made = []
    for spec in (PAPER1, PAPER2):
        made.append(build_paper(spec, with_answers=False))
        made.append(build_paper(spec, with_answers=True))
    made.append(build_guide())
    for f in made:
        print('OK', f)


if __name__ == '__main__':
    main()
