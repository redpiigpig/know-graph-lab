# -*- coding: utf-8 -*-
"""每次上課的課堂講義（HTML 章節 → 可列印 Word）。

從 public/content/works/{slug}/chapters-*/chNN.html 取出某一次上課要用的章節，
加上課程封面（課名／第幾次／日期／單元／教師／班級／教室）後輸出成 docx。
與 works_series_docx.py 分工：那支做「整本書」，本支做「一次上課的份量」，
並額外支援 <table>（本書講義大量使用對照表）。

成品是 docx，依 docs/repo-hygiene.md 不進 git，輸出到 Drive：
  G:\\我的雲端硬碟\\資料\\知識圖工作室\\教學\\{課程資料夾}\\講義\\

用法：
  python scripts/course_handout_docx.py            # 第 1 次
  python scripts/course_handout_docx.py 1 2 3      # 指定幾次
  python scripts/course_handout_docx.py all        # 全部八次
"""
import re
import sys
from pathlib import Path

from lxml import html as lhtml
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
PUB = ROOT / 'public'
DRIVE = Path(r'G:\我的雲端硬碟\資料\知識圖工作室\教學')

KAI = 'DFKai-SB'        # 標楷體：標題與引文
MING = 'PMingLiU'       # 明體：正文
HEI = 'Microsoft JhengHei'
GRAY = RGBColor(0x60, 0x60, 0x60)

COURSE_WR2 = {
    'slug': 'world-religions-intro',
    'chapters_dir': 'chapters-wr2',
    'folder': '115-1_世界宗教文化導論',
    'title': '世界宗教文化導論',
    'code': 'PPA001',
    'teacher': '張辰瑋',
    'klass': '玄奘大學宗教與文化學系‧二年制在職專班1年A班（雙週）',
    'place': '妙然401　週日 13:00–17:00',
    'sessions': [
        (1, '2026年9月20日', [1, 2], '宗教是什麼、宗教裡有什麼'),
        (2, '2026年10月4日', [3, 4], '信仰、體制與範疇；宗教怎麼分類'),
        (3, '2026年10月18日', [5, 6], '泛神論'),
        (4, '2026年11月1日', [7, 8], '多神論'),
        (5, '2026年11月15日', [9, 10], '一神論'),
        (6, '2026年11月29日', [11, 12], '實用神論'),
        (7, '2026年12月13日', [13, 14], '現代世界的宗教'),
        (8, '2026年12月27日', [15, 16], '臺灣宗教（含期末考）'),
    ],
}

# PPA005 基督宗教概論（雙週班）。日期依開課學期填；book 模式不使用日期。
COURSE_CH = {
    'slug': 'christianity-intro',
    'chapters_dir': 'chapters',
    'folder': '115-1_基督宗教概論',
    'title': '基督宗教概論',
    'code': 'PPA005',
    'teacher': '張辰瑋',
    'klass': '玄奘大學宗教與文化學系‧二年制在職專班1年A班（雙週）',
    'place': '待定',
    'sessions': [
        (1, '待定', [1, 2], '四項核心傳統；拿撒勒人耶穌'),
        (2, '待定', [3, 4], '初代教會；正典的形成'),
        (3, '待定', [5, 6], '釋經傳統；大公會議'),
        (4, '待定', [7, 8], '教父與東方教會；中世紀西方教會'),
        (5, '待定', [9, 10], '宗教改革；近代基督教'),
        (6, '待定', [11, 12], '二十世紀與當代；禮儀與聖事'),
        (7, '待定', [13, 14], '聖職與體制；藝術與物質文化'),
        (8, '待定', [15, 16], '臺灣的基督宗教；生死觀（含期末考）'),
    ],
}

# 大學國文（十六週）；book 模式不使用日期與場地。
COURSE_SL = {
    'slug': 'sinographic-literature',
    'chapters_dir': 'chapters',
    'folder': '宗教系國文講義',
    'title': '宗教系國文講義',
    'code': '大學國文',
    'teacher': '張辰瑋',
    'klass': '玄奘大學宗教與文化學系',
    'place': '待定',
    'sessions': [(i, '待定', [2 * i - 1, 2 * i], '') for i in range(1, 9)],
}

COURSES = {'wr2': COURSE_WR2, 'ch': COURSE_CH, 'sl': COURSE_SL}
COURSE = COURSE_WR2


# ── Word 基本工具 ───────────────────────────────────────────────────────────
def run_ea(par, text, font=MING, size=11.0, bold=False, italic=False,
           sup=False, color=None):
    r = par.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if sup:
        r.font.superscript = True
    if color:
        r.font.color.rgb = color
    return r


def emit_inline(par, el, font, size, bold=False, italic=False, color=None):
    """遞迴輸出行內內容：strong→粗、em→斜、sup→上標、a→純文字。"""
    if el.text:
        run_ea(par, el.text, font, size, bold, italic, color=color)
    for ch in el:
        tag = ch.tag if isinstance(ch.tag, str) else ''
        b = bold or tag in ('strong', 'b')
        i = italic or tag in ('em', 'i')
        if tag == 'sup':
            run_ea(par, ch.text_content(), font, max(size - 2.5, 7), sup=True)
        elif tag == 'br':
            run_ea(par, '　', font, size)
        else:
            emit_inline(par, ch, font, size, b, i, color)
        if ch.tail:
            run_ea(par, ch.tail, font, size, bold, italic, color=color)


def para(doc_or_cell, el=None, text=None, font=MING, size=11.0, bold=False,
         indent=None, first_indent=Cm(0.74), spacing=1.5, before=Pt(0),
         after=Pt(4), align=None, color=None):
    p = doc_or_cell.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = spacing
    pf.space_before = before
    pf.space_after = after
    if indent is not None:
        pf.left_indent = indent
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if align is not None:
        p.alignment = align
    if el is not None:
        emit_inline(p, el, font, size, bold, color=color)
    elif text:
        run_ea(p, text, font, size, bold, color=color)
    return p


def heading(doc, text, level, center=False):
    # 套 Word 內建 Heading 樣式（目錄欄位靠它抓章節），字型另外指定
    sizes = {1: 17, 2: 13.5, 3: 11.5}
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.space_before = Pt(4 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 6)
    p.paragraph_format.line_spacing = 1.3
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ea(p, text, KAI, sizes[level], bold=True)
    return p


def shade(cell, hexcolor):
    e = OxmlElement('w:shd')
    e.set(qn('w:val'), 'clear')
    e.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(e)


def emit_table(doc, el):
    """HTML <table> → Word 表格（支援 rowspan；colspan 少用故不處理）。"""
    rows = el.cssselect('tr')
    if not rows:
        return
    ncols = 0
    for tr in rows:
        ncols = max(ncols, sum(int(c.get('colspan', 1)) for c in tr.cssselect('th,td')))
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # rowspan：記錄每欄還被上一列佔用幾列
    pending = [0] * ncols
    for ri, tr in enumerate(rows):
        cells = tr.cssselect('th,td')
        ci = 0
        for c in cells:
            while ci < ncols and pending[ci] > 0:
                pending[ci] -= 1
                ci += 1
            if ci >= ncols:
                break
            cell = t.rows[ri].cells[ci]
            cell.text = ''
            is_head = (c.tag == 'th')
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.25
            if is_head:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            emit_inline(p, c, HEI if is_head else MING, 9.5, bold=is_head)
            if is_head:
                shade(cell, 'E8EDF3')
            rs = int(c.get('rowspan', 1))
            if rs > 1:
                pending[ci] = rs - 1
                for k in range(1, rs):
                    if ri + k < len(rows):
                        t.rows[ri].cells[ci].merge(t.rows[ri + k].cells[ci])
            ci += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def emit_list(doc, el, ordered):
    marks = '一二三四五六七八九十十一十二十三十四十五'
    for i, li in enumerate(el.cssselect('li')):
        prefix = f'{marks[i] if ordered and i < 10 else str(i + 1)}、' if ordered else '‧ '
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(0.9)
        pf.first_line_indent = Cm(-0.45)
        pf.line_spacing = 1.5
        pf.space_after = Pt(3)
        run_ea(p, prefix, MING, 10.5)
        emit_inline(p, li, MING, 10.5)


def emit_footnotes(doc, el):
    for item in el.cssselect('.fn-item'):
        num = (item.cssselect('.fn-num') or [None])[0]
        body = (item.cssselect('.fn-body') or [None])[0]
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Cm(0.9)
        pf.first_line_indent = Cm(-0.9)
        pf.line_spacing = 1.35
        pf.space_after = Pt(2)
        run_ea(p, f'[{num.text_content().strip() if num is not None else ""}] ',
               MING, 9, bold=True)
        if body is not None:
            for a in body.cssselect('a.footnote-backref'):
                a.getparent().remove(a)
            emit_inline(p, body, MING, 9, color=GRAY)


def emit_chapter(doc, sec):
    for ch in sec:
        tag = ch.tag if isinstance(ch.tag, str) else ''
        cls = ch.get('class') or ''
        if tag == 'h2':
            heading(doc, ch.text_content().strip(), 1, center=True)
        elif tag == 'h3':
            heading(doc, ch.text_content().strip(), 2)
        elif tag == 'h4':
            heading(doc, ch.text_content().strip(), 3)
        elif tag == 'p':
            para(doc, el=ch, size=11)
        elif tag == 'blockquote':
            for sub in ch.iter('p'):
                para(doc, el=sub, font=KAI, size=11, indent=Cm(1.0),
                     first_indent=None, spacing=1.5, before=Pt(4), after=Pt(6))
        elif tag == 'ul':
            emit_list(doc, ch, ordered=False)
        elif tag == 'ol':
            emit_list(doc, ch, ordered=True)
        elif tag == 'table':
            emit_table(doc, ch)
        elif tag == 'div' and 'footnotes' in cls:
            emit_footnotes(doc, ch)


def add_toc(doc):
    """插入 Word 目錄欄位（開檔後按 F9 或轉 PDF 時會回填頁碼）。"""
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-2" \h \z \u')
    inner = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '（目錄：在 Word 中按 Ctrl+A 再按 F9 更新頁碼）'
    r.append(t); inner.append(r); fld.append(inner)
    p._p.append(fld)


def add_page_numbers(doc, label):
    par = doc.sections[0].footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ea(par, f'{label}　－ ', MING, 9, color=GRAY)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    par._p.append(fld)
    run_ea(par, ' －', MING, 9, color=GRAY)


# ── 產生一次上課的講義 ──────────────────────────────────────────────────────
def build(session):
    no, date, chapters, unit = session
    cdir = PUB / 'content/works' / COURSE['slug'] / COURSE['chapters_dir']

    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Cm(2.2)
    s.left_margin = s.right_margin = Cm(2.4)
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'
    st.element.rPr.rFonts.set(qn('w:eastAsia'), MING)
    st.font.size = Pt(11)

    # ── 封面 ──
    def cover(txt, font, size, bold, before, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(0)
        run_ea(p, txt, font, size, bold, color=color)

    cover(COURSE['klass'], HEI, 11, False, 120, GRAY)
    cover(f"{COURSE['title']}（{COURSE['code']}）", KAI, 26, True, 22)
    cover(f'第 {no} 次上課講義　　{date}', KAI, 14, False, 26)
    cover(f'第{chapters[0]}–{chapters[-1]}章　{unit}', KAI, 13, False, 10, GRAY)
    cover(COURSE['place'], HEI, 10.5, False, 40, GRAY)
    cover(f"授課教師　{COURSE['teacher']}", KAI, 12, False, 14)
    cover('線上講義全書：redpiigpig.com/works', HEI, 9.5, False, 60, GRAY)

    # ── 本次進度 ──
    doc.add_page_break()
    heading(doc, '本次進度', 2)
    titles = []
    for n in chapters:
        f = cdir / f'ch{n:02d}.html'
        tree = lhtml.fromstring(f.read_text(encoding='utf-8'))
        h2 = tree.cssselect('h2')
        titles.append(h2[0].text_content().strip() if h2 else f'第{n}章')
    for t in titles:
        para(doc, text=t, font=KAI, size=12, first_indent=None,
             indent=Cm(0.8), after=Pt(6))
    para(doc, text='每章末附「本章重點」、「討論問題」、註釋與參考資料，'
                   '課後請先看過討論問題再讀正文。',
         size=10, first_indent=None, indent=Cm(0.8), before=Pt(14), color=GRAY)

    # ── 章節本文 ──
    for n in chapters:
        f = cdir / f'ch{n:02d}.html'
        tree = lhtml.fromstring(f.read_text(encoding='utf-8'))
        sec = tree if tree.tag == 'section' else (tree.cssselect('section') or [tree])[0]
        doc.add_page_break()
        emit_chapter(doc, sec)

    add_page_numbers(doc, f"{COURSE['title']}　第 {no} 次")

    outdir = DRIVE / COURSE['folder'] / '講義'
    outdir.mkdir(parents=True, exist_ok=True)
    y, m, d = re.match(r'(\d+)年(\d+)月(\d+)日', date).groups()
    ymd = f'{y}{int(m):02d}{int(d):02d}'
    out = outdir / f'第{no}次_{ymd}_第{chapters[0]}-{chapters[-1]}章.docx'
    doc.save(out)
    return out


def build_book():
    """紙本版講義全書：封面＋目錄＋十六章。"""
    cdir = PUB / 'content/works' / COURSE['slug'] / COURSE['chapters_dir']

    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Cm(2.4)
    s.left_margin = s.right_margin = Cm(2.6)
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'
    st.element.rPr.rFonts.set(qn('w:eastAsia'), MING)
    st.font.size = Pt(11)
    for name, size in (('Heading 1', 17), ('Heading 2', 13.5), ('Heading 3', 11.5)):
        stl = doc.styles[name]
        stl.font.name = 'Times New Roman'
        stl.element.rPr.rFonts.set(qn('w:eastAsia'), KAI)
        stl.font.size = Pt(size)
        stl.font.bold = True
        stl.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

    def cover(txt, font, size, bold, before, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(0)
        run_ea(p, txt, font, size, bold, color=color)

    cover(COURSE['klass'], HEI, 11, False, 150, GRAY)
    cover(COURSE['title'], KAI, 32, True, 26)
    cover('神聖的經驗與探詢——一部以信仰型態為軸的世界宗教入門', KAI, 13, False, 18, GRAY)
    cover(f"授課教師　{COURSE['teacher']}", KAI, 13, False, 70)
    cover('全書十六章　配合十六週、雙週上課八次的學期使用', HEI, 10, False, 40, GRAY)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run_ea(p, '目次', KAI, 17, bold=True, color=RGBColor(0x1F, 0x39, 0x64))
    add_toc(doc)

    for n in range(1, 17):
        f = cdir / f'ch{n:02d}.html'
        tree = lhtml.fromstring(f.read_text(encoding='utf-8'))
        sec = tree if tree.tag == 'section' else (tree.cssselect('section') or [tree])[0]
        doc.add_page_break()
        emit_chapter(doc, sec)

    add_page_numbers(doc, COURSE['title'])
    outdir = DRIVE / COURSE['folder'] / '講義'
    outdir.mkdir(parents=True, exist_ok=True)
    out = outdir / f"{COURSE['title']}（紙本版‧全十六章）.docx"
    doc.save(out)
    return out


if __name__ == '__main__':
    args = sys.argv[1:] or ['1']
    course = next((a.split('=')[1] for a in args if a.startswith('--course=')), 'wr2')
    COURSE = COURSES[course]
    args = [a for a in args if not a.startswith('--')] or ['1']
    if args == ['book']:
        print('✔', build_book())
        raise SystemExit
    nos = [s[0] for s in COURSE['sessions']] if args == ['all'] else [int(a) for a in args]
    for no in nos:
        sess = next(s for s in COURSE['sessions'] if s[0] == no)
        print('✔', build(sess))
