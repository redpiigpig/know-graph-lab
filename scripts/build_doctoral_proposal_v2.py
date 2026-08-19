from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\user\Desktop\know-graph-lab")
SOURCE = ROOT / "data" / "doctoral_thesis_proposal_v2_draft.md"
OUTPUT = ROOT / "output" / "documents" / "張辰瑋_博士論文研究計畫_第二版草案.docx"

TITLE_ZH = "入世轉向的兩條譜系"
SUBTITLE_ZH = "台灣人間佛教與以長老教會為核心之本土神學的歷史比較及神學比較"
FIGURES_ZH = "以太虛、印順、傳道、昭慧與黃彰輝、宋泉盛、王憲治、黃伯和為核心"
TITLE_EN = "Two Genealogies of This-Worldly Engagement"
SUBTITLE_EN = "A Historical and Theological Comparison of Humanistic Buddhism and Presbyterian-Centered Taiwanese Contextual Theology"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "A6A6A6"
TEXT = "222222"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "B8C2CC")


def set_east_asia_font(run, name="Microsoft JhengHei") -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr, fld_char_2])


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "開啟文件後更新欄位即可顯示目錄"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, placeholder, end])


def add_inline_text(paragraph, text: str, *, size: float | None = None) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if bold else part
        run = paragraph.add_run(clean)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(TEXT)
        set_east_asia_font(run)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, (label, width) in enumerate(zip(headers, widths)):
        cell = hdr.cells[i]
        set_cell_width(cell, width)
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(9)
        set_east_asia_font(r)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        if ridx % 2 == 1:
            for c in cells:
                set_cell_shading(c, LIGHT_GRAY)
        for i, (value, width) in enumerate(zip(row, widths)):
            cell = cells[i]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(value)
            r.font.size = Pt(8.5)
            set_east_asia_font(r)
    doc.add_paragraph()


def figures_table(doc: Document) -> None:
    headers = ["傳統", "人物", "歷史位置", "主要場域／媒介", "比較焦點"]
    rows = [
        ["人間佛教", "太虛", "近代佛教改革與人生佛教的問題開端", "僧教育、佛教團體、刊物與改革方案", "佛教如何回應現代國家與社會危機"],
        ["人間佛教", "印順", "以佛教史與經典判攝重建人間佛教", "著作、講學、僧團及台灣佛教知識網絡", "歷史敘事如何賦予入世實踐正當性"],
        ["人間佛教", "傳道", "印順思想的制度化、出版化與地方實踐", "妙心寺、中華佛教百科文獻基金會、社會運動", "思想如何透過寺院與組織進入公共領域"],
        ["人間佛教", "昭慧", "人間佛教公共倫理的系統化與行動化", "弘誓學院、戒律詮釋、性別與動物倫理運動", "規範倫理如何連結佛教傳統與社會倡議"],
        ["本土神學", "黃彰輝", "實況化神學與普世神學教育的開創", "神學教育、普世教會網絡、處境化方法", "教會如何在具體歷史處境中重述信仰"],
        ["本土神學", "宋泉盛", "亞洲故事、文化經驗與受苦人民的神學重構", "敘事神學、亞洲神學、普世出版與教學", "人民故事如何成為神學知識與救贖論資源"],
        ["本土神學", "王憲治", "鄉土神學的政治化、制度化與公共實踐", "南神、長老教會神學顧問、鄉土神學研究組織", "人民、土地、權力與上帝如何構成歷史政治論述"],
        ["本土神學", "黃伯和", "出頭天／自決神學與信仰再告白的系統發展", "研究中心、出版、人才培育、跨宗教與國際網絡", "本土神學如何延續、擴張並回應新公共議題"],
    ]
    add_table(doc, headers, rows, [950, 900, 2380, 2700, 2430])


def schedule_table(doc: Document) -> None:
    headers = ["階段", "主要工作", "材料與方法", "預期成果"]
    rows = [
        ["第一年上", "修訂研究問題、概念界定與史料目錄；完成太虛、印順與黃彰輝、宋泉盛之初步文獻回顧", "專書、論文、年譜、期刊與檔案盤點；建立人物—概念—機構年表", "緒論、研究回顧與方法章草稿"],
        ["第一年下", "重建兩條譜系的早期歷史處境與思想形成", "思想史、概念史、文本脈絡分析；比對版本與出版時間", "第二、三、五章初稿"],
        ["第二年上", "研究傳道、昭慧、王憲治、黃伯和的制度網絡與公共實踐", "寺院、神學院、教會、基金會、期刊與社運資料；接受史與制度史", "第四、六章初稿；機構與事件資料庫"],
        ["第二年下", "進行訪談並完成跨宗教事件個案的史料校證", "半結構訪談、口述歷史、新聞與組織檔案互證", "訪談逐字稿、事件年表與個案分析"],
        ["第三年上", "展開歷時與共時的比較歷史分析，並進行有限度神學比較", "比較矩陣、概念流變、因果機制與不對稱性分析", "第七章初稿；全論文整合稿"],
        ["第三年下", "撰寫結論，評估「人間宗教」命題的解釋力與界限；全面修訂", "回到史料檢核論證；補訪、格式與引註校訂", "第八章、博士論文完整稿與口試版本"],
    ]
    add_table(doc, headers, rows, [1050, 3100, 3150, 2060])


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Title", 22, DARK_BLUE, 0, 10),
        ("Subtitle", 13, BLUE, 0, 8),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.208

    if "Bibliography" not in styles:
        bib = styles.add_style("Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bib = styles["Bibliography"]
    bib.font.name = "Calibri"
    bib._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    bib.font.size = Pt(10)
    bib.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bib.paragraph_format.left_indent = Inches(0.28)
    bib.paragraph_format.first_line_indent = Inches(-0.28)
    bib.paragraph_format.space_after = Pt(4)
    bib.paragraph_format.line_spacing = 1.15


def configure_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("博士論文研究計畫｜第二版草案")
    hr.font.size = Pt(8.5)
    hr.font.color.rgb = RGBColor.from_string(MID_GRAY)
    set_east_asia_font(hr)
    add_page_field(section.footer.paragraphs[0])


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("玄奘大學宗教與文化學系博士班")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_east_asia_font(r)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(TITLE_ZH)
    r.font.size = Pt(24)
    r.bold = True
    set_east_asia_font(r)

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(SUBTITLE_ZH)
    r.font.size = Pt(16)
    r.bold = True
    set_east_asia_font(r)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(FIGURES_ZH)
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_east_asia_font(r)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(TITLE_EN)
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(MID_GRAY)
    set_east_asia_font(r)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run(SUBTITLE_EN)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(MID_GRAY)
    set_east_asia_font(r)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("博士論文研究計畫（第二版草案）")
    r.bold = True
    r.font.size = Pt(13)
    set_east_asia_font(r)

    for _ in range(3):
        doc.add_paragraph()

    for label in ("研究生：張辰瑋", "日期：2026 年 7 月 24 日", "狀態：未發表草案，供指導與修訂使用"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(label)
        r.font.size = Pt(11)
        if label.startswith("狀態"):
            r.font.color.rgb = RGBColor.from_string(MID_GRAY)
        set_east_asia_font(r)
    doc.add_page_break()


def add_contents_page(doc: Document) -> None:
    p = doc.add_paragraph("目錄", style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = False
    p.paragraph_format.space_after = Pt(16)
    entries = [
        "摘要／Abstract",
        "一、研究背景、動機與問題意識",
        "二、研究回顧",
        "三、研究對象、範圍與分期",
        "四、研究方法與材料",
        "五、主要史料與分析程序",
        "六、預期研究成果與貢獻",
        "七、論文章節大綱",
        "八、研究進度規劃",
        "九、核心參考文獻",
    ]
    for entry in entries:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(entry)
        r.font.size = Pt(11)
        if entry[0].isdigit() or entry[0] in "一二三四五六七八九":
            r.bold = True
        set_east_asia_font(r)
    doc.add_page_break()


def add_markdown(doc: Document, source: str) -> None:
    previous_blank = True
    first_h1 = True
    for raw in source.splitlines():
        line = raw.rstrip()
        if not line.strip():
            previous_blank = True
            continue
        if line == "[[TABLE:FIGURES]]":
            figures_table(doc)
            previous_blank = False
            continue
        if line == "[[TABLE:SCHEDULE]]":
            schedule_table(doc)
            previous_blank = False
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:].strip(), style="Heading 3")
            p.paragraph_format.keep_with_next = True
        elif line.startswith("## "):
            p = doc.add_paragraph(line[3:].strip(), style="Heading 2")
            p.paragraph_format.keep_with_next = True
        elif line.startswith("# "):
            p = doc.add_paragraph(line[2:].strip(), style="Heading 1")
            if not first_h1 and line not in ("# Abstract",):
                p.paragraph_format.page_break_before = True
            first_h1 = False
        elif line.startswith("REF: "):
            p = doc.add_paragraph(style="Bibliography")
            add_inline_text(p, line[5:].strip(), size=10)
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_text(p, re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_text(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            add_inline_text(p, line)
            if previous_blank and len(line) < 50 and line.endswith("："):
                p.paragraph_format.keep_with_next = True
        previous_blank = False


def set_document_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = f"{TITLE_ZH}：{SUBTITLE_ZH}"
    props.subject = "博士論文研究計畫第二版草案"
    props.author = "張辰瑋"
    props.keywords = "人間佛教；台灣本土神學；比較宗教史；思想史；神學比較"
    props.comments = "第二版未發表草案；歷史學方法為主，神學比較為輔。"


def prevent_table_row_splitting(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    configure_header_footer(doc)
    set_document_properties(doc)
    add_cover(doc)
    add_contents_page(doc)
    add_markdown(doc, text)
    prevent_table_row_splitting(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
