"""口述訪談紀錄 → Word (.docx) 匯出工具

讀取 public/content/interviews/*.txt（與 /thesis?tab=interviews 線上頁面同源），
套用與 Vue 元件 pages/thesis/interview/[name].vue 一致的樣式規則，
輸出每篇一個 .docx，預設輸出到 c:/tmp/interviews-docx/。

用法：
  # 全部轉檔
  python scripts/interviews_to_word.py

  # 單篇（檔名片段比對，模糊匹配）
  python scripts/interviews_to_word.py 釋昭慧
  python scripts/interviews_to_word.py "04.06"

  # 自訂輸出目錄
  python scripts/interviews_to_word.py --out D:/exports

需要：pip install python-docx
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parent.parent
SRC_DIR = ROOT / "public" / "content" / "interviews"
DEFAULT_OUT = Path("c:/tmp/interviews-docx")

CN_FONT = "Noto Serif TC"
EN_FONT = "Georgia"

# ── 樣式參數（對照 [name].vue 的 :deep(.i-*) CSS）────────────────────
TITLE_PT = 22
META_PT = 11
SECTION_PT = 17
SUBSECTION_PT = 14
PARA_PT = 12
Q_LABEL_PT = 10
A_LABEL_PT = 10

COLOR_TITLE = RGBColor(0x11, 0x18, 0x27)
COLOR_META = RGBColor(0x6B, 0x72, 0x80)
COLOR_SECTION = RGBColor(0x37, 0x41, 0x51)
COLOR_SUBSECTION = RGBColor(0x4B, 0x55, 0x63)
COLOR_PARA = RGBColor(0x1F, 0x29, 0x37)
COLOR_Q_LABEL = RGBColor(0x6B, 0x72, 0x80)
COLOR_Q_BODY = RGBColor(0x37, 0x41, 0x51)
COLOR_A_LABEL = RGBColor(0x7C, 0x3A, 0xED)  # 紫色
COLOR_A_BODY = RGBColor(0x11, 0x18, 0x27)

Q_SHADE = "F9FAFB"      # 問句淡灰底
Q_BORDER = "D1D5DB"     # 問句左側 border
SECTION_RULE = "E5E7EB"  # 章節下底線

FOOTNOTE_PATTERN = re.compile(r"\[(\d+)\]\(#footnote\d+\)")
ANSWER_PATTERN = re.compile(r"^([一-鿿]{1,5}(?:法師|教授|主教|和尚|居士|博士|老師|牧師|女士|先生|主任|院長|住持))：(.+)$")
META_PATTERN = re.compile(r"^(受訪者|訪問者|訪問時間|訪問地點|訪談時間|訪談地點|採訪者|採訪時間|採訪地點)：")
SECTION_PATTERN = re.compile(r"^[一二三四五六七八九十]+、")
SUBSECTION_PATTERN = re.compile(r"^（[一二三四五六七八九十]+）")
NUMBERED_SUB_PATTERN = re.compile(r"^\d+\.\s+\S")  # 例: "1.  印順導師的人格典範"


# ── docx 低階輔助 ───────────────────────────────────────────
def _set_cell_or_para_shading(element, fill_hex: str) -> None:
    """套用底色到段落 (<w:pPr><w:shd .../>)。element 是 <w:pPr>。"""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    element.append(shd)


def _add_left_border(pPr, color_hex: str, size: int = 18) -> None:
    """段落左側邊框 (對應 CSS border-left)。"""
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))      # 1/8 pt 單位
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_bottom_border(pPr, color_hex: str, size: int = 8) -> None:
    """段落下底線 (對應 CSS border-bottom)。"""
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _ensure_pPr(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    return pPr


def _set_run_fonts(run, size_pt: int, color: RGBColor, *,
                   bold: bool = False, italic: bool = False,
                   superscript: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.font.name = EN_FONT
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rFonts.set(qn("w:cs"), EN_FONT)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if superscript:
        run.font.superscript = True


def _add_runs_with_footnotes(paragraph, text: str, *,
                             size_pt: int, color: RGBColor,
                             italic: bool = False, bold: bool = False) -> None:
    """把 `...[39](#footnote39)...` 拆成正文 + 上標 [39] runs。"""
    cursor = 0
    for m in FOOTNOTE_PATTERN.finditer(text):
        if m.start() > cursor:
            run = paragraph.add_run(text[cursor:m.start()])
            _set_run_fonts(run, size_pt, color, italic=italic, bold=bold)
        fn_run = paragraph.add_run(f"[{m.group(1)}]")
        _set_run_fonts(fn_run, max(size_pt - 2, 8), color, superscript=True)
        cursor = m.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        _set_run_fonts(run, size_pt, color, italic=italic, bold=bold)


# ── 段落樣式 ────────────────────────────────────────────────
def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    _set_run_fonts(run, TITLE_PT, COLOR_TITLE, bold=True)


def add_meta(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_run_fonts(run, META_PT, COLOR_META)


def add_section(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(10)
    pPr = _ensure_pPr(p)
    _add_bottom_border(pPr, SECTION_RULE, size=12)
    run = p.add_run(text)
    _set_run_fonts(run, SECTION_PT, COLOR_SECTION, bold=True)


def add_subsection(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_fonts(run, SUBSECTION_PT, COLOR_SUBSECTION, bold=True)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(PARA_PT * 2)
    p.paragraph_format.line_spacing = 1.9
    p.paragraph_format.space_after = Pt(4)
    _add_runs_with_footnotes(p, text, size_pt=PARA_PT, color=COLOR_PARA)


def add_question(doc: Document, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.8
    p.paragraph_format.left_indent = Cm(0.3)
    pPr = _ensure_pPr(p)
    _set_cell_or_para_shading(pPr, Q_SHADE)
    _add_left_border(pPr, Q_BORDER, size=24)

    label = p.add_run("問　")
    _set_run_fonts(label, Q_LABEL_PT, COLOR_Q_LABEL, bold=True)
    _add_runs_with_footnotes(p, body, size_pt=PARA_PT, color=COLOR_Q_BODY, italic=True)


def add_answer(doc: Document, label_text: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.9
    p.paragraph_format.first_line_indent = Pt(0)

    label = p.add_run(f"{label_text}　")
    _set_run_fonts(label, A_LABEL_PT, COLOR_A_LABEL, bold=True)
    _add_runs_with_footnotes(p, body, size_pt=PARA_PT, color=COLOR_A_BODY)


def add_blank(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


# ── 內文解析（對照 [name].vue formatInterview）────────────────────
def render(doc: Document, raw: str) -> None:
    lines = raw.splitlines()
    is_first = True
    for line in lines:
        t = line.strip()
        if not t:
            add_blank(doc)
            continue

        if is_first:
            is_first = False
            add_title(doc, t)
            continue

        if META_PATTERN.match(t):
            add_meta(doc, t)
            continue

        if SECTION_PATTERN.match(t):
            add_section(doc, t)
            continue

        if SUBSECTION_PATTERN.match(t) or NUMBERED_SUB_PATTERN.match(t):
            add_subsection(doc, t)
            continue

        if t.startswith("筆者："):
            add_question(doc, t[len("筆者："):])
            continue

        m = ANSWER_PATTERN.match(t)
        if m:
            add_answer(doc, m.group(1), m.group(2))
            continue

        # 短行（≤25 字、無結尾標點）→ 視為小標題
        if len(t) <= 25 and not re.search(r"[。，、；：？！…]$", t) and not t.startswith(("[", "【", "（", "〔")):
            add_section(doc, t)
            continue

        add_paragraph(doc, t)


def configure_document(doc: Document) -> None:
    # 邊距與預設字型
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(PARA_PT)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rFonts.set(qn("w:cs"), EN_FONT)


def convert_one(src: Path, out_dir: Path) -> Path:
    raw = src.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    render(doc, raw)
    out_path = out_dir / f"{src.stem}.docx"
    doc.save(out_path)
    return out_path


def main() -> int:
    ap = argparse.ArgumentParser(description="把 /thesis 口述訪談紀錄轉成 Word")
    ap.add_argument("filter", nargs="?", default="",
                    help="檔名片段（模糊匹配），不給就轉全部")
    ap.add_argument("--out", default=str(DEFAULT_OUT),
                    help=f"輸出目錄（預設 {DEFAULT_OUT}）")
    ap.add_argument("--src", default=str(SRC_DIR),
                    help=f"來源目錄（預設 {SRC_DIR}）")
    args = ap.parse_args()

    src_dir = Path(args.src)
    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    if not src_dir.exists():
        print(f"[ERROR] 找不到來源目錄：{src_dir}", file=sys.stderr)
        return 1

    files = sorted(src_dir.glob("*.txt"))
    if args.filter:
        needle = args.filter.lower()
        files = [f for f in files if needle in f.name.lower()]

    if not files:
        print("[WARN] 沒有符合的檔案。", file=sys.stderr)
        return 1

    print(f"[INFO] 將轉換 {len(files)} 篇 → {out_dir}")
    for src in files:
        try:
            out = convert_one(src, out_dir)
            print(f"  [OK]   {src.name}  ->  {out.name}")
        except Exception as e:  # noqa: BLE001
            print(f"  [FAIL] {src.name}: {e}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    sys.exit(main())
