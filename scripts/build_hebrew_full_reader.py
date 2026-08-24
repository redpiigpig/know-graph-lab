"""Build the complete JIS-B5 Hebrew original-language reader (DOCX + PDF-ready).

Design preset: ``compact_reference_guide``.
Opening pattern: ``editorial_cover``.
The preset is deliberately overridden to JIS B5 (182 x 257 mm) because the
reader is intended as a bound study volume rather than US Letter handouts.
"""

from __future__ import annotations

import json
import math
import os
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "output" / "source-cache" / "original-readers" / "hebrew-full" / "hebrew-reader-50-lessons.json"
# QA renders may point at a scratch gloss master; production always uses the
# checked-in one.
INTERLINEAR_PATH = Path(
    os.environ.get("HBO_INTERLINEAR")
    or ROOT / "output" / "source-cache" / "original-readers" / "hebrew-full" / "interlinear.json"
)
OUTPUT_DIR = ROOT / "output" / "original-readers"
OUTPUT_PATH = OUTPUT_DIR / "hebrew-original-reader-50-lessons.docx"

PAGE_WIDTH_MM = 182
PAGE_HEIGHT_MM = 257
MARGIN_INSIDE_MM = 24
MARGIN_OUTSIDE_MM = 17
MARGIN_TOP_MM = 18
MARGIN_BOTTOM_MM = 20
USABLE_WIDTH_MM = PAGE_WIDTH_MM - MARGIN_INSIDE_MM - MARGIN_OUTSIDE_MM

# compact_reference_guide numeric tokens, with named JIS-B5 and Hebrew-script
# overrides documented above. The page geometry is the sole form-factor
# override; the preset's paragraph rhythm, heading ladder, list rhythm, table
# indent, and base cell margins remain exact.
BODY_SIZE_PT = 11.5
BODY_LINE_MULTIPLE = 1.25
TITLE_SIZE_PT = 24
H1_SIZE_PT = 17
H2_SIZE_PT = 14
H3_SIZE_PT = 12.5
TABLE_SIZE_PT = 9.6
HEBREW_BODY_PT = 15
HEBREW_MEMORY_PT = 16
PARA_AFTER_PT = 6

# Interlinear layer.  Every running-text word carries a Traditional-Chinese
# gloss directly beneath it, and each unit closes with the whole-sentence
# meaning.  The Hebrew face is set slightly below the running-text size because
# a gloss row sits under it; the pair still reads larger than the previous
# translation-only rhythm.
INTERLINEAR_HEBREW_PT = 14
INTERLINEAR_GLOSS_PT = 9.6
SENSE_PT = 10.4
TRANSLATION_PT = 10.2
CAPTION_PT = 9.0
LABEL_PT = 8.2
INTERLINEAR_GUTTER_MM = 3.4
INTERLINEAR_LINE_GAP_PT = 3.5
CELL_PAD_DXA = 80
CELL_SIDE_PAD_DXA = 120
TABLE_INDENT_DXA = 120

FONT_ZH = "MingLiU"
# LibreOffice does not reliably resolve the variable Noto Sans TC build that is
# installed on this workstation.  Using it for small UI labels caused embedded
# MS Gothic/Tahoma substitutions in the print PDF.  MingLiU is installed as a
# conventional TrueType collection, covers both Traditional Chinese and Latin,
# and is therefore also the stable UI face for this private print edition.
FONT_UI = FONT_ZH
FONT_HEBREW = "Noto Serif Hebrew"
FONT_TRANSLIT = "Noto Serif"
INK = "29241F"
MUTED = "6E675F"
PAPER = "FFFDF8"
ACCENT = "8A4E2F"
ACCENT_DARK = "3A2720"
GOLD = "D4A653"
PALE = "F3EDE3"
PALE_2 = "FAF6EF"
RULE = "D5C9BA"


def set_rfonts(r_pr, font: str) -> None:
    """Pin every Word script slot and remove theme fallbacks."""
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), font)
    for key in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        fonts.attrib.pop(qn(f"w:{key}"), None)


def set_run_font(run, font: str, size: float, *, bold=False, italic=False, color=INK) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    set_rfonts(r_pr, font)


def add_mixed_script_text(paragraph, text: str, font: str, size: float, *, bold=False, italic=False, color=INK) -> None:
    """Add a mainly Chinese/Latin string while giving embedded Hebrew its face."""
    for chunk in filter(None, re.split(r"([\u0590-\u05FF]+)", text)):
        is_hebrew = bool(re.search(r"[\u05D0-\u05EA]", chunk))
        run = paragraph.add_run(chunk)
        set_run_font(run, FONT_HEBREW if is_hebrew else font, size, bold=bold, italic=italic, color=color)
        if is_hebrew:
            r_pr = run._r.get_or_add_rPr()
            rtl = r_pr.find(qn("w:rtl"))
            if rtl is None:
                rtl = OxmlElement("w:rtl")
                r_pr.append(rtl)
            rtl.set(qn("w:val"), "1")


def set_rtl(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    for run in paragraph.runs:
        r_pr = run._r.get_or_add_rPr()
        rtl = r_pr.find(qn("w:rtl"))
        if rtl is None:
            rtl = OxmlElement("w:rtl")
            r_pr.append(rtl)
        rtl.set(qn("w:val"), "1")


def set_keep(paragraph, *, next_paragraph=False, together=False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if next_paragraph:
        p_pr.append(OxmlElement("w:keepNext"))
    if together:
        p_pr.append(OxmlElement("w:keepLines"))


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_row_split(row) -> None:
    """Keep each logical record together; never strand a partial cell line."""
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:cantSplit"))
    if node is None:
        node = OxmlElement("w:cantSplit")
        tr_pr.append(node)
    node.set(qn("w:val"), "true")


def set_cell_margins(cell, top=CELL_PAD_DXA, start=CELL_SIDE_PAD_DXA, bottom=CELL_PAD_DXA, end=CELL_SIDE_PAD_DXA) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), color)


def set_table_geometry(table, widths_mm: list[float], *, indent_dxa=TABLE_INDENT_DXA) -> None:
    widths_dxa = [round(width / 25.4 * 1440) for width in widths_mm]
    target_dxa = round(sum(widths_mm) / 25.4 * 1440)
    widths_dxa[-1] += target_dxa - sum(widths_dxa)

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(target_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_dxa in widths_dxa:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width_dxa))
        grid.append(node)
    for index, width_dxa in enumerate(widths_dxa):
        for cell in table.columns[index].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")


def set_borders(table, *, color=RULE, size="3", outside=True, inside=True) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        enabled = outside if edge in ("top", "left", "bottom", "right") else inside
        node.set(qn("w:val"), "single" if enabled else "nil")
        if enabled:
            node.set(qn("w:sz"), size)
            node.set(qn("w:color"), color)


def paragraph_rule(paragraph, color=RULE, size="12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run_font(run, FONT_UI, 7.5, color=MUTED)


def configure(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = Mm(MARGIN_TOP_MM)
    section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    section.left_margin = Mm(MARGIN_INSIDE_MM)
    section.right_margin = Mm(MARGIN_OUTSIDE_MM)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(9)
    section.different_first_page_header_footer = True

    settings = document.settings._element
    if settings.find(qn("w:mirrorMargins")) is None:
        settings.append(OxmlElement("w:mirrorMargins"))

    styles = document.styles
    # Resolve the compact-reference preset all the way down to OOXML rather
    # than leaving any built-in style on Word's Calibri/Cambria theme.  Direct
    # Hebrew and transliteration runs still override this base face below.
    for style in styles:
        r_pr = style._element.get_or_add_rPr()
        set_rfonts(r_pr, FONT_ZH)
    doc_defaults = styles.element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.element.insert(0, doc_defaults)
    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        doc_defaults.insert(0, r_pr_default)
    default_r_pr = r_pr_default.find(qn("w:rPr"))
    if default_r_pr is None:
        default_r_pr = OxmlElement("w:rPr")
        r_pr_default.append(default_r_pr)
    set_rfonts(default_r_pr, FONT_ZH)

    normal = styles["Normal"]
    normal.font.name = FONT_ZH
    normal.font.size = Pt(BODY_SIZE_PT)
    set_rfonts(normal._element.get_or_add_rPr(), FONT_ZH)
    normal.paragraph_format.space_after = Pt(PARA_AFTER_PT)
    normal.paragraph_format.line_spacing = BODY_LINE_MULTIPLE

    for name, size, color, before, after in (
        ("Title", TITLE_SIZE_PT, INK, 0, 8),
        ("Heading 1", H1_SIZE_PT, ACCENT, 18, 10),
        ("Heading 2", H2_SIZE_PT, ACCENT, 14, 7),
        ("Heading 3", H3_SIZE_PT, ACCENT_DARK, 10, 5),
    ):
        style = styles[name]
        style.font.name = FONT_ZH
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        set_rfonts(style._element.get_or_add_rPr(), FONT_ZH)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    # The default template supplies genuine numbering definitions for these
    # two list styles; their paragraph geometry is resolved to the selected
    # compact_reference_guide preset instead of inheriting Word defaults.
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT_ZH
        style.font.size = Pt(BODY_SIZE_PT)
        set_rfonts(style._element.get_or_add_rPr(), FONT_ZH)
        style.paragraph_format.left_indent = Mm(9.525)  # 0.375 in
        style.paragraph_format.first_line_indent = Mm(-4.7752)  # -0.188 in
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(header.add_run("聖經希伯來文原文讀本  ·  五十課"), FONT_UI, 7.5, color=MUTED)
    paragraph_rule(header, color=RULE, size="3")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("私人研讀版  ·  "), FONT_UI, 7.5, color=MUTED)
    set_page_field(footer)
    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(first_footer.add_run("PRIVATE STUDY EDITION  ·  2026"), FONT_UI, 7.5, color=MUTED)

    props = document.core_properties
    props.title = "聖經希伯來文原文讀本：五十課"
    props.subject = "1,000詞、100節背誦、25章、25篇禱文／文章與完整逾越節禮文"
    props.author = "Know Graph Lab"
    props.keywords = "Biblical Hebrew, niqqud, WLC, BBH2, Haggadah, JIS B5"


def page_break(document: Document) -> None:
    document.add_page_break()


def add_label(document: Document, text: str, *, page_break_before=False):
    p = document.add_paragraph()
    p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    set_run_font(run, FONT_UI, LABEL_PT, bold=True, color=ACCENT)
    set_keep(p, next_paragraph=True)
    return p


def add_body(document: Document, text: str, *, size=BODY_SIZE_PT, color=INK, bold=False, italic=False, align=None):
    p = document.add_paragraph()
    if align is not None:
        p.alignment = align
    set_run_font(p.add_run(text), FONT_ZH, size, color=color, bold=bold, italic=italic)
    return p


def add_hebrew(
    document: Document,
    text: str,
    *,
    size=HEBREW_BODY_PT,
    color=INK,
    bold=False,
    background=None,
    line_spacing=1.38,
    space_after=None,
    table_indent_dxa=TABLE_INDENT_DXA,
):
    if background:
        table = document.add_table(rows=1, cols=1)
        set_table_geometry(table, [USABLE_WIDTH_MM], indent_dxa=table_indent_dxa)
        set_borders(table, outside=False, inside=False)
        cell = table.cell(0, 0)
        set_cell_margins(cell, top=110, bottom=110, start=130, end=130)
        shade(cell, background)
        p = cell.paragraphs[0]
    else:
        p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = line_spacing
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, FONT_HEBREW, size, color=color, bold=bold)
    set_rtl(p)
    # Short verses and prayers should stay together.  Very long Haggadah
    # segments must be allowed to break naturally; forcing keepLines on a
    # paragraph longer than a B5 page made LibreOffice flow text into the
    # footer instead of honoring the bottom margin.
    if len(text) <= 420:
        set_keep(p, together=True)
    return p


# --------------------------------------------------------------------------- #
# interlinear layer
# --------------------------------------------------------------------------- #


_interlinear_master: dict[str, dict] | None = None
_hebrew_metrics_font = None


def interlinear_master() -> dict[str, dict]:
    """Word-by-word Chinese glosses, keyed by unit id (see build_hebrew_interlinear.py)."""
    global _interlinear_master
    if _interlinear_master is None:
        if not INTERLINEAR_PATH.exists():
            raise SystemExit(
                f"缺少逐詞對譯主檔 {INTERLINEAR_PATH}；先跑 scripts/build_hebrew_interlinear.py"
            )
        _interlinear_master = json.loads(INTERLINEAR_PATH.read_text(encoding="utf-8"))["units"]
    return _interlinear_master


def hebrew_width_mm(text: str, size_pt: float) -> float:
    """Measure with the real print face so a packed row can never overrun the
    text block.  Niqqud and cantillation are zero-advance marks in Noto Serif
    Hebrew, so the sum of advances is the true set width."""
    global _hebrew_metrics_font
    if _hebrew_metrics_font is None:
        from PIL import ImageFont  # imported lazily: only the DOCX build needs metrics

        _hebrew_metrics_font = ImageFont.truetype(r"C:\Windows\Fonts\NotoSerifHebrew-Regular.ttf", 1000)
    return _hebrew_metrics_font.getlength(text) / 1000 * size_pt / 72 * 25.4


def gloss_width_mm(text: str, size_pt: float) -> float:
    """Chinese glosses are full-width; Latin/digits inside them are half-width."""
    units = sum(0.5 if character.isascii() else 1.0 for character in text)
    return units * size_pt / 72 * 25.4


def align_glosses(printed_text: str, unit_id: str) -> list[dict] | None:
    """Return glossed tokens for exactly the string this page prints.

    Print strips a leading Hebrew title from some prayer and Haggadah segments,
    so the printed run can be a suffix of the glossed unit.  Align by word
    sequence rather than assuming both layers start at the same token.
    """
    record = interlinear_master().get(unit_id)
    if not record:
        return None
    tokens = record["tokens"]
    printed = [
        {"word": word, "trailing": trailing}
        for word, trailing in _printed_tokens(printed_text)
    ]
    if not printed:
        return []
    words = [token["word"] for token in tokens]
    target = [token["word"] for token in printed]
    for offset in range(0, max(1, len(words) - len(target) + 1)):
        if words[offset : offset + len(target)] == target:
            return [
                {**tokens[offset + index], "trailing": printed[index]["trailing"]}
                for index in range(len(target))
            ]
    return None


def _printed_tokens(text: str) -> list[tuple[str, str]]:
    """Same tokenisation contract as scripts/build_hebrew_interlinear.py."""
    maqqef = "\u05be"
    trailing_re = re.compile(r"[\u05c0\u05c3,:;.!?。，：；]+$")
    output: list[tuple[str, str]] = []
    for chunk in text.split():
        pieces = chunk.split(maqqef)
        for index, piece in enumerate(pieces):
            joined = index < len(pieces) - 1
            match = trailing_re.search(piece)
            trailing = match.group(0) if match else ""
            word = piece[: len(piece) - len(trailing)] if trailing else piece
            if not re.search(r"[\u05d0-\u05ea]", word):
                if output:
                    output[-1] = (output[-1][0], output[-1][1] + piece)
                continue
            output.append((word, trailing + (maqqef if joined else "")))
    return output


def pack_interlinear(tokens: list[dict], available_mm: float, *, lead_mm: float = 0.0) -> list[list[dict]]:
    """Greedy right-to-left packing of word/gloss pairs into full-width rows."""
    lines: list[list[dict]] = []
    current: list[dict] = []
    used = lead_mm
    for token in tokens:
        surface = token["word"] + token.get("trailing", "")
        width = max(
            hebrew_width_mm(surface, INTERLINEAR_HEBREW_PT),
            gloss_width_mm(token.get("glossZh", ""), INTERLINEAR_GLOSS_PT),
        ) + INTERLINEAR_GUTTER_MM
        if current and used + width > available_mm:
            lines.append(current)
            current, used = [], 0.0
        current.append({**token, "widthMm": min(width, available_mm)})
        used += width
    if current:
        lines.append(current)
    return lines


def add_interlinear_unit(
    container,
    tokens: list[dict],
    *,
    sense: str = "",
    lead: str = "",
    available_mm: float = USABLE_WIDTH_MM,
    hebrew_color: str = INK,
    gloss_color: str = MUTED,
    sense_color: str = INK,
    indent_dxa: int = TABLE_INDENT_DXA,
) -> None:
    """Render one verse/segment as stacked word blocks running right to left,
    then close it with the whole-sentence meaning."""
    lead_mm = 7.0 if lead else 0.0
    lines = pack_interlinear(tokens, available_mm, lead_mm=lead_mm)
    for line_index, line in enumerate(lines):
        cells_mm = [token["widthMm"] for token in line]
        if line_index == 0 and lead:
            cells_mm.insert(0, lead_mm)
        # Absorb the leftover into one trailing filler column instead of
        # stretching the word blocks.  A short final row then still begins at
        # the right margin rather than floating in the middle of the measure.
        slack = available_mm - sum(cells_mm)
        filler = slack > 1.0
        if filler:
            cells_mm = [*cells_mm, slack]
        elif slack > 0:
            cells_mm[-1] += slack
        table = container.add_table(rows=1, cols=len(cells_mm))
        set_table_geometry(table, cells_mm, indent_dxa=indent_dxa)
        set_borders(table, outside=False, inside=False)
        set_table_rtl(table)
        prevent_row_split(table.rows[0])
        for cell_index, cell in enumerate(table.rows[0].cells):
            set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            top = cell.paragraphs[0]
            top.alignment = WD_ALIGN_PARAGRAPH.CENTER
            top.paragraph_format.space_after = Pt(0)
            top.paragraph_format.space_before = Pt(INTERLINEAR_LINE_GAP_PT if line_index else 0)
            top.paragraph_format.line_spacing = 1.18
            bottom = cell.add_paragraph()
            bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bottom.paragraph_format.space_before = Pt(0)
            bottom.paragraph_format.space_after = Pt(0)
            bottom.paragraph_format.line_spacing = 1.0
            if line_index == 0 and lead and cell_index == 0:
                set_run_font(top.add_run(lead), FONT_UI, LABEL_PT, bold=True, color=ACCENT)
                continue
            token_index = cell_index - 1 if (line_index == 0 and lead) else cell_index
            if token_index >= len(line):
                continue  # trailing filler column
            token = line[token_index]
            set_run_font(
                top.add_run(token["word"] + token.get("trailing", "")),
                FONT_HEBREW,
                INTERLINEAR_HEBREW_PT,
                color=hebrew_color,
            )
            set_rtl(top)
            set_run_font(bottom.add_run(token.get("glossZh", "")), FONT_ZH, INTERLINEAR_GLOSS_PT, color=gloss_color)
            if line_index < len(lines) - 1:
                set_keep(bottom, next_paragraph=True)
    if sense:
        p = container.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(9)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.left_indent = Mm(5)
        p.paragraph_format.first_line_indent = Mm(-5)
        set_run_font(p.add_run("整句　"), FONT_UI, LABEL_PT, bold=True, color=ACCENT)
        add_mixed_script_text(p, sense, FONT_ZH, SENSE_PT, color=sense_color)
        set_keep(p, together=True)


def set_table_rtl(table) -> None:
    """Lay the columns out right to left so word order matches the Hebrew."""
    tbl_pr = table._tbl.tblPr
    node = tbl_pr.find(qn("w:bidiVisual"))
    if node is None:
        node = OxmlElement("w:bidiVisual")
        tbl_pr.append(node)
    node.set(qn("w:val"), "1")


def has_hebrew(text: str) -> bool:
    return bool(re.search(r"[א-ת]", text))


def add_divider(document: Document) -> None:
    """Some prayer/Haggadah sources carry a rule-only segment (e.g. "-----").
    Print it as an actual rule rather than as literal dashes with no gloss."""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    paragraph_rule(p, color=RULE, size="6")


def add_cover(document: Document, data: dict) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [USABLE_WIDTH_MM])
    set_borders(table, outside=False, inside=False)
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=500, bottom=500, start=350, end=350)
    shade(cell, ACCENT_DARK)
    for index, (text, font, size, color, bold) in enumerate((
        ("ORIGINAL-LANGUAGE READER", FONT_UI, 8, GOLD, True),
        (data["title"], FONT_ZH, 25, "FFF8ED", True),
        ("מִקְרָא עִבְרִי מְנֻקָּד", FONT_HEBREW, 20, "FFF8ED", False),
    )):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(text), font, size, color=color, bold=bold)
        if index == 2:
            set_rtl(p)
    document.add_paragraph().paragraph_format.space_after = Pt(26)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(data["subtitle"]), FONT_ZH, 12, bold=True, color=INK)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("附完整逾越節禮文流程"), FONT_ZH, 10.5, color=ACCENT)
    document.add_paragraph().paragraph_format.space_after = Pt(26)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_rule(p, color=GOLD, size="24")
    set_run_font(p.add_run("JIS B5  182 × 257 mm  ·  私人研讀"), FONT_UI, 8.5, color=MUTED)


def add_front_matter(document: Document, data: dict) -> None:
    add_label(document, "Reader architecture")
    document.add_heading("這一本怎麼使用", level=1)
    add_body(document, "全書五十課。每課固定收二十個詞、兩節背誦經文與一篇完整主讀文；第1–25課讀二十五章《希伯來聖經》，第26–50課讀二十五篇禱文或拉比文章。冊末另附完整逾越節 Haggadah 流程。")
    cards = [
        ("1", "先學本課詞表", "第1–33課就是 BBH2 第3–35章的原章詞表，詞數依課本而定；第34–50課以頻率與專名延伸補足一千詞。"),
        ("2", "背兩節經文", "每天先聽、再讀、最後遮住中文默寫；五十課恰好一百節。"),
        ("3", "讀完整原文", "聖經正文保留 WLC 母音點與 cantillation；禱文及文章保留或明示編者附點。"),
    ]
    table = document.add_table(rows=1, cols=3)
    set_table_geometry(table, [47, 47, 47])
    set_borders(table, outside=False, inside=False)
    for index, (number, title, text) in enumerate(cards):
        cell = table.cell(0, index)
        set_cell_margins(cell, top=160, bottom=160, start=130, end=130)
        shade(cell, PALE)
        p = cell.paragraphs[0]
        set_run_font(p.add_run(number), FONT_UI, 8, bold=True, color=ACCENT)
        p = cell.add_paragraph()
        set_run_font(p.add_run(title), FONT_ZH, 10.5, bold=True)
        p = cell.add_paragraph()
        set_run_font(p.add_run(text), FONT_ZH, 8, color=MUTED)
    document.add_heading("原文與音標原則", level=2)
    for text in (
        "聖經希伯來文一律列母音點；不得以現代希伯來文的無母音拼寫替代。",
        "一千詞的音標使用 Pratico–Van Pelt BBH2 課本系統；音標是學習層，不取代附點原文。",
        "人名、地名、民族名與神名另行標記並收在冊末索引。",
        "每課所列登入後網址與線上讀本共用 lesson ID；沒有校訂錄音時不以現代以色列語裝置 TTS 冒充。",
    ):
        p = document.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(text), FONT_ZH, 9.2)
    document.add_heading("版本與責任", level=2)
    add_body(document, "聖經底本為 Open Scriptures Hebrew Bible 的 WLC 4.20 文字層；繁中對照全部採《和合本修訂版》（2010，RCUV2 上帝版，© 香港聖經公會），依使用者取得的授權供私人研讀。禱文與拉比文章來源逐篇列於來源表。", size=8.8, color=MUTED)


def add_toc(document: Document, data: dict) -> None:
    add_label(document, "Contents")
    document.add_heading("五十課目錄", level=1)
    table = document.add_table(rows=1, cols=3)
    set_table_geometry(table, [14, 93, 34])
    set_borders(table, outside=True, inside=True)
    headers = ("課", "主讀文", "類型")
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        shade(cell, ACCENT)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(header), FONT_UI, 7.5, bold=True, color="FFFFFF")
    set_repeat_header(table.rows[0])
    for lesson in data["lessons"]:
        cells = table.add_row().cells
        values = (f"{lesson['lesson']:02d}", lesson["title"], "完整章" if lesson["reading"]["kind"] == "bible_chapter" else "禱文／文章")
        for i, value in enumerate(values):
            set_cell_margins(cells[i], top=50, bottom=50)
            if lesson["lesson"] % 2 == 0:
                shade(cells[i], PALE_2)
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(p.add_run(value), FONT_ZH if i == 1 else FONT_UI, 7.5, color=INK)


def proper_name_label(item: dict) -> str:
    labels = {
        "person": "人名",
        "place": "地名",
        "people_or_nation": "民族／國名",
        "divine_name_or_title": "神名／稱號",
        "divine": "神名／稱號",
        "festival_or_sacred_time": "節期／聖日",
    }
    return "、".join(labels.get(value, value) for value in item.get("properNameTypes", []))


def add_lesson_opener(document: Document, lesson: dict) -> None:
    add_label(
        document,
        f"Lesson {lesson['lesson']:02d}  ·  {lesson['reading']['kind'].replace('_', ' ')}",
        page_break_before=True,
    )
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    set_run_font(p.add_run(f"第 {lesson['lesson']:02d} 課"), FONT_UI, 11, bold=True, color=ACCENT)
    heading = document.add_heading(lesson["title"], level=1)
    paragraph_rule(heading, color=GOLD, size="14")
    route = document.add_paragraph()
    set_run_font(route.add_run("線上跟讀（登入後）  "), FONT_ZH, 8, bold=True, color=ACCENT)
    set_run_font(route.add_run(lesson["audioRoute"]), FONT_TRANSLIT, 7.6, color=MUTED)


def add_vocabulary(document: Document, lesson: dict) -> None:
    document.add_heading(f"本課 {len(lesson['vocabulary'])} 詞", level=2)
    table = document.add_table(rows=1, cols=5)
    widths = [9, 29, 30, 42, 31]
    set_table_geometry(table, widths)
    set_borders(table, outside=True, inside=True)
    headers = ("#", "附點詞形", "BBH2 音標", "繁中義", "詞類／專名")
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        shade(cell, ACCENT)
        set_cell_margins(cell, top=65, bottom=65)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(header), FONT_UI, LABEL_PT + 0.4, bold=True, color="FFFFFF")
    set_repeat_header(table.rows[0])

    for item in lesson["vocabulary"]:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        values = (
            str(item["lessonSlot"]),
            item["pointed"],
            item["textbookTransliteration"],
            item["glossZh"],
            "／".join(filter(None, [item.get("partOfSpeech", ""), proper_name_label(item)])),
        )
        for index, value in enumerate(values):
            set_cell_margins(cells[index], top=45, bottom=45, start=60, end=60)
            if item["lessonSlot"] % 2 == 0:
                shade(cells[index], PALE_2)
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index in (0, 1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            font = FONT_HEBREW if index == 1 else (FONT_TRANSLIT if index == 2 else FONT_ZH)
            size = 12.5 if index == 1 else (9.2 if index in (2, 4) else TABLE_SIZE_PT)
            if index == 3:
                add_mixed_script_text(p, value, font, size, color=INK)
            else:
                set_run_font(p.add_run(value), font, size, color=INK if index != 4 else MUTED)
            if index == 1:
                set_rtl(p)


def add_memory(document: Document, lesson: dict) -> None:
    document.add_heading("本課背誦經文", level=2)
    for item in lesson["memoryVerses"]:
        reference = document.add_paragraph()
        reference.paragraph_format.space_before = Pt(4)
        reference.paragraph_format.space_after = Pt(2)
        set_run_font(reference.add_run(f"背誦 {item['slot']}　"), FONT_UI, LABEL_PT, bold=True, color=ACCENT)
        set_run_font(reference.add_run(item["ref"]), FONT_TRANSLIT, CAPTION_PT, color=MUTED)
        set_keep(reference, next_paragraph=True)
        tokens = align_glosses(item["text"], f"bible:{item['ref']}")
        if tokens is None:
            raise SystemExit(f"逐詞對譯缺背誦經文 {item['ref']}")
        add_interlinear_unit(document, tokens, sense=item["translationZh"])


def add_bible_reading(document: Document, reading: dict) -> None:
    add_label(document, reading["ref"])
    document.add_heading("完整主讀文", level=1)
    add_hebrew(document, reading["titleHe"], size=17, color=ACCENT, bold=True)
    add_body(
        document,
        f"底本：{reading['version']}  ·  全章 {reading['verseCount']} 節  ·  逐詞繁中義在下，整句取和合本修訂版",
        size=CAPTION_PT,
        color=MUTED,
    )
    seen_translation_refs: set[str] = set()
    for verse in reading["verses"]:
        crosswalk = verse.get("translationCrosswalk") or {}
        translation_ref = str(crosswalk.get("translationRef") or "")
        is_combined_continuation = bool(
            crosswalk.get("combinedVerseRange")
            and translation_ref
            and translation_ref in seen_translation_refs
        )
        if translation_ref:
            seen_translation_refs.add(translation_ref)
        sense = (
            f"〔RCUV {crosswalk.get('translationRange')} 合併節；譯文見上一節〕"
            if is_combined_continuation
            else verse["translationZh"]
        )
        tokens = align_glosses(verse["text"], f"bible:{verse['ref']}")
        if tokens is None:
            raise SystemExit(f"逐詞對譯缺 {verse['ref']}，或與排印正文對不上")
        add_interlinear_unit(document, tokens, sense=sense, lead=str(verse["verse"]))


def clean_title_from_text(text: str, title_he: str) -> str:
    cleaned = text.strip()
    if title_he and cleaned.startswith(title_he):
        cleaned = cleaned[len(title_he):].lstrip(" \n:־–—")
    return cleaned


def add_prayer_reading(document: Document, reading: dict) -> None:
    add_label(document, reading["ref"])
    document.add_heading("完整主讀文", level=1)
    add_hebrew(document, reading["title_he"], size=17, color=ACCENT, bold=True)
    add_body(document, reading["summaryZh"], size=CAPTION_PT + 0.6, color=MUTED)
    for segment in reading["segments"]:
        text = clean_title_from_text(segment["text"], reading["title_he"])
        if not text:
            continue
        if not has_hebrew(text):
            add_divider(document)
            continue
        unit_id = f"prayer:{segment['id']}"
        tokens = align_glosses(text, unit_id)
        if tokens is None:
            raise SystemExit(f"逐詞對譯缺禱文段落 {unit_id}")
        record = interlinear_master().get(unit_id, {})
        add_interlinear_unit(document, tokens, sense=record.get("senseZh", ""))
    source = add_body(document, f"來源：{reading['source']}  ·  {reading['ref']}", size=CAPTION_PT - 0.6, color=MUTED)
    source.paragraph_format.line_spacing = 1.05
    source.paragraph_format.space_before = Pt(2)
    source.paragraph_format.space_after = Pt(0)


def add_practice(document: Document, lesson: dict, *, page_break_before=False) -> None:
    # Every lesson gets a stable practice sheet.  Prayer/article lessons use
    # the compact rhythm because some contain many proper names, but no longer
    # begin in whatever fragment remains below the preceding reading.  That
    # former flow produced pages containing only checklist items 4--6 or a
    # single orphaned item.
    compact = lesson["reading"]["kind"] != "bible_chapter"
    practice_label = add_label(document, "Close reading & practice", page_break_before=page_break_before)
    practice_label.paragraph_format.space_before = Pt(14)
    practice_heading = document.add_heading("細讀、專名與練習", level=1)
    if compact:
        practice_heading.paragraph_format.space_before = Pt(8)
        practice_heading.paragraph_format.space_after = Pt(5)
    proper = [item for item in lesson["vocabulary"] if item.get("isProperName")]
    if proper:
        proper_heading = document.add_heading("本課專名", level=2)
        if compact:
            proper_heading.paragraph_format.space_before = Pt(7)
            proper_heading.paragraph_format.space_after = Pt(4)
        for item in proper:
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Mm(6)
            p.paragraph_format.first_line_indent = Mm(-6)
            if compact:
                p.paragraph_format.space_after = Pt(2)
            set_run_font(p.add_run(item["pointed"] + "  "), FONT_HEBREW, 13.5, color=ACCENT)
            set_run_font(p.add_run(item["textbookTransliteration"] + "  "), FONT_TRANSLIT, 9.4, italic=True, color=MUTED)
            set_run_font(p.add_run(item["glossZh"] + "（" + proper_name_label(item) + "）"), FONT_ZH, 10.4)
            set_rtl(p)
    completion_heading = document.add_heading("完成本課", level=2)
    set_keep(completion_heading, next_paragraph=True)
    if compact:
        completion_heading.paragraph_format.space_before = Pt(7)
        completion_heading.paragraph_format.space_after = Pt(4)
    reading_title = lesson["title"]
    prompts = (
        "不看中文，準確朗讀二十個附點詞；說出每個詞的主要義。",
        "把兩節背誦經文各抄寫一次，圈出母音或重音與預期不同的詞。",
        f"讀完〈{reading_title}〉全文；在主讀文中標出本課詞彙。",
        "選三個動詞辨認詞幹／時式，或選三個名詞辨認性、數、狀態。",
        "登入線上讀本跟讀；沒有校訂音檔時只按課本音標自讀，不啟用現代希伯來文 TTS。",
        "用一句繁中寫出本篇主旨，再以一個希伯來關鍵詞作標題。",
    )
    for index, text in enumerate(prompts, 1):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Mm(8)
        p.paragraph_format.first_line_indent = Mm(-8)
        if compact:
            p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(f"□ {index}. "), FONT_UI, 10.2, bold=True, color=ACCENT)
        set_run_font(p.add_run(text), FONT_ZH, 10.2 if compact else 10.6)
    # A fixed note block overflowed on lessons with many proper names and made
    # rule-only pages.  Keep a useful writing area only where it fits on the
    # practice sheet; dense proper-name lessons already use the full page.
    if not compact and len(proper) <= 5:
        notes_heading = document.add_heading("讀後筆記", level=2)
        set_keep(notes_heading, next_paragraph=True)
        note_lines = max(2, 8 - len(proper))
        for _ in range(note_lines):
            p = document.add_paragraph(" ")
            paragraph_rule(p, color=RULE, size="3")


def add_haggadah(document: Document, haggadah: dict) -> None:
    add_label(document, "Complete Passover Haggadah", page_break_before=True)
    document.add_heading("附錄：完整逾越節禮文流程", level=1)
    add_hebrew(document, haggadah["title_he"], size=20, color=ACCENT, bold=True)
    add_body(document, "本附錄不計入25篇禱文／文章。依傳統十五步次序完整排列；禮儀動作與不同日況的變體保留在正文中。", size=9.2)
    add_hebrew(document, haggadah["stepConvention"]["traditionalFifteen"], size=13.5, background=PALE)
    # Preserve all fifteen named stages without wasting a separate B5 page on
    # one-line actions.  Major liturgical divisions open a fresh page; short
    # adjacent stages flow together under distinct headings.
    # Step 5 is deliberately allowed to follow the very short steps 3--4.
    # Starting it on a new page stranded those two actions on a mostly blank
    # sheet.  Nirtzah also flows after the short close of step 14; forcing it
    # to a new page stranded step 14's final blessing on a mostly empty page.
    # The indices still begin on a fresh page below.
    major_page_starts = {1}
    for display_ordinal, step in enumerate(haggadah["steps"], 1):
        step_label = add_label(
            document,
            f"Haggadah · {display_ordinal:02d}",
            page_break_before=display_ordinal in major_page_starts,
        )
        step_heading = document.add_heading(step["title_zh"], level=1)
        step_title = add_hebrew(document, step["title_he"], size=18, color=ACCENT, bold=True)
        # The stage label, Chinese heading, and Hebrew name form one semantic
        # heading block.  Each paragraph must keep with the next so a short
        # remainder at the bottom of a B5 page cannot strand the entire block
        # above the first prayer/rubric (seen previously at Haggadah step 06).
        set_keep(step_label, next_paragraph=True)
        set_keep(step_heading, next_paragraph=True)
        set_keep(step_title, next_paragraph=True)
        for segment in step["segments"]:
            text = clean_title_from_text(segment["text"], step["title_he"])
            if not text:
                continue
            if not has_hebrew(text):
                add_divider(document)
                continue
            unit_id = f"haggadah:{segment['id']}"
            tokens = align_glosses(text, unit_id)
            if tokens is None:
                raise SystemExit(f"逐詞對譯缺 Haggadah 段落 {unit_id}")
            record = interlinear_master().get(unit_id, {})
            rubric = segment.get("kind") == "rubric_or_variant"
            add_interlinear_unit(
                document,
                tokens,
                sense=record.get("senseZh", ""),
                hebrew_color=MUTED if rubric else INK,
                sense_color=MUTED if rubric else INK,
            )


def _table_header(document: Document, widths: list[float], headers: tuple[str, ...]):
    table = document.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_borders(table, outside=True, inside=True)
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        shade(cell, ACCENT)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(paragraph.add_run(header), FONT_UI, 7, bold=True, color="FFFFFF")
    set_repeat_header(table.rows[0])
    return table


def _table_row(table, values: list[tuple[str, str, float]]):
    """Append one row; each value is (text, font, size)."""

    row = table.add_row()
    prevent_row_split(row)
    for index, (text, font, size) in enumerate(values):
        cell = row.cells[index]
        set_cell_margins(cell, top=45, bottom=45)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT if font == FONT_ZH else WD_ALIGN_PARAGRAPH.CENTER
        )
        set_run_font(paragraph.add_run(text), font, size)
        if font == FONT_HEBREW:
            set_rtl(paragraph)
    return row


def _lesson_label(entry: dict) -> str:
    lesson = entry.get("lesson")
    return f"第{lesson}課" if lesson else "—"


def _attestation_label(entry: dict) -> str:
    if entry.get("attestation") == "post_biblical":
        return "後期文獻"
    if entry.get("formSource") == "lexicon":
        return "詞典引用形"
    reference = entry.get("firstOccurrence") or ""
    frequency = entry.get("frequency")
    return f"{reference}（{frequency}）" if reference else "—"


def add_reference_tables(document: Document, data: dict) -> None:
    """Render the numeral, kinship, calendar and proper-name appendix tables."""

    tables = (data.get("referenceTables") or {}).get("tables") or []
    for table in tables:
        add_label(document, table["titleZh"], page_break_before=True)
        add_hebrew(document, table["titleHe"], size=18, color=ACCENT, bold=True)
        document.add_heading(table["titleZh"], level=1)
        # These paragraphs quote Hebrew inline (בֵּית אָב, חָם／חָמוֹת), and the
        # Chinese face has no Hebrew glyphs, so they go through the mixed-script
        # writer rather than add_body.
        add_mixed_script_text(document.add_paragraph(), table["intro"], FONT_ZH, 9.2)
        for group in table["groups"]:
            document.add_heading(group["titleZh"], level=2)
            for caption in ("note", "source"):
                if group.get(caption):
                    add_mixed_script_text(
                        document.add_paragraph(), group[caption], FONT_ZH, 8.6, color=MUTED
                    )
            shape = group["shape"]
            if shape == "gender_pair":
                grid = _table_header(
                    document,
                    [12, 28, 22, 28, 22, 29],
                    ("數", "陽性形", "音標", "陰性形", "音標", "繁中"),
                )
                for entry in group["entries"]:
                    _table_row(
                        grid,
                        [
                            (entry["value"], FONT_UI, 8),
                            (entry["masculine"]["pointed"], FONT_HEBREW, 11),
                            (entry["masculine"]["transliteration"], FONT_TRANSLIT, 7.8),
                            (entry["feminine"]["pointed"], FONT_HEBREW, 11),
                            (entry["feminine"]["transliteration"], FONT_TRANSLIT, 7.8),
                            (entry["glossZh"], FONT_ZH, 8.4),
                        ],
                    )
            elif shape == "month":
                grid = _table_header(
                    document,
                    [16, 30, 26, 30, 39],
                    ("序位", "附點形", "音標", "繁中月名", "首見（次數）"),
                )
                for entry in group["entries"]:
                    _table_row(
                        grid,
                        [
                            (entry.get("order", ""), FONT_UI, 8),
                            (entry["pointed"], FONT_HEBREW, 11),
                            (entry["transliteration"], FONT_TRANSLIT, 7.8),
                            (entry["glossZh"], FONT_ZH, 8.4),
                            (_attestation_label(entry), FONT_ZH, 7.6),
                        ],
                    )
            elif shape == "name":
                grid = _table_header(
                    document,
                    [30, 26, 40, 30, 15],
                    ("附點形", "音標", "繁中", "首見（次數）", "課"),
                )
                for entry in group["entries"]:
                    _table_row(
                        grid,
                        [
                            (entry["pointed"], FONT_HEBREW, 10.6),
                            (entry["transliteration"], FONT_TRANSLIT, 7.6),
                            (entry["glossZh"], FONT_ZH, 8.2),
                            (_attestation_label(entry), FONT_ZH, 7.4),
                            (_lesson_label(entry), FONT_UI, 7.4),
                        ],
                    )
            else:
                grid = _table_header(
                    document,
                    [20, 30, 26, 36, 29],
                    ("項", "附點形", "音標", "繁中", "首見（次數）"),
                )
                for entry in group["entries"]:
                    _table_row(
                        grid,
                        [
                            (entry.get("value", ""), FONT_UI, 8),
                            (entry["pointed"], FONT_HEBREW, 11),
                            (entry["transliteration"], FONT_TRANSLIT, 7.8),
                            (entry["glossZh"], FONT_ZH, 8.4),
                            (_attestation_label(entry), FONT_ZH, 7.6),
                        ],
                    )
            for entry in group["entries"]:
                if not entry.get("note"):
                    continue
                head = entry.get("pointed") or entry.get("masculine", {}).get("pointed", "")
                # add_body would set the whole line in the Chinese face, which has
                # no Hebrew glyphs and silently falls back to a UI font.
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(2)
                add_mixed_script_text(paragraph, f"{head}　{entry['note']}", FONT_ZH, 8.0, color=MUTED)


def add_back_indices(document: Document, data: dict) -> None:
    add_label(document, "Colophon", page_break_before=True)
    heading = document.add_heading("來源與成品檢核", level=1)
    heading.paragraph_format.page_break_before = True
    for text in (
        "50課；每課固定20詞；總計1,000詞。",
        "每課2節背誦；總計100節。",
        "第1–25課為25個完整聖經章；第26–50課為25篇完整禱文或文章。",
        "冊末逾越節禮文按完整流程另列，不抵充25篇；其後另附數字、親屬、曆法與分類專名四張對照表。",
        "聖經希伯來文保留完整母音點與 cantillation；全部詞彙列 BBH2 課本式音標。",
        "線上音訊與紙本共用 lesson ID；正式錄音須經校訂，不以裝置TTS計入完成。",
    ):
        p = document.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(text), FONT_ZH, 9.2)
    document.add_heading("主要來源", level=2)
    add_body(document, "希伯來聖經：Open Scriptures Hebrew Bible / Westminster Leningrad Codex text, WLC 4.20。", size=8.5, color=MUTED)
    add_body(document, "繁中聖經對照：《和合本修訂版》（2010，RCUV2 上帝版，© 香港聖經公會；私人授權使用）。", size=8.5, color=MUTED)
    add_body(document, "詞彙：使用者已授權的 Pratico–Van Pelt BBH2 排序與本計畫頻率延伸；音標欄使用 BBH2 系統。", size=8.5, color=MUTED)
    add_body(document, "禱文、拉比文章與 Haggadah：逐篇來源見資料層；私人使用授權已由使用者確認。", size=8.5, color=MUTED)


def build(data: dict) -> Path:
    document = Document()
    configure(document)
    add_cover(document, data)
    page_break(document)
    add_front_matter(document, data)
    page_break(document)
    add_toc(document, data)

    for lesson in data["lessons"]:
        add_lesson_opener(document, lesson)
        add_vocabulary(document, lesson)
        add_memory(document, lesson)
        if lesson["reading"]["kind"] == "bible_chapter":
            add_bible_reading(document, lesson["reading"])
        else:
            add_prayer_reading(document, lesson["reading"])
        add_practice(document, lesson)

    add_haggadah(document, data["haggadah"])
    add_reference_tables(document, data)
    add_back_indices(document, data)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


def main() -> None:
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    # A layout-only proof render may cut the book down to a few lessons; the
    # count gate below is skipped for those, never for a release build.
    subset = os.environ.get("HBO_LESSON_SUBSET")
    if subset:
        wanted = {int(value) for value in subset.split(",")}
        data = {**data, "lessons": [lesson for lesson in data["lessons"] if lesson["lesson"] in wanted]}
        global OUTPUT_PATH
        OUTPUT_PATH = OUTPUT_DIR / "hebrew-original-reader-layout-proof.docx"
        print(build(data))
        return
    counts = data["counts"]
    expected = {
        "lessons": 50,
        "vocabulary": 1000,
        "memoryVerses": 100,
        "bibleChapters": 25,
        "prayersOrArticles": 25,
    }
    for key, value in expected.items():
        if counts.get(key) != value:
            raise ValueError(f"{key}: expected {value}, got {counts.get(key)}")
    print(build(data))


if __name__ == "__main__":
    main()
