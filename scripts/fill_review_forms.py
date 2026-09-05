# -*- coding: utf-8 -*-
"""填「學期進度審查表」（05）與「期末進度審查表」（06）的學生端欄位。

一學期幾次：**一次**。《玄奘大學研究生助學金實施辦法》第六條第二款第二目
「學期補助結束後，應繳交論文進度審查表」，05 表註 3 也寫「本審查表請於每學期
第十八週前擲回研發處」。研究助學金補助十個月＝兩學期，所以 05 交兩次
（115-1 第一次、115-2 第二次），06 期末審查表在補助最後一學期期末交一次。

表上「完成度%」「審查意見」「審查結果」「與學生互動情形」都是**指導教授填的**，
這裡一律留白，只填學生資料與論文題目。

🚨 合併儲存格在 row.cells 會重複出現同一個 tc，照索引寫會寫到隔壁格去。
🚨 **有選項的格子要就地勾選，不要用 write() 覆蓋。** write() 先清空整格，
   多段落的選項欄會被壓成一行、少數選項會整個消失（04 表的「學生與指導教授
   互動情形」六個項目就是這樣被我刪成一個，而學制欄的「□碩專班」也不見了）。
🚨 產出檔留 Drive 不進 repo（含學號）。

  python -X utf8 scripts/fill_review_forms.py
"""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

BASE = Path("G:/我的雲端硬碟/玄奘/博一上/獎學金/玄奘助學金")
NAME, SID, GRADE = "張辰瑋", "DB1153002", "一年級"
ADVISOR = "釋昭慧"
TITLE = "從彼岸向此岸的轉向：台灣佛教與基督教公共性之宗教史比較研究（1920–2020）"
CJK, EN = "標楷體", "Times New Roman"

JOBS = [
    (BASE / "05-學年度進度審查表(docx) (1).docx",
     BASE / "05-進度審查表（已填‧115-1第一次）.docx", 1, 1),
    (BASE / "05-學年度進度審查表(docx) (1).docx",
     BASE / "05-進度審查表（已填‧115-2第二次）.docx", 2, 2),
    (BASE / "06-期末進度審查表(docx).docx",
     BASE / "06-期末進度審查表（已填）.docx", None, None),
]


def write(cell, text, *, size=12):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.font.name = EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)


def edit_par(par, old, new):
    """就地把段落裡的 old 換成 new。跨 run 的文字要先併起來再換——
    Word 常把一行切成好幾個 run，逐 run 比對永遠對不上。"""
    if old not in par.text or not par.runs:
        return False
    par.runs[0].text = par.text.replace(old, new)
    for r in par.runs[1:]:
        r.text = ""
    return True


def uniq(row):
    seen, out = [], []
    for c in row.cells:
        if id(c._tc) in seen:
            continue
        seen.append(id(c._tc))
        out.append(c)
    return out


def fill(doc, semester, nth):
    for p in doc.paragraphs:
        t = p.text
        if "學年度第" in t and "學期研究生助學金" in t and p.runs:
            p.runs[0].text = (t.replace("玄奘大學", "玄奘大學 115")
                               .replace("學年度第　學期", f"學年度第 {semester} 學期")
                               .replace("    學年度", " 學年度"))
            for r in p.runs[1:]:
                r.text = ""
        elif "學年度研究生助學金" in t and p.runs:          # 06 表沒有學期欄
            p.runs[0].text = t.replace("玄奘大學", "玄奘大學 115").replace("    學年度", " 學年度")
            for r in p.runs[1:]:
                r.text = ""
        elif "第　　次進度審查表" in t and p.runs:
            p.runs[0].text = t.replace("第　　次", f"第 {nth} 次")
            for r in p.runs[1:]:
                r.text = ""

    pairs = {"學生姓名": NAME, "學號": SID, "年級": GRADE,
             "指導教授": ADVISOR, "論文題目": TITLE}
    # 🚨「指導教授」在表上出現兩次：上面是資料欄，最底下那一列是**簽章欄**。
    #    照字面比對會把名字打進簽章格，那格要留白給老師親簽。每個欄位只填第一次。
    filled = set()
    for row in doc.tables[0].rows:
        cells = uniq(row)
        for k, cell in enumerate(cells):
            label = cell.text.strip().replace("\n", "")
            for key, val in pairs.items():
                if label == key and key not in filled and k + 1 < len(cells):
                    write(cells[k + 1], val, size=10 if key == "論文題目" else 12)
                    filled.add(key)
            if label.startswith("系級") and k + 1 < len(cells):
                for par in cells[k + 1].paragraphs:
                    edit_par(par, "□博士班", "■博士班")
            if label.startswith("論文形式") and k + 1 < len(cells):
                for par in cells[k + 1].paragraphs:
                    edit_par(par, "□學術研究論文研究計畫", "■學術研究論文研究計畫")


def main():
    for src, out, sem, nth in JOBS:
        doc = Document(src)
        fill(doc, sem, nth)
        doc.save(out)
        print(f"  → {out.name}")


if __name__ == "__main__":
    main()
