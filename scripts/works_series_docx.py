# -*- coding: utf-8 -*-
"""/works 叢書（HTML 章節書）→ 設計排版 Word。

支援三套叢書（同一 book-head/chapter HTML schema）：
  genesis-philosophy（創生哲學 15 卷：含章首故事引子三變體、論證分析圖、章末摘要）
  world-religions-intro／sinographic-literature（兩本教科書講義）

版式：封面（叢書名/卷題/副題/主旨）／目次（Word TOC 欄位，F9 更新得頁碼）／
章起新頁、章題置中標楷體；章首引子取主選變體整段標楷體；論證圖縮排小字；
本節依據對話等出處小字；頁尾置中頁碼。輸出 docx 與 html 同目錄同名。

用法：python works_series_docx.py [slug ...]（預設跑全部三套）
"""
import json
import re
import sys
from pathlib import Path

from lxml import html as lhtml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from dadaodao_book_export import add_page_numbers, style_fonts

ROOT = Path(__file__).resolve().parent.parent
PUB = ROOT / 'public'
KAI = 'DFKai-SB'
MING = 'PMingLiU'
GRAY = RGBColor(0x66, 0x66, 0x66)

SLUGS = ['genesis-philosophy', 'world-religions-intro', 'sinographic-literature']


def run_ea(par, text, font=MING, size=11.0, bold=False, italic=False, sup=False, color=None):
    r = par.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if sup:
        r.font.superscript = True
    if color:
        r.font.color.rgb = color
    return r


def emit_inline(par, el, font, size, bold=False, italic=False, color=None):
    """遞迴輸出元素內文：strong→粗、em→斜、sup→上標、a→純文字。"""
    if el.text:
        run_ea(par, el.text, font, size, bold, italic, color=color)
    for ch in el:
        tag = ch.tag if isinstance(ch.tag, str) else ''
        b = bold or tag == 'strong'
        i = italic or tag == 'em'
        if tag == 'sup':
            if ch.text_content():
                run_ea(par, ch.text_content(), font, max(size - 3, 7), sup=True, color=color)
        elif tag == 'br':
            par.add_run().add_break()
        else:
            emit_inline(par, ch, font, size, b, i, color=color)
        if ch.tail:
            run_ea(par, ch.tail, font, size, bold, italic, color=color)


def para(doc, el=None, text=None, font=MING, size=11.0, bold=False, indent=None,
         first_indent=Cm(0.85), spacing=1.7, align=None, color=None,
         before=None, after=Pt(4)):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if indent is not None:
        pf.left_indent = indent
    pf.line_spacing = spacing
    if before is not None:
        pf.space_before = before
    if after is not None:
        pf.space_after = after
    if align is not None:
        p.alignment = align
    if el is not None:
        emit_inline(p, el, font, size, bold, color=color)
    elif text:
        run_ea(p, text, font, size, bold, color=color)
    return p


def heading(doc, text_or_el, level, center=False):
    h = doc.add_heading(level=level)
    if center:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_before = Pt(20 if level == 1 else 14)
    h.paragraph_format.space_after = Pt(12 if level == 1 else 6)
    size = {1: 17, 2: 13.5, 3: 12}.get(level, 12)
    if isinstance(text_or_el, str):
        run_ea(h, text_or_el, KAI, size, bold=True)
    else:
        emit_inline(h, text_or_el, KAI, size, bold=True)
    return h


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ea(p, '目　次', KAI, 18, bold=True)
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-1" \h \z \u')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '（在 Word 中按 F9 更新目次即顯示各章頁碼）'
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def emit_body_children(doc, sec, skip=()):
    """章／序內容主迴圈（fable 與 subtitle 已另行處理時以 skip 略過）。"""
    for ch in sec:
        tag = ch.tag if isinstance(ch.tag, str) else ''
        cls = ch.get('class', '') or ''
        if ch in skip or tag in ('h1', 'h2') or 'chapter-subtitle' in cls or 'chapter-fable' in cls:
            continue
        if tag == 'h3':
            heading(doc, ch, 2)
        elif tag == 'h4':
            heading(doc, ch, 3)
        elif tag == 'p':
            if 'section-source' in cls:
                para(doc, el=ch, size=8.5, first_indent=None, spacing=1.3, color=GRAY, before=Pt(2), after=Pt(10))
            elif 'fable-bridge' in cls:
                para(doc, el=ch, font=KAI, size=10.5, indent=Cm(1.0), first_indent=None, spacing=1.6, color=GRAY)
            else:
                para(doc, el=ch)
        elif tag == 'blockquote':
            for sub in ch.iter('p'):
                para(doc, el=sub, font=KAI, size=10.5, indent=Cm(1.0), first_indent=None, spacing=1.6)
        elif tag in ('ul', 'ol'):
            for li in ch.iter('li'):
                p = para(doc, size=10, indent=Cm(0.8), first_indent=None, spacing=1.5, after=Pt(2))
                run_ea(p, '• ', MING, 10)
                emit_inline(p, li, MING, 10)
        elif cls == 'argmap' or 'argmap' in cls:
            for step in ch:
                scls = step.get('class', '') or ''
                txt = step.text_content().strip()
                if not txt:
                    continue
                if 'arg-op' in scls:
                    para(doc, text=txt, size=9, first_indent=None, spacing=1.1,
                         align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=Pt(1))
                else:
                    bold = 'arg-conclusion' in scls
                    para(doc, el=step, size=9.5, indent=Cm(0.8), first_indent=None,
                         spacing=1.4, bold=bold, color=None if bold else GRAY, after=Pt(3))
        elif tag == 'div':
            emit_body_children(doc, ch, skip=skip)


def build_book(book, kicker='', author='張辰瑋'):
    src = PUB / book['file'].lstrip('/')
    tree = lhtml.fromstring(src.read_text(encoding='utf-8'))

    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Cm(2.5)
    s.left_margin = s.right_margin = Cm(2.8)
    style_fonts(doc, 'Heading 1', KAI, size=17, bold=True)
    style_fonts(doc, 'Heading 2', KAI, size=13.5, bold=True)
    style_fonts(doc, 'Heading 3', KAI, size=12, bold=True)
    style_fonts(doc, 'Normal', MING, size=11)

    head = tree.cssselect('header.book-head')
    kick = ttl = sub = thesis = ''
    if head:
        g = lambda sel: (head[0].cssselect(sel) or [None])[0]
        kick = g('.book-kicker').text_content().strip() if g('.book-kicker') is not None else kicker
        ttl = g('.book-title').text_content().strip() if g('.book-title') is not None else book['title']
        sub = g('.book-sub').text_content().strip() if g('.book-sub') is not None else ''
        thesis = g('.book-thesis').text_content().strip() if g('.book-thesis') is not None else ''
    # 封面
    for txt, font, size, bold, before, align in [
        (kick, KAI, 13, False, 160, WD_ALIGN_PARAGRAPH.CENTER),
        (ttl, KAI, 26, True, 16, WD_ALIGN_PARAGRAPH.CENTER),
        (sub, KAI, 13, False, 14, WD_ALIGN_PARAGRAPH.CENTER),
        (author, KAI, 13, False, 60, WD_ALIGN_PARAGRAPH.CENTER),
    ]:
        if not txt:
            continue
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(before)
        run_ea(p, txt, font, size, bold)
    if thesis:
        p = para(doc, text=thesis, size=10, indent=Cm(1.2), first_indent=None,
                 spacing=1.6, color=GRAY, before=Pt(40))
        p.paragraph_format.right_indent = Cm(1.2)
    doc.add_page_break()
    add_toc(doc)

    first = True
    for sec in tree.cssselect('section'):
        cls = sec.get('class', '') or ''
        if 'chapter-recap' in cls:
            heading(doc, '本章摘要', 2)
            emit_body_children(doc, sec)
            continue
        h2 = sec.cssselect('h2')
        title = h2[0] if h2 else None
        doc.add_page_break()
        if title is not None:
            heading(doc, title, 1, center=True)
        sub_el = sec.cssselect('.chapter-subtitle')
        if sub_el:
            para(doc, el=sub_el[0], font=KAI, size=12, first_indent=None,
                 align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=Pt(16))
        # 章首故事引子：取主選（第一個）變體，整段標楷體
        fable = sec.cssselect('.chapter-fable .fable-variant')
        if fable:
            v = fable[0]
            cap = v.get('data-fable-title', '')
            if cap:
                para(doc, text=f'章首引子──{cap}', font=KAI, size=10.5, bold=True,
                     indent=Cm(1.0), first_indent=None, spacing=1.5, before=Pt(6), after=Pt(4))
            for pel in v.iter('p'):
                para(doc, el=pel, font=KAI, size=10.5, indent=Cm(1.0),
                     first_indent=None, spacing=1.65,
                     color=GRAY if 'fable-bridge' in (pel.get('class') or '') else None)
            para(doc, text='＊　＊　＊', size=9, first_indent=None,
                 align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=Pt(10))
        emit_body_children(doc, sec)
        first = False

    add_page_numbers(doc)
    out = src.with_suffix('.docx')
    doc.save(str(out))
    return out


def main(slugs):
    for slug in slugs:
        mf = PUB / f'content/works/{slug}-books.json'
        d = json.loads(mf.read_text(encoding='utf-8'))
        books = d.get('books') or [b for g in d.get('groups', []) for b in g['books']]
        print(f'== {slug}（{len(books)} 冊）')
        for b in books:
            out = build_book(b)
            print(f'  {b["id"]} {b["title"]} → {out.name} ({out.stat().st_size:,} bytes)')


if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    main(sys.argv[1:] or SLUGS)
