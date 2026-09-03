# -*- coding: utf-8 -*-
"""玄奘大學課程教學大綱 → Word。

**以學校原本的表單 docx 為範本，只替換內容**，因此版式、欄寬、字型、
核取方塊樣式都與原檔完全一致（不重畫表格）。範本＝使用者的
「114.2基督宗教概論2026.03.03.docx」。

核取方塊的三種狀態（照原檔的做法）：
  已勾  = <w:sym w:font="Wingdings 2" w:char="F052"/>（打勾符號，無文字）
  ☒     = 字面 '☒'（Segoe UI Symbol）
  未勾  = 字面 '☐'（Segoe UI Symbol）

成品是 docx，依 docs/repo-hygiene.md 不進 git，輸出到 Drive：
  G:\\我的雲端硬碟\\資料\\知識圖工作室\\教學\\{課程資料夾}\\

用法：
  python scripts/course_syllabus_docx.py                 # 由範本重出整份（會覆蓋既有檔）
  python scripts/course_syllabus_docx.py --schedule-only # 只改既有檔的「授課進度與內容」

⚠ 檔案若正在 Word 中開啟會寫入失敗；使用者可能正在手改該檔，
  要改既有檔時一律走 --schedule-only，不要整份覆蓋掉別人的編輯。
"""
import copy
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TEMPLATE = Path(r'C:\Users\user\Desktop\114.2基督宗教概論2026.03.03.docx')
DRIVE = Path(r'G:\我的雲端硬碟\資料\知識圖工作室\教學')

# ── 課程資料 ────────────────────────────────────────────────────────────────
WORLD_RELIGIONS = {
    'folder': '115-1_世界宗教文化導論',
    'filename': '115.1世界宗教文化導論',
    'year': '115',
    'term': '1',
    'course_name': '世界宗教文化導論',
    'course_code': 'PPA001',
    'objective': (
        '本課程是宗教與文化學系的入門必修，目標是讓學生建立兩項基本能力。'
        '其一是說得清楚宗教是什麼：理解宗教作為人類對神聖的一種經驗與探詢，'
        '以及由此形成的一系列組織與活動；掌握宗教學史上幾種經典定義及其爭議，'
        '並理解十九世紀歐洲原本以拜神界定宗教，如何在接觸儒教、佛教等傳統之後被迫修正。'
        '其二是分得出宗教的類型：課程以泛神論、多神論、一神論、實用神論四大信仰型態為主軸，'
        '輔以語言民族型分類（泛亞伯拉罕諸教、泛印度諸教等），'
        '將全世界信徒人數一千萬以上的宗教安置在同一張地圖上，'
        '並涵蓋馬爾杜克信仰、阿頓信仰、吠陀信仰、早期耶和華信仰等古代宗教，'
        '以及希臘、羅馬、北歐、迦南等地的傳統宗教，'
        '理解各型態之間的歷史關聯與相互轉化，而非把世界宗教讀成彼此無關的一格一格。'
        '課程最後兩次上課回到現代世界的宗教處境與臺灣宗教現況，'
        '使學生能運用全學期所學的架構分析自己身處的宗教環境。'
    ),
    # 教學方法：填要打勾的項目（其餘自動改為 ☐）
    'methods_on': [
        '講述', '媒體融入教學', '問題導向學習', '合作學習',
        '即時互動', '對話教學法', '個別指導',
    ],
    # 學習評量：{項目關鍵字: 百分比}，未列者一律 ☐ 0%
    'assessment_on': {
        '出席': '25%',
        '課堂參與': '25%',
        '口頭報告': '25%',
        '期末考-筆試': '25%',
    },
    # 授課進度：(日期, 章號)。章名一律取自講義章節檔的 <h2>，不另行改寫；
    # 章號 None 表示非授課週（自由學習等），此時第二欄直接寫該字串。
    'chapters_html': 'public/content/works/world-religions-intro/chapters-wr2',
    'schedule': [
        ('9月20日', 1), ('9月20日', 2),
        ('10月4日', 3), ('10月4日', 4),
        ('10月18日', 5), ('10月18日', 6),
        ('11月1日', 7), ('11月1日', 8),
        ('11月15日', 9), ('11月15日', 10),
        ('11月29日', 11), ('11月29日', 12),
        ('12月13日', 13), ('12月13日', 14),
        ('12月27日', 15), ('12月27日', '第十六章　臺灣宗教（下）：當代圖像——四種型態在同一座島上，期末考'),
        ('1月10日', '自由學習：個別與老師討論自評'),
        ('1月10日', '自由學習：個別與老師討論自評'),
    ],
    'foreign_text': '有',
    'textbooks': [
        '休斯頓．史密士（Huston Smith）著，劉安雲譯。《人的宗教：人類偉大的智慧傳統》。台北：立緒文化，1998。',
        '涂爾幹（Émile Durkheim）著，芮傳明、趙學元譯。《宗教生活的基本形式》。台北：桂冠圖書，1992。',
        '伊利亞德（Mircea Eliade）著，楊素娥譯。《聖與俗：宗教的本質》。台北：桂冠圖書，2001。',
        '詹姆斯（William James）著，蔡怡佳、劉宏信譯。《宗教經驗之種種》。台北：立緒文化，2001。',
        '韋伯（Max Weber）著，康樂、簡惠美譯。《宗教社會學》。台北：遠流出版，1993。',
        '瞿海源。《台灣宗教變遷的社會政治分析》。台北：桂冠圖書，1997。',
    ],
    'course_url': '',   # 留空：不列線上資源
}

# ── 學期週次表 ──────────────────────────────────────────────────────────────
# 由「PPA001 首堂 9/20（日）且屬雙週」反推：學期第一週為 9/7–9/13。
# 九個雙週日（9/20…1/10）與使用者給的日期完全吻合，因此單週六與每週三照同一張表推。
WED = ['9月9日', '9月16日', '9月23日', '9月30日', '10月7日', '10月14日',
       '10月21日', '10月28日', '11月4日', '11月11日', '11月18日', '11月25日',
       '12月2日', '12月9日', '12月16日', '12月23日', '12月30日', '1月6日']
SAT_ODD = ['9月12日', '9月26日', '10月10日', '10月24日', '11月7日',
           '11月21日', '12月5日', '12月19日', '1月2日']   # ⚠ 10/10 為國慶日，需確認調課


def weekly(chapters, midterm_at=9, exam_tail='，期末考'):
    """每週兩節、十八週：八章 → 期中考 → 八章 → 期末考。"""
    rows, ch = [], 1
    for w in range(1, 19):
        if w == midterm_at:
            rows.append((WED[w - 1], '期中考'))
        elif w == 18:
            rows.append((WED[w - 1], '期末考'))
        else:
            rows.append((WED[w - 1], ch))
            ch += 1
    return rows


def biweekly_sat(free='自由學習：個別與老師討論自評'):
    """單週六四節、九次：每次兩章，第九次自由學習（每次佔兩列）。"""
    rows, ch = [], 1
    for i, d in enumerate(SAT_ODD):
        if i == 8:
            rows += [(d, free), (d, free)]
        else:
            rows += [(d, ch), (d, ch + 1)]
            ch += 2
    return rows


WR_INTRO_DAY = {
    'folder': '115-1_世界宗教文化導論',
    'filename': '115.1世界宗教文化導論（BBE275宗教系1A）',
    'year': '115', 'term': '1',
    'course_name': '世界宗教文化導論',
    'course_code': 'BBE275',
    'klass': '宗教與文化學系1年A班',
    'elective': '專業必修',
    'hours': '2', 'credits': '2',
    'objective': WORLD_RELIGIONS['objective'],
    'methods_on': WORLD_RELIGIONS['methods_on'],
    'assessment_on': {'出席': '20%', '課堂參與': '20%',
                      '期中考-筆試': '30%', '期末考-筆試': '30%'},
    'chapters_html': 'public/content/works/world-religions-intro/chapters-wr2',
    'schedule': weekly(None),
    'foreign_text': '有',
    'textbooks': WORLD_RELIGIONS['textbooks'],
    'course_url': '',
}

CHINESE = {
    'folder': '宗教系國文講義',
    'filename': '115.1國文',
    'year': '115', 'term': '1',
    'course_name': '國文',
    'course_code': 'PPA066',
    'klass': '宗教與文化學系二年制在職專班1年A班',
    'elective': '通識必修',
    'hours': '2', 'credits': '2',
    'objective': (
        '本課程是宗教與文化學系的大學國文，以「漢字書寫圈」取代民族國家文學史的框架，'
        '並以宗教為貫穿軸線。授課範圍涵蓋中國、日本、韓半島、越南與臺灣的漢文書寫：'
        '這些地區共用同一套文字、同一批典故與同一個文本庫，卻各自發展出書寫本地語言的辦法'
        '（假名、吏讀、諺文、喃字、白話字），只讀其中一國會看不見這個結構。'
        '課程的第二條線是宗教：現存最早的漢字文獻是占卜紀錄，最早的正典化工程是經學，'
        '世界史上規模最大的翻譯工程是佛典漢譯，白話文學的源頭是講經與變文，'
        '戲曲從祭壇長出來，而宗教至今仍是漢文最穩固的使用場所。'
        '學生除了讀作品，還要學會用一組固定欄位介紹一部作品'
        '（作者與時代、成書與版本、文體與語言、內容、宗教脈絡、今日何處可讀），'
        '並追問它「被誰、在什麼場合、用什麼方式使用」——抄、誦、演、考、教。'
    ),
    'methods_on': ['講述', '媒體融入教學', '問題導向學習', '合作學習',
                   '即時互動', '對話教學法', '個別指導', '實地考察/參訪'],
    'assessment_on': {'出席': '25%', '課堂參與': '25%',
                      '心得或作業撰寫': '25%', '期末考-筆試': '25%'},
    'chapters_html': 'public/content/works/sinographic-literature/chapters',
    'schedule': biweekly_sat(),
    'foreign_text': '無',
    'textbooks': [
        '裘錫圭。《文字學概要》。臺北：萬卷樓，1995。',
        '郭慶藩輯。《莊子集釋》。北京：中華書局。',
        '黃征、張涌泉。《敦煌變文校注》。北京：中華書局，1997。',
        '王國維。《宋元戲曲史》。上海：商務印書館，1915。',
        '賴永祥。《教會史話》。臺南：人光出版社。',
        '張妙娟。《開啟心眼——〈臺灣府城教會報〉與長老教會的基督徒教育》。臺南：人光出版社，2005。',
    ],
    'course_url': '',
}

CHRISTIANITY = {
    'folder': '115-1_基督宗教概論',
    'filename': '115.1基督宗教概論',
    'year': '115', 'term': '1',
    'course_name': '基督宗教概論',
    'course_code': 'BBE150',
    'klass': '宗教與文化學系2年A班',
    'elective': '專業必修',
    'hours': '2', 'credits': '2',
    'objective': (
        '本課程從宗教學而非神學的立場介紹基督宗教：不問「這是不是真的、我該怎麼信」，'
        '而問「它是怎麼形成的、在人的生活裡怎麼運作」。'
        '課程先確立四項核心傳統（天主教、東正教、東方正統、新教）與四條分歧軸線'
        '（權威在哪裡、救恩怎麼運作、教會是什麼、物質能否承載神聖），'
        '再依序處理拿撒勒人耶穌的史料與處境、初代教會、正典的形成與兩千年的讀法、'
        '大公會議與教義的成形、教父與東方基督教世界、中世紀西方教會、宗教改革、'
        '近代的四條回應路線，以及二十世紀的梵二、普世運動與南方基督教。'
        '課程刻意補上中文教材慣常壓縮的三塊：東方基督教世界（衣索比亞、敘利亞、'
        '亞美尼亞、景教）、聖行（禮儀、聖事、齋期、喪禮）、以及藝術與物質文化。'
        '最後兩章回到臺灣的基督宗教與生死觀，'
        '使學生能用同一組欄位介紹任何一個基督宗教傳統，並指出某個說法漏掉了什麼。'
    ),
    'methods_on': ['講述', '媒體融入教學', '問題導向學習', '合作學習',
                   '即時互動', '對話教學法', '個別指導'],
    'assessment_on': {'出席': '20%', '課堂參與': '20%',
                      '期中考-筆試': '30%', '期末考-筆試': '30%'},
    'chapters_html': 'public/content/works/christianity-intro/chapters',
    'schedule': weekly(None),
    'foreign_text': '有',
    'textbooks': [
        '陶理（Tim Dowley）主編，李伯明、林牧野譯。《基督教二千年史》。香港：海天書樓，1997。',
        '大衛‧班特利‧哈特（David Bentley Hart）著，黃恩霖譯。《基督教的故事》。臺中：好讀出版，2013。',
        '哈維‧寇克斯（Harvey G. Cox）著，孫尚揚譯。《基督宗教》。臺北：麥田出版。',
        '林鴻信。《基督宗教思想史》。臺北：國立臺灣大學出版中心，2017。',
        '奧爾森（Roger E. Olson）著，吳瑞誠、徐成德譯。《神學的故事》。臺北：校園書房出版社。',
        '鄭仰恩。《定根本土的臺灣基督教》。臺南：人光出版社，2005。',
    ],
    'course_url': '',
}

COURSES = {
    'world-religions-intro': WORLD_RELIGIONS,
    'world-religions-day': WR_INTRO_DAY,
    'chinese': CHINESE,
    'christianity': CHRISTIANITY,
}


# ── 儲存格編輯工具（保留原有格式）────────────────────────────────────────────
def _clone_run_format(src_run, text):
    """複製 src_run 的格式，換成新文字，回傳新的 <w:r>。"""
    r = copy.deepcopy(src_run._r)
    for child in list(r):
        if child.tag == qn('w:t') or child.tag == qn('w:sym'):
            r.remove(child)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def set_cell(cell, lines):
    """把儲存格內容換成 lines（list[str]，一行一段），沿用原第一個 run 的格式。"""
    if isinstance(lines, str):
        lines = [lines]
    src = None
    for p in cell.paragraphs:
        if p.runs:
            src = p.runs[0]
            break
    keep = cell.paragraphs[0]
    for p in cell.paragraphs[1:]:
        p._p.getparent().remove(p._p)
    for r in list(keep.runs):
        r._r.getparent().remove(r._r)
    if src is None:
        for line in lines:
            keep.add_run(line)
        return
    keep._p.append(_clone_run_format(src, lines[0]))
    for line in lines[1:]:
        newp = copy.deepcopy(keep._p)
        for child in list(newp):
            if child.tag == qn('w:r'):
                newp.remove(child)
        newp.append(_clone_run_format(src, line))
        keep._p.addnext(newp)
        keep = cell.paragraphs[-1]
    return


def set_checkbox(cell, checked):
    """改寫核取方塊狀態，保留後方的標籤文字。"""
    p = cell.paragraphs[0]
    runs = p._p.findall(qn('w:r'))
    if not runs:
        return
    first = runs[0]
    label_run = runs[1] if len(runs) > 1 else None
    for child in list(first):
        if child.tag in (qn('w:t'), qn('w:sym')):
            first.remove(child)
    rPr = first.find(qn('w:rPr'))
    if checked:
        # 照原檔：Wingdings 2 的打勾符號，字型要跟著換
        if rPr is not None:
            fonts = rPr.find(qn('w:rFonts'))
            if fonts is None:
                fonts = OxmlElement('w:rFonts')
                rPr.insert(0, fonts)
            for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
                fonts.set(qn(a), 'Wingdings 2')
            fonts.set(qn('w:eastAsia'), 'Wingdings 2')
        sym = OxmlElement('w:sym')
        sym.set(qn('w:font'), 'Wingdings 2')
        sym.set(qn('w:char'), 'F052')
        first.append(sym)
    else:
        if rPr is not None:
            fonts = rPr.find(qn('w:rFonts'))
            if fonts is None:
                fonts = OxmlElement('w:rFonts')
                rPr.insert(0, fonts)
            for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
                fonts.set(qn(a), 'Segoe UI Symbol')
            fonts.set(qn('w:eastAsia'), 'Segoe UI Symbol')
        t = OxmlElement('w:t')
        t.text = '☐'
        first.append(t)
    # 確保方塊與標籤之間有兩個空白（原檔慣例）
    if label_run is not None:
        tt = label_run.find(qn('w:t'))
        if tt is not None and tt.text and tt.text.strip() == '':
            tt.set(qn('xml:space'), 'preserve')
            tt.text = '  '


def cell_label(cell):
    """取儲存格中核取方塊後方的標籤文字。"""
    return cell.paragraphs[0].text.replace('☐', '').replace('☒', '').strip()


def chapter_title(c, n):
    """章名一律取自講義章節檔的 <h2>，不在授課表裡另行改寫。"""
    cdir = Path(__file__).resolve().parent.parent / c['chapters_html']
    html = (cdir / f'ch{n:02d}.html').read_text(encoding='utf-8')
    return re.sub(r'<[^>]+>', '', re.search(r'<h2>(.*?)</h2>', html).group(1)).strip()


def write_schedule(c, T2):
    for i, (date, topic) in enumerate(c['schedule']):
        if isinstance(topic, int):
            topic = chapter_title(c, topic)
        row = T2.rows[i + 2]
        set_cell(row.cells[0], [f'{i + 1:02d}', date])
        set_cell(row.cells[1], topic)


def update_schedule(c):
    """只改既有檔的「授課進度與內容」，其餘欄位（可能已被手動編輯）原樣保留。"""
    out = DRIVE / c['folder'] / f"{c['filename']}.docx"
    if not out.exists():
        raise SystemExit(f'找不到 {out}，請先跑一次完整產生')
    doc = Document(str(out))
    write_schedule(c, doc.tables[2])
    doc.save(str(out))
    return out


# ── 產生 ────────────────────────────────────────────────────────────────────
def build(c):
    outdir = DRIVE / c['folder']
    outdir.mkdir(parents=True, exist_ok=True)
    out = outdir / f"{c['filename']}.docx"
    shutil.copyfile(TEMPLATE, out)

    doc = Document(str(out))
    T0, T1, T2, T3, T4 = doc.tables[:5]

    # ── 課程基本資料 ──
    set_cell(T0.rows[1].cells[1], c['year'])
    set_cell(T0.rows[1].cells[3], c['term'])
    set_cell(T0.rows[4].cells[1], c['course_name'])
    set_cell(T0.rows[4].cells[5], c['course_code'])
    set_cell(T0.rows[5].cells[1], c['course_name'])
    if c.get('klass'):
        set_cell(T0.rows[2].cells[3], c['klass'])
    if c.get('hours'):
        set_cell(T0.rows[3].cells[1], c['hours'])
    if c.get('credits'):
        set_cell(T0.rows[3].cells[3], c['credits'])
    if c.get('elective'):
        set_cell(T0.rows[3].cells[5], c['elective'])

    # ── 課程目標 ──
    set_cell(T1.rows[1].cells[0], c['objective'])

    # ── 教學方法（r3–r9，4 欄）──
    on = c['methods_on']
    for ri in range(3, 10):
        for ci in range(4):
            cell = T1.rows[ri].cells[ci]
            label = cell_label(cell)
            if not label:
                continue
            set_checkbox(cell, any(k in label for k in on))

    write_schedule(c, T2)

    # ── 學習評量方式（r2–r24）──
    for ri in range(2, len(T3.rows)):
        cell = T3.rows[ri].cells[0]
        label = cell_label(cell)
        if not label:
            continue
        hit = next((v for k, v in c['assessment_on'].items() if k in label), None)
        set_checkbox(cell, hit is not None)
        set_cell(T3.rows[ri].cells[1], hit or '0%')

    # ── 學習參考資源 ──
    set_cell(T4.rows[1].cells[1], c['foreign_text'])
    set_cell(T4.rows[2].cells[1], c['textbooks'])
    set_cell(T4.rows[3].cells[1], c['course_url'] or '')

    doc.save(str(out))
    return out


if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    args = [a for a in sys.argv[1:] if not a.startswith('--')]
    only = '--schedule-only' in sys.argv
    for k in (args or list(COURSES)):
        c = COURSES[k]
        # Word 開檔時會留下 ~$ 開頭的鎖檔（檔名會被截去前幾個字，故用 glob 比對）
        locks = list((DRIVE / c['folder']).glob('~$*.docx'))
        if locks:
            raise SystemExit(f'⚠ {c["filename"]}.docx 正在 Word 中開啟（鎖檔 {locks[0].name}）。'
                             '請先在 Word 關閉這個檔再執行，以免蓋掉尚未存檔的編輯。')
        print('✔', update_schedule(c) if only else build(c))
