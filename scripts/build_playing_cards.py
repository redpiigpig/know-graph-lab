#!/usr/bin/env python3
"""可裁切的印刷撲克牌：佛教人物、基督宗教人物各五十四張。

版面、卡框、雙面鏡像全部沿用單字卡那一套（`build_flashcards`），同一把裁刀、
同一組印表機設定通吃，所以這裡不重算任何尺寸——只換卡面。

與單字卡的三個差別：

1. **背面是統一的牌背。** 撲克牌之所以能玩，靠的就是每張背面長得一樣；資訊
   全部印在正面。單字卡是正面問、背面答，兩者的正反面意義相反。
2. **框色按花色不按課次**，用的是四色牌（黑桃黑、紅心紅、方塊藍、梅花綠）。
   四色牌是既有的撲克牌慣例，也讓「同一點數的四張」在桌面上一眼分得開。
3. **點數才是分類軸**：K 佛陀、Q 菩薩……A 當代弘傳。花色只作四色分組，
   不表位階——這一條印在說明卡上，免得被讀成價值排序。

54 張填不滿 7 張紙（56 格），剩下的兩格做成說明卡：一張點數分類對照，
一張使用與授權說明。空著會裁出兩張正面全白的牌。
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))

import build_flashcards as base  # noqa: E402

ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data/playingCards"
OUTPUT_DIR = ROOT / "output/playing-cards"

FONT_INDEX = "Segoe UI Symbol"   # 唯一同時有 ♠♥♦♣ 與數字的內建字型
FONT_ZH = base.FONT_ZH
FONT_LATIN = "Noto Serif"        # 梵語轉寫的 ā ṇ ṃ ś ṭ，MingLiU 與標楷體都沒有

SUITS = {
    "spade":   {"glyph": "♠", "zh": "黑桃", "color": "1B1B1B"},
    "heart":   {"glyph": "♥", "zh": "紅心", "color": "C8102E"},
    "diamond": {"glyph": "♦", "zh": "方塊", "color": "1E6FD9"},
    "club":    {"glyph": "♣", "zh": "梅花", "color": "1F6B4A"},
}
SUIT_ORDER = ("spade", "heart", "diamond", "club")
RANK_ORDER = ("K", "Q", "J", "10", "9", "8", "7", "6", "5", "4", "3", "2", "A")
SUIT_GLYPHS = "".join(item["glyph"] for item in SUITS.values())
JOKER_COLOR = "7B3FA0"
NOTE_COLOR = "4A4A4A"

ICON_MM = 26
INDEX_PT = 17
ORIG_PT = 9.5
DATES_PT = 9.5
LABEL_PT = 10
# 中文名從三個字（藥師佛）到十三個字（若望二十三世與若望保祿二世）都有，
# 一律設同一級會溢出，所以按長度縮。
ZH_STEPS = ((4, 20), (6, 18), (8, 16), (10, 13), (99, 11))


def index_font(text: str) -> str:
    """角標字型。Segoe UI Symbol 是唯一有 ♠♥♦♣ 的內建字型，卻沒有中日韓字，
    所以「大鬼」「小鬼」「※」交給中文字型——不換就靜靜回退到 NotoSansJP-Thin。"""

    return FONT_INDEX if all(ch in SUIT_GLYPHS or ch.isascii() for ch in text) else FONT_ZH


def zh_size(text: str) -> float:
    for limit, size in ZH_STEPS:
        if len(text) <= limit:
            return size
    return ZH_STEPS[-1][1]


def openmoji():
    """借希伯來 matcher 的 OpenMoji 載入，不重寫一份下載與檔名解析。"""

    import importlib.util

    spec = importlib.util.spec_from_file_location(
        "match_flashcard_images", ROOT / "scripts/match_flashcard_images.py"
    )
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    module.ensure_openmoji()
    index = {}
    for entry in json.loads((base.CACHE / "flashcards/openmoji.json").read_text("utf-8")):
        index.setdefault(entry["annotation"], entry["hexcode"])
    return index


def variants():
    """`emoji_variant_images` 的 PNG 抓取，同樣不重寫一份。"""

    import importlib.util

    spec = importlib.util.spec_from_file_location(
        "emoji_variant_images", ROOT / "scripts/emoji_variant_images.py"
    )
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def icon_path(index: dict, name: str) -> Path:
    """圖以 OpenMoji 的**本名**指定，查無就當場停下來。

    名字打錯若靜靜略過，卡上就是一張沒有圖、或更糟、隔壁那張圖——這一類錯印在
    紙上看起來完全正常，拿牌的人無從察覺。
    """

    if ":" in name:
        # 別套繪者的畫法（`noto-v1:wheel-of-dharma`）。OpenMoji 的法輪畫在一個
        # 紫色方框裡，跟其餘留白的圖擺在一起像貼錯；概念沒變，只換畫法。
        prefix, slug = name.split(":", 1)
        return variants().png_for(prefix, slug)
    hexcode = index.get(name)
    if hexcode is None:
        raise SystemExit(f"OpenMoji 查無此圖：{name!r}")
    path = base.IMAGE_DIR / f"{hexcode}.png"
    if not path.exists():
        raise SystemExit(f"OpenMoji 有名字沒圖檔：{name!r}（{hexcode}）")
    return path


def load_deck(path: Path, index: dict) -> tuple[dict, list[dict]]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    ranks = payload["ranks"]
    by_key = {(card["rank"], card["suit"]): card for card in payload["cards"]}
    missing = [(rank, suit) for rank in RANK_ORDER for suit in SUIT_ORDER
               if (rank, suit) not in by_key]
    if missing or len(payload["cards"]) != 52:
        raise SystemExit(f"牌組不完整：{len(payload['cards'])} 張，缺 {missing}")

    cards: list[dict] = []
    for rank in RANK_ORDER:
        for suit in SUIT_ORDER:
            card = by_key[(rank, suit)]
            cards.append({
                "index": f"{rank}{SUITS[suit]['glyph']}",
                "color": SUITS[suit]["color"],
                "icon": icon_path(index, ranks[rank]["icon"]),
                "zh": card["zh"],
                "orig": card["orig"],
                "dates": card["dates"],
                "label": ranks[rank]["label"],
            })
    for joker in payload["jokers"]:
        cards.append({
            "index": joker["rank"],
            "color": JOKER_COLOR,
            "icon": icon_path(index, joker["icon"]),
            "zh": joker["zh"],
            "orig": joker["orig"],
            "dates": joker["dates"],
            "label": joker["label"],
        })
    cards.extend(note_cards(payload))
    return payload, cards


def note_cards(payload: dict) -> list[dict]:
    """把 54 張補滿 56 格的兩張說明卡。空格會裁出正面全白的牌。"""

    ranks = payload["ranks"]
    table = "　".join(f"{rank} {ranks[rank]['label']}" for rank in RANK_ORDER)
    suits = "　".join(f"{item['glyph']} {item['zh']}" for item in SUITS.values())
    return [
        {"index": "※", "color": NOTE_COLOR, "icon": None,
         "zh": "點數分類對照", "orig": table, "dates": "大鬼‧小鬼另立一類",
         "label": "說明卡", "wrap": True},
        {"index": "※", "color": NOTE_COLOR, "icon": None,
         "zh": "花色與用法", "orig": suits,
         "dates": "花色不表位階，只作四色分組；點數才是分類軸。"
                  "插圖 OpenMoji（CC BY-SA 4.0）畫的是該點數的類別，不是人物肖像。",
         "label": "說明卡", "wrap": True},
    ]


def fill_front(cell, card: dict, place: tuple[int, int, int]) -> None:
    cell = base.framed(cell, card["color"], place)

    top = cell.add_paragraph()
    top.alignment = WD_ALIGN_PARAGRAPH.LEFT
    top.paragraph_format.space_after = Pt(0)
    base.write(top, card["index"], index_font(card["index"]), INDEX_PT,
               color=card["color"], bold=True)

    if card["icon"]:
        picture = cell.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.paragraph_format.space_after = Pt(5)
        picture.add_run().add_picture(str(card["icon"]), width=Mm(ICON_MM))
    else:
        base.blank(cell, 14)

    name = cell.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(3)
    base.write(name, card["zh"], FONT_ZH, zh_size(card["zh"]), bold=True)

    orig = cell.add_paragraph()
    orig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    orig.paragraph_format.space_after = Pt(2)
    base.write(orig, card["orig"], FONT_ZH if card.get("wrap") else FONT_LATIN,
               8.5 if card.get("wrap") else ORIG_PT, color=base.MUTED)

    dates = cell.add_paragraph()
    dates.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dates.paragraph_format.space_after = Pt(0)
    base.write(dates, card["dates"], FONT_ZH,
               8 if card.get("wrap") else DATES_PT, color=base.MUTED)

    base.blank(cell, 8)
    foot = cell.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.paragraph_format.space_after = Pt(0)
    base.write(foot, card["label"], FONT_ZH, LABEL_PT, color=card["color"])

    bottom = cell.add_paragraph()
    bottom.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    bottom.paragraph_format.space_after = Pt(0)
    base.write(bottom, card["index"], index_font(card["index"]), INDEX_PT,
               color=card["color"], bold=True)


def fill_back(cell, deck: dict, back_icon: Path, place: tuple[int, int, int]) -> None:
    """每一張都一樣——牌背有差異的就不是撲克牌了。"""

    cell = base.framed(cell, deck["backColor"], place)
    base.blank(cell, 10)
    picture = cell.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.space_after = Pt(6)
    picture.add_run().add_picture(str(back_icon), width=Mm(34))
    title = cell.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    base.write(title, deck["title"], FONT_ZH, 15, color=deck["backColor"], bold=True)
    motto = cell.add_paragraph()
    motto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    motto.paragraph_format.space_after = Pt(0)
    base.write(motto, deck["backMotto"], FONT_ZH, 9, color=base.MUTED)
    base.blank(cell, 10)


def add_cover(document: Document, deck: dict, total: int) -> None:
    cuts_h = "、".join(f"{base.MARGIN_V_MM + base.CARD_H_MM * i:.0f}" for i in range(3))
    cuts_v = "、".join(f"{base.CARD_W_MM * i:.2f}".rstrip("0").rstrip(".") for i in range(1, 4))
    ranks = deck["ranks"]
    lines = [
        (deck["title"], 30, base.INK, 12),
        (f"{deck['subtitle']}・{total // 8} 組雙面・每頁 8 張", 13, base.MUTED, 20),
        ("列印：A4 橫式，雙面列印選「沿長邊翻頁」，縮放設為 100%（不要選「符合頁面大小」）。",
         11, base.INK, 6),
        (f"裁切：不印裁切線。自紙張上緣量，橫向切在 {cuts_h} mm；自左緣量，縱向切在 {cuts_v} mm。"
         f"成品每張 {base.CARD_W_MM:.2f}×{base.CARD_H_MM:.0f} mm。"
         f"卡框比裁切線內縮 {base.FRAME_INSET_MM:.0f} mm，裁歪一兩毫米只會讓白邊不等寬，不會切到框。",
         11, base.INK, 6),
        ("框色按花色四色分：♠ 黑、♥ 紅、♦ 藍、♣ 綠，鬼牌紫、說明卡灰。"
         "四色牌讓同一點數的四張在桌面上一眼分得開。", 11, base.INK, 6),
        ("正面：點數與花色、類別插圖、中文名、原文名、生卒年、點數類別。"
         "背面全牌統一，所以這副牌可以照一般撲克牌玩法使用。", 11, base.INK, 6),
        ("點數即分類：" + "；".join(f"{r} {ranks[r]['label']}" for r in RANK_ORDER) + "。"
         "花色不表位階，只作四色分組。", 11, base.INK, 20),
        ("插圖畫的是該點數的類別，不是人物肖像——同一點數的四張共用一張圖是刻意的。",
         10, base.MUTED, 10),
        ("圖片來源：OpenMoji 17.0.0（openmoji.org），CC BY-SA 4.0；"
         "法輪一圖取自 Noto Emoji 舊版（Apache 2.0）。", 9.5, base.MUTED, 4),
    ]
    lines.extend((text, 9.5, base.MUTED, 4) for text in deck["sources"])
    for text, size, color, after in lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.space_before = Pt(0)
        base.write(paragraph, text, base.FONT_UI, size, color=color, bold=size >= 30)
    base.page_break(document)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.write(note, "（此頁留白，供雙面列印對齊）", base.FONT_UI, 9.5, color=base.MUTED)


def build(deck: dict, cards: list[dict], back_icon: Path, output: Path) -> Path:
    document = Document()
    base.configure(document)
    add_cover(document, deck, len(cards))

    per_page = base.COLS * base.ROWS
    for start in range(0, len(cards), per_page):
        page = cards[start : start + per_page]
        for side in ("front", "back"):
            base.page_break(document)
            table = base.new_grid(document)
            for offset, card in enumerate(page):
                row, column = divmod(offset, base.COLS)
                target_column = column if side == "front" else base.COLS - 1 - column
                place = (target_column, row, 1000 + start * 2 + offset * 2 + (side == "back"))
                cell = table.cell(row, target_column)
                if side == "front":
                    fill_front(cell, card, place)
                else:
                    fill_back(cell, deck, back_icon, place)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output


DECKS = {
    "buddhist":  {"data": DATA_DIR / "buddhist.json",  "output": "buddhist-playing-cards.docx",
                  "backColor": "8B5A2B"},
    "christian": {"data": DATA_DIR / "christian.json", "output": "christian-playing-cards.docx",
                  "backColor": "8C1D2C"},
}


def main() -> None:
    parser = argparse.ArgumentParser(description="產生撲克牌（A4 橫式、每頁 8 張、雙面）")
    parser.add_argument("--deck", choices=sorted(DECKS), default="buddhist")
    args = parser.parse_args()

    spec = DECKS[args.deck]
    index = openmoji()
    payload, cards = load_deck(spec["data"], index)
    payload["backColor"] = spec["backColor"]
    back_icon = icon_path(index, payload["backIcon"])
    path = build(payload, cards, back_icon, OUTPUT_DIR / spec["output"])

    sheets = len(cards) // (base.COLS * base.ROWS)
    print(f"  {payload['title']}：{len(cards)} 張（52 + 2 鬼牌 + 2 說明卡），"
          f"正反共 {sheets * 2} 頁，每張 {base.CARD_W_MM:.2f}×{base.CARD_H_MM:.0f} mm")
    print(path)


if __name__ == "__main__":
    main()
