# -*- coding: utf-8 -*-
"""115-1 週課表 → .docx，交給使用者列印／存查。

橫向單頁彩色課表（14 節 × 七天，跨節的課合併儲存格），配色與 /works 那張課表
artifact 一致：教課藍、修課綠、家教橙。GRID 與 artifact 的 EVENTS 是同一份資料，
改這裡也要一起改。

★ 標的是自己授課的科目；其餘是修課與家教。
台神課與家教不對齊玄奘節次，格子畫的是概略位置，實際時間看格內標示。

學期行事曆（假日授課與校外行程）只放 Google 日曆「115-1 玄奘教課」，不印進這份檔案；
真的要另出一頁時加 --calendar。

用法：python scripts/semester_schedule_docx.py [--calendar]
輸出：G:\\我的雲端硬碟\\玄奘\\博一上\\115-1 週課表.docx
"""
import datetime
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = r'G:\我的雲端硬碟\玄奘\博一上\115-1 週課表.docx'

W = '一二三四五六日'
PERIODS = [(1, '08:30–09:20'), (2, '09:25–10:15'), (3, '10:25–11:15'), (4, '11:20–12:10'),
           (5, '12:10–13:10'), (6, '13:10–14:00'), (7, '14:10–15:00'), (8, '15:10–16:00'),
           (9, '16:10–17:00'), (10, '17:10–18:00'), (11, '18:30–19:15'), (12, '19:15–20:00'),
           (13, '20:10–20:55'), (14, '20:55–21:40')]

# 三類的（底色, 主文字色, 註記色）；與 artifact 的 teach／study／tutor 同一組色
KIND = {
    'teach': ('DFEDF3', RGBColor(0x17, 0x51, 0x6A), RGBColor(0x3E, 0x6E, 0x85)),
    'study': ('E6EFDD', RGBColor(0x3F, 0x5F, 0x30), RGBColor(0x5C, 0x77, 0x4D)),
    'tutor': ('F6E7DE', RGBColor(0x8A, 0x45, 0x2A), RGBColor(0xA0, 0x66, 0x4C)),
}
HEAD_FILL = 'E9EEF1'
REST_FILL = 'F2F5F7'

# (星期, 起節, 迄節, 類別, 名稱, 教師與教室, 課號與班別)
GRID = [
    (1, 2, 4, 'study', '宗教研究基本問題與研究方法', '根瑟馬庫斯　妙然 208',
     'BBJ001‧博士班 1A‧專必‧3 學分'),
    (1, 9, 10, 'tutor', '小三家教', '16:00–17:30', ''),
    (1, 12, 13, 'tutor', '國中家教', '19:00–20:30', ''),
    (2, 1, 2, 'study', '初階宗教學日文文獻選讀', '倪杰　妙然 M201',
     'BBA224‧碩士班 1A‧專選‧2 學分'),
    (2, 7, 9, 'study', '希伯來文 II', '曾宗盛　台神（士林）',
     '14:30–17:20‧神研道碩選修‧3 學分'),
    (3, 1, 2, 'teach', '★世界宗教文化導論', '妙然 401',
     'BBE275‧宗教系 1A‧專必‧2 學分'),
    (3, 3, 4, 'study', '唯識思想專題研討', '釋昭慧、丹增南卓　妙然 206',
     'BBJ029‧博士班 1A‧專選‧2 學分'),
    (3, 8, 9, 'teach', '★基督宗教概論', '妙然 401',
     'BBE150‧宗教系 2A‧專必‧2 學分'),
    (3, 10, 11, 'tutor', '小五家教', '17:00–18:30', ''),
    (3, 12, 13, 'tutor', '國中家教', '19:00–20:30', ''),
    (5, 10, 11, 'tutor', '小五家教', '17:00–18:30', ''),
    (5, 12, 13, 'tutor', '國中家教', '19:00–20:30', ''),
    (6, 1, 4, 'study', '〔單週〕宗教學理論與方法（一）', '根瑟馬庫斯　妙然 308',
     'JJA051‧碩專班 1A‧專必‧2 學分'),
    (6, 6, 9, 'teach', '〔單週〕★國文', '妙然 206',
     'PPA066‧二專 1A‧通必‧2 學分'),
    (7, 6, 9, 'teach', '〔雙週〕★世界宗教文化導論', '妙然 401',
     'PPA001‧二專 1A‧專必‧2 學分'),
]

# 假日授課：單週六國文八次、雙週日世界宗教八次，第 17／18 週自學
SAT = ['2026-09-12', '2026-09-26', '2026-10-10', '2026-10-24',
       '2026-11-07', '2026-11-21', '2026-12-05', '2026-12-19']
SUN = ['2026-09-20', '2026-10-04', '2026-10-18', '2026-11-01',
       '2026-11-15', '2026-11-29', '2026-12-13', '2026-12-27']

OTHER = [
    ('2026-09-05', '09:00–12:00', '弘誓學院地藏法會', '佛教弘誓學院'),
    ('2026-09-05', '18:30–20:30', '濟南教會演講', '濟南基督長老教會'),
    ('2026-09-06', '整日', '天母國三家教', '天母'),
    ('2026-09-13', '整日', '天母國三家教', '天母'),
    ('2026-09-18', '整日', '大專研究佛學研討會（協助）', ''),
    ('2026-09-19', '整日', '天母國三家教', '天母'),
    ('2026-10-03', '整日', '天母國三家教（待 Judy 確認）', '天母'),
    ('2026-10-09', '整日', '天母國三家教', '天母'),
    ('2026-10-11', '整日', '天母國三家教', '天母'),
    ('2026-10-16', '整日（–10/17）', '與 Soe San 出遊', ''),
]


def calendar_rows():
    rows = [(d, '13:10–17:00', f'國文（PPA066）　第 {i} 次', '玄奘大學 妙然 206')
            for i, d in enumerate(SAT, 1)]
    rows.append(('2027-01-02', '13:10–17:00', '國文（PPA066）　自學', '玄奘大學 妙然 206'))
    rows += [(d, '13:10–17:00', f'世界宗教文化導論（PPA001）　第 {i} 次', '玄奘大學 妙然 401')
             for i, d in enumerate(SUN, 1)]
    rows.append(('2027-01-10', '13:10–17:00', '世界宗教文化導論（PPA001）　自學', '玄奘大學 妙然 401'))
    rows += OTHER
    rows.sort(key=lambda r: (r[0], r[1]))
    return rows


def _font(run, name, size, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color


def shade(cell, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def head(doc, txt, size=17):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(txt)
    r.bold = True
    _font(r, '微軟正黑體', size)


def note(doc, txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _font(p.add_run(txt), '新細明體', 8, RGBColor(0x60, 0x60, 0x60))


def legend(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for kind, label in (('teach', ' 教課　'), ('study', ' 修課　'), ('tutor', ' 家教　')):
        _font(p.add_run('■'), '微軟正黑體', 9, KIND[kind][1])
        _font(p.add_run(label), '微軟正黑體', 8.5, RGBColor(0x40, 0x40, 0x40))
    _font(p.add_run('教課 4 門 8 學分‧修課 5 門 12 學分‧家教 6 場／週‧教室皆在妙然樓'),
          '新細明體', 8, RGBColor(0x60, 0x60, 0x60))


def cell_lines(c, lines, center=True, bold_first=True):
    """lines = [(文字, 字型, 級數, 顏色), …]；第一段寫進儲存格既有的 paragraph。"""
    c.text = ''
    for i, (txt, font, size, color) in enumerate(lines):
        p = c.paragraphs[0] if i == 0 else c.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1 if i == 0 else 0)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(txt)
        r.bold = bold_first and i == 0
        _font(r, font, size, color)


def cell_text(c, main, sub='', bold=False, size=8.5, center=True):
    lines = [(main, '微軟正黑體', size, None)]
    if sub:
        lines.append((sub, '新細明體', 7, RGBColor(0x55, 0x55, 0x55)))
    cell_lines(c, lines, center, bold_first=bold)


def week_table(doc):
    t = doc.add_table(rows=len(PERIODS) + 1, cols=8)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_text(t.cell(0, 0), '節次 / 時間', bold=True, size=9)
    shade(t.cell(0, 0), HEAD_FILL)
    for i, d in enumerate(W):
        cell_text(t.cell(0, i + 1), d, bold=True, size=10)
        shade(t.cell(0, i + 1), HEAD_FILL)
    for i, (n, tm) in enumerate(PERIODS):
        cell_text(t.cell(i + 1, 0), '午休' if n == 5 else f'第 {n} 節', tm, size=8)
        if n == 5:
            for c in range(8):
                shade(t.cell(i + 1, c), REST_FILL)
    for day, p0, p1, kind, main, sub, tag in GRID:
        fill, ink, sub_ink = KIND[kind]
        c = t.cell(p0, day).merge(t.cell(p1, day)) if p0 != p1 else t.cell(p0, day)
        lines = [(main, '微軟正黑體', 8.5, ink)]
        if sub:
            lines.append((sub, '新細明體', 7, sub_ink))
        if tag:
            lines.append((tag, '新細明體', 6.5, sub_ink))
        cell_lines(c, lines)
        shade(c, fill)
    t.columns[0].width = Cm(2.2)
    for c in range(1, 8):
        t.columns[c].width = Cm(3.3)


def calendar_table(doc):
    rows = calendar_rows()
    t = doc.add_table(rows=len(rows) + 1, cols=5)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['日期', '星期', '時間', '項目', '地點']):
        cell_text(t.cell(0, i), h, bold=True, size=9.5)
        shade(t.cell(0, i), HEAD_FILL)
    for i, (d, tm, name, loc) in enumerate(rows, 1):
        dt = datetime.date.fromisoformat(d)
        cell_text(t.cell(i, 0), f'{dt.month}/{dt.day}', size=9)
        cell_text(t.cell(i, 1), W[dt.weekday()], size=9)
        cell_text(t.cell(i, 2), tm, size=8.5)
        cell_text(t.cell(i, 3), name, size=9, center=False)
        cell_text(t.cell(i, 4), loc, size=8.5, center=False)
    for w, c in zip([Cm(1.8), Cm(1.3), Cm(3.0), Cm(11.5), Cm(6.0)], t.columns):
        c.width = w
    return len(rows)


def build(with_calendar=False):
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, m, Cm(1.2))

    st = doc.styles['Normal']
    st.font.name = '新細明體'
    st.font.size = Pt(9)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), '新細明體')

    head(doc, '辰瑋　115 學年度第 1 學期　一週課表')
    legend(doc)
    week_table(doc)
    note(doc, '★ 為授課；其餘為修課與家教。第 5 節為午休 12:10–13:10，第 10 節後夜間休息 18:00–18:30。'
              '節次時間依〈玄奘大學課堂時間表〉；台神課與家教不對齊玄奘節次，以格內標示時間為準。'
              '單／雙週依學期週次：第一週 9/7–9/13。')

    n = 0
    if with_calendar:
        doc.add_page_break()
        head(doc, '辰瑋　115 學年度第 1 學期　學期行事曆')
        note(doc, '假日授課與校外行程；每週固定的平日課程與家教不列於此，見前頁週課表。')
        n = calendar_table(doc)

    doc.save(OUT)
    return OUT, len(GRID), n


if __name__ == '__main__':
    sys.stdout.reconfigure(encoding='utf-8')
    p, g, n = build('--calendar' in sys.argv)
    print(f'OK {p}  (週課表 {g} 格' + (f'；行事曆 {n} 筆)' if n else ')'))
