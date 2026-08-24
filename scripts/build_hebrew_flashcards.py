#!/usr/bin/env python3
"""Build the printable Hebrew vocabulary flashcard deck.

The sheet follows the English tutoring deck already in use: A4 landscape, eight
cards to a page, front sheet then back sheet, duplex.  Two things differ.

No cutting lines are printed.  A hairline that prints a millimetre off shows up
on the finished card as a crooked edge, so instead the grid is an exact quarter
of the page wide and half of it tall — 74.25 x 105 mm — and a guillotine set to
those two measurements cuts every card square.

The back sheet mirrors the column order (4-3-2-1 instead of 1-2-3-4).  Printed
duplex, that is what puts each meaning behind its own word; without it every
card would carry a neighbour's translation.

Pictures come from ``hebrew-card-images.json`` and are optional by design: a
third of the deck has one, and a card with no honest picture prints without one
rather than borrowing an approximate image.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
VOCAB = ROOT / "data/originalReaders/vocabulary/hebrew-1000.json"
GLOSSES = ROOT / "output/source-cache/original-readers/hebrew-full/hebrew-gloss-zh-reviewed-by-lemma.json"
IMAGES = ROOT / "output/source-cache/flashcards/hebrew-card-images.json"
IMAGE_DIR = ROOT / "output/source-cache/flashcards/openmoji-618"
OUTPUT_DIR = ROOT / "output/flashcards"
OUTPUT = OUTPUT_DIR / "hebrew-flashcards-1000.docx"

PAGE_W_MM = 297.0
PAGE_H_MM = 210.0
COLS = 4
ROWS = 2
CARD_W_MM = PAGE_W_MM / COLS      # 74.25
# Both figures are measured, not derived.  The renderer reserves more vertical
# space than the declared margins account for, and it moves the second row to a
# page of its own well before the arithmetic says the rows should stop fitting;
# a 10 mm margin fails where 5 mm holds, and 97 mm rows fail where 94 mm hold.
# So: 5 mm margin, 94 mm rows.  Horizontal cuts fall at 5, 99 and 193 mm,
# vertical cuts at 74.25, 148.5 and 222.75 mm — an even grid a guillotine can
# follow with no printed rules.
CARD_H_MM = 94.0
MARGIN_V_MM = 5.0

FONT_HEBREW = "Noto Serif Hebrew"
FONT_ZH = "MingLiU"
FONT_UI = "MingLiU"
INK = "1B1B1B"
MUTED = "8A8A8A"

HEBREW_PT = 54
# The glosses run from two characters to twenty-seven, so the meaning line is
# sized to fit the card rather than set at one size and allowed to overflow.
ZH_STEPS = ((6, 24), (10, 20), (16, 16), (22, 13), (99, 11))
POS_PT = 12
LESSON_PT = 10
IMAGE_MM = 32

POS_ZH = {
    "noun": "名詞",
    "verb": "動詞",
    "adjective": "形容詞",
    "adverb": "副詞",
    "pronoun": "代名詞",
    "preposition": "介系詞",
    "conjunction": "連接詞",
    "particle": "質詞",
    "proper_name": "專名",
    "particle_or_preposition": "質詞／介系詞",
    "interrogative_particle": "疑問質詞",
    "prepositional_phrase": "介系詞片語",
    "adverbial_phrase": "副詞片語",
    "conjunction_phrase": "連接詞片語",
}


def set_rfonts(run, font: str) -> None:
    """Pin every Word script slot so Hebrew never falls back to a UI font."""

    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{slot}"), font)
    for theme in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        fonts.attrib.pop(qn(f"w:{theme}"), None)


def write(paragraph, text: str, font: str, size: float, *, color: str = INK, bold: bool = False, rtl: bool = False):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    set_rfonts(run, font)
    if rtl:
        r_pr = run._element.get_or_add_rPr()
        mark = OxmlElement("w:rtl")
        mark.set(qn("w:val"), "1")
        r_pr.append(mark)
        p_pr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        p_pr.append(bidi)
    return run


def blank(cell, points: float) -> None:
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    write(paragraph, "", FONT_UI, points)


def strip_borders(table) -> None:
    """Print no rules at all; the cut is measured, not traced."""

    properties = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        borders.append(element)
    properties.append(borders)


def new_grid(document: Document):
    table = document.add_table(rows=ROWS, cols=COLS)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    strip_borders(table)
    for row in table.rows:
        # One height declaration only.  Setting row.height and then appending a
        # second w:trHeight leaves two competing rules in the XML, and the
        # renderer honoured the wrong one — the front sheet lost its bottom row.
        row.height = Mm(CARD_H_MM)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            cell.width = Mm(CARD_W_MM)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            margins = OxmlElement("w:tcMar")
            for edge, value in (("top", 170), ("start", 140), ("bottom", 170), ("end", 140)):
                node = OxmlElement(f"w:{edge}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            cell._tc.get_or_add_tcPr().append(margins)
    return table


def clear(cell) -> None:
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)


def fill_front(cell, entry: dict) -> None:
    clear(cell)
    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    write(paragraph, entry["pointed"], FONT_HEBREW, HEBREW_PT, rtl=False)
    blank(cell, 20)
    footer = cell.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    write(footer, f"第 {entry['lesson']} 課", FONT_UI, LESSON_PT, color=MUTED)


def zh_size(gloss: str) -> float:
    for limit, size in ZH_STEPS:
        if len(gloss) <= limit:
            return size
    return ZH_STEPS[-1][1]


def fill_back(cell, entry: dict, gloss: str, picture: Path | None) -> None:
    clear(cell)
    if picture:
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run().add_picture(str(picture), width=Mm(IMAGE_MM))
    else:
        blank(cell, 20)
    meaning = cell.add_paragraph()
    meaning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meaning.paragraph_format.space_after = Pt(3)
    write(meaning, gloss, FONT_ZH, zh_size(gloss))
    pos = cell.add_paragraph()
    pos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pos.paragraph_format.space_after = Pt(0)
    write(pos, POS_ZH.get(entry["partOfSpeech"], entry["partOfSpeech"]), FONT_UI, POS_PT, color=MUTED)
    blank(cell, 10)
    footer = cell.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    write(footer, f"第 {entry['lesson']} 課", FONT_UI, LESSON_PT, color=MUTED)


def configure(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(PAGE_W_MM)
    section.page_height = Mm(PAGE_H_MM)
    for attribute in ("left_margin", "right_margin", "header_distance", "footer_distance", "gutter"):
        setattr(section, attribute, Mm(0))
    section.top_margin = Mm(MARGIN_V_MM)
    section.bottom_margin = Mm(MARGIN_V_MM)
    style = document.styles["Normal"]
    style.font.name = FONT_UI
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0



def add_cover(document: Document, total: int, sheets: int, with_picture: int) -> None:
    """One sheet of instructions, then a blank back so the duplex pairing holds."""

    lines = [
        ("聖經希伯來文單字卡", 30, INK, 14),
        (f"{total} 張・{sheets} 組雙面・每頁 8 張", 13, MUTED, 22),
        ("列印：A4 橫式，雙面列印選「沿長邊翻頁」，縮放設為 100%（不要選「符合頁面大小」）。", 11, INK, 6),
        (f"裁切：不印裁切線。自紙張上緣量，橫向切在 {MARGIN_V_MM:.0f}、{MARGIN_V_MM + CARD_H_MM:.0f}、{MARGIN_V_MM + CARD_H_MM * 2:.0f} mm；"
         f"自左緣量，縱向切在 {CARD_W_MM:.2f}、{CARD_W_MM * 2:.1f}、{CARD_W_MM * 3:.2f} mm。成品每張 {CARD_W_MM:.2f}×{CARD_H_MM:.0f} mm。", 11, INK, 6),
        ("正面：附點原文與課次。背面：繁體中文詞義、詞性與課次。", 11, INK, 6),
        (f"插圖：{with_picture} 張有圖，其餘留白。多義詞取最常見的義項；找不到誠實對應的圖就不放，"
         "不以近似圖充數。", 11, INK, 22),
        ("圖片來源：OpenMoji 17.0.0（openmoji.org），CC BY-SA 4.0。", 9.5, MUTED, 4),
        ("詞表：Pratico–Van Pelt《Basics of Biblical Hebrew》2/e 詞序，其後接語料頻率延伸。", 9.5, MUTED, 4),
        ("原文與出現次數：Westminster Leningrad Codex（OSHB）。", 9.5, MUTED, 4),
    ]
    for text, size, color, after in lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.space_before = Pt(0)
        write(paragraph, text, FONT_UI, size, color=color, bold=size >= 30)
    page_break(document)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write(note, "（此頁留白，供雙面列印對齊）", FONT_UI, 9.5, color=MUTED)


def page_break(document: Document) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    spacer.paragraph_format.line_spacing = Pt(1)
    run = spacer.add_run()
    run.font.size = Pt(1)
    run.add_break(WD_BREAK.PAGE)


def build(entries: list[dict], glosses: dict, images: dict, limit: int) -> Path:
    document = Document()
    configure(document)
    if limit:
        entries = entries[:limit]
    sheets = -(-len(entries) // (COLS * ROWS))
    with_picture = sum(1 for entry in entries if f"{entry['strong']}|{entry['pointed']}" in images)
    add_cover(document, len(entries), sheets, with_picture)

    first_sheet = False
    for start in range(0, len(entries), COLS * ROWS):
        page = entries[start : start + COLS * ROWS]
        for side in ("front", "back"):
            if not first_sheet:
                # A full-height paragraph here would push the second row of cards
                # onto a page of its own, so the break carries no visible line.
                page_break(document)
            first_sheet = False
            table = new_grid(document)
            for index in range(COLS * ROWS):
                row, column = divmod(index, COLS)
                if index >= len(page):
                    continue
                entry = page[index]
                # Duplex alignment: the back sheet runs right to left.
                target = table.cell(row, column if side == "front" else COLS - 1 - column)
                if side == "front":
                    fill_front(target, entry)
                else:
                    key = f"{entry['strong']}|{entry['pointed']}"
                    record = images.get(key)
                    picture = IMAGE_DIR / record["file"] if record else None
                    fill_back(target, entry, glosses[(entry["strong"], entry["pointed"])], picture)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


def main() -> None:
    parser = argparse.ArgumentParser(description="產生希伯來文 1000 詞單字卡（A4 橫式、每頁 8 張、雙面）")
    parser.add_argument("--limit", type=int, default=0, help="只做前 N 張（試印用）")
    args = parser.parse_args()

    entries = sorted(json.loads(VOCAB.read_text(encoding="utf-8")), key=lambda item: item["ordinal"])
    glosses = {
        (item["strong"], item["pointed"]): item["glossZh"]
        for item in json.loads(GLOSSES.read_text(encoding="utf-8"))["items"]
    }
    images = json.loads(IMAGES.read_text(encoding="utf-8"))["images"]

    missing = [entry["pointed"] for entry in entries if (entry["strong"], entry["pointed"]) not in glosses]
    if missing:
        raise SystemExit(f"缺繁中詞義 {len(missing)} 筆，例如 {missing[:5]}")

    path = build(entries, glosses, images, args.limit)
    total = args.limit or len(entries)
    sheets = -(-total // (COLS * ROWS))
    with_picture = sum(
        1 for entry in entries[: total] if f"{entry['strong']}|{entry['pointed']}" in images
    )
    print(f"  卡片 {total} 張，正反共 {sheets * 2} 頁（每頁 8 張，{CARD_W_MM:.2f}×{CARD_H_MM:.1f} mm）")
    print(f"  有圖 {with_picture}，留白 {total - with_picture}")
    print(path)


if __name__ == "__main__":
    main()
