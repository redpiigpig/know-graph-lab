#!/usr/bin/env python3
"""Typeset the Koine Greek reader as two JIS B5 print masters, one per volume.

Everything comes from ``greek-reader-two-volumes.json`` and ``interlinear.json``;
this script only sets type.  It shares the Hebrew reader's low-level DOCX
machinery — page geometry, styles, table helpers, run fonts — so both volumes of
the series look like one series, and differs only where the language does:
Greek runs left to right, and its face is Palatino Linotype, which carries the
full polytonic repertoire and installs as a conventional TrueType file that
LibreOffice resolves without substituting.

Each lesson prints its vocabulary table, its two memory units and its whole
reading with a Traditional-Chinese gloss under every Greek word.  The five
reference tables print at the back of **both** volumes: they are a cross-index
of the whole work, and each volume has to be usable on its own.  The liturgy
belongs to 下冊 alone, in celebration order, each utterance labelled with who
says it.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))

from proper_name_categories import PRINT_ORDER  # noqa: E402
from build_hebrew_full_reader import (  # noqa: E402  - shared typesetting machinery
    ACCENT,
    ACCENT_DARK,
    add_contents,
    cover_colors,
    GOLD,
    CAPTION_PT,
    FONT_UI,
    FONT_ZH,
    H1_SIZE_PT,
    H2_SIZE_PT,
    INK,
    LABEL_PT,
    MUTED,
    PALE,
    RULE,
    TABLE_SIZE_PT,
    TITLE_SIZE_PT,
    TRANSLATION_PT,
    USABLE_WIDTH_MM,
    add_body,
    add_label,
    add_mixed_script_text,
    configure,
    gloss_width_mm,
    page_break,
    paragraph_rule,
    prevent_row_split,
    set_borders,
    set_cell_margins,
    set_keep,
    set_repeat_header,
    set_run_font,
    set_table_geometry,
    shade,
)


ROOT = Path(__file__).resolve().parents[1]
CACHE = ROOT / "output" / "source-cache" / "original-readers" / "greek-full"
MASTER_PATH = CACHE / "greek-reader-two-volumes.json"
INTERLINEAR_PATH = CACHE / "interlinear.json"
OUTPUT_DIR = ROOT / "output" / "original-readers"
OUTPUT_STEM = "greek-original-reader-vol"

FONT_GREEK = "Palatino Linotype"
# Chinese characters and CJK punctuation, which the Greek face cannot set.
CJK_RE = re.compile(r"([\u3000-\u303F\u3400-\u4DBF\u4E00-\u9FFF\uFF00-\uFFEF]+)")
GREEK_METRICS = Path(r"C:\Windows\Fonts\pala.ttf")

INTERLINEAR_GREEK_PT = 13.5
INTERLINEAR_GLOSS_PT = 9.4
INTERLINEAR_GUTTER_MM = 3.2
INTERLINEAR_LINE_GAP_PT = 3.5
MEMORY_GREEK_PT = 14
SENSE_PT = 10.4

_greek_metrics = None


def greek_width_mm(text: str, size_pt: float) -> float:
    """Measure with the real print face so a packed row can never overrun.

    Polytonic marks are precomposed in Palatino Linotype, so the sum of
    advances is the true set width.
    """
    global _greek_metrics
    if _greek_metrics is None:
        from PIL import ImageFont

        _greek_metrics = ImageFont.truetype(str(GREEK_METRICS), 1000)
    return _greek_metrics.getlength(text) / 1000 * size_pt / 72 * 25.4


def pack(
    tokens: list[dict],
    available_mm: float,
    *,
    lead_mm: float = 0.0,
    measure=None,
    script_pt: float = INTERLINEAR_GREEK_PT,
) -> list[list[dict]]:
    """Greedy left-to-right packing of word/gloss pairs into full-width rows.

    ``measure`` is how wide a word of the *source* script is; it defaults to
    Greek because that is the book this was written for, and the Latin reader
    passes its own so a Noto Serif word is not measured with Palatino metrics.
    """
    measure = measure or greek_width_mm
    lines: list[list[dict]] = []
    current: list[dict] = []
    used = lead_mm
    for token in tokens:
        surface = token["word"] + token.get("trailing", "")
        width = max(
            measure(surface, script_pt),
            gloss_width_mm(token.get("glossZh", ""), INTERLINEAR_GLOSS_PT),
        ) + INTERLINEAR_GUTTER_MM
        if current and used + width > available_mm:
            lines.append(current)
            current, used = [], 0.0
        current.append({**token, "widthMm": min(width, available_mm)})
        used += width
    if current:
        lines.append(current)
    return lines


def add_interlinear(
    document: Document,
    tokens: list[dict],
    *,
    lead: str = "",
    sense: str = "",
    greek_pt: float = INTERLINEAR_GREEK_PT,
    available_mm: float = USABLE_WIDTH_MM,
    measure=None,
    render=None,
) -> None:
    """One unit as stacked word blocks, closed by the whole-sentence meaning.

    ``measure``/``render`` default to Greek. The Latin reader passes its own
    pair; everything else about the layout — the packing, the gutters, the
    gloss size, the keep-with-next — is deliberately the same in both books.
    """
    if not tokens:
        return
    render = render or (lambda paragraph, text, size: add_greek_run(paragraph, text, size))
    lead_mm = 8.0 if lead else 0.0
    lines = pack(tokens, available_mm, lead_mm=lead_mm, measure=measure, script_pt=greek_pt)
    for line_index, line in enumerate(lines):
        cells_mm = [token["widthMm"] for token in line]
        if line_index == 0 and lead:
            cells_mm.insert(0, lead_mm)
        # Absorb the leftover into a trailing filler column rather than
        # stretching the word blocks, so a short final row still starts at the
        # left margin instead of floating in the middle of the measure.
        slack = available_mm - sum(cells_mm)
        if slack > 1.0:
            cells_mm = [*cells_mm, slack]
        elif slack > 0:
            cells_mm[-1] += slack
        table = document.add_table(rows=1, cols=len(cells_mm))
        set_table_geometry(table, cells_mm)
        set_borders(table, outside=False, inside=False)
        prevent_row_split(table.rows[0])
        for cell_index, cell in enumerate(table.rows[0].cells):
            set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            top = cell.paragraphs[0]
            top.alignment = WD_ALIGN_PARAGRAPH.CENTER
            top.paragraph_format.space_after = Pt(0)
            top.paragraph_format.space_before = Pt(INTERLINEAR_LINE_GAP_PT if line_index else 0)
            top.paragraph_format.line_spacing = 1.16
            bottom = cell.add_paragraph()
            bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bottom.paragraph_format.space_before = Pt(0)
            bottom.paragraph_format.space_after = Pt(0)
            bottom.paragraph_format.line_spacing = 1.0
            if line_index == 0 and lead and cell_index == 0:
                set_run_font(top.add_run(lead), FONT_UI, LABEL_PT, bold=True, color=ACCENT)
                continue
            token_index = cell_index - 1 if (line_index == 0 and lead) else cell_index
            if token_index >= len(line):
                continue
            token = line[token_index]
            render(top, token["word"] + token.get("trailing", ""), greek_pt)
            set_run_font(
                bottom.add_run(token.get("glossZh", "")),
                FONT_ZH,
                INTERLINEAR_GLOSS_PT,
                color=MUTED,
            )
            if line_index < len(lines) - 1:
                set_keep(bottom, next_paragraph=True)
    if sense:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(9)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.left_indent = Mm(6)
        p.paragraph_format.first_line_indent = Mm(-6)
        set_run_font(p.add_run("整句　"), FONT_UI, LABEL_PT, bold=True, color=ACCENT)
        add_mixed_script_text(p, sense, FONT_ZH, SENSE_PT, color=INK)
        set_keep(p, together=True)


def add_greek_run(paragraph, text: str, size: float, *, color=None) -> None:
    """Set Greek in the Greek face, and any CJK-range mark in the Chinese one.

    The Wikisource canons print editorial supplements in CJK angle brackets —
    「〈πρὸ〉」 — and Palatino Linotype has no glyph for those, so LibreOffice
    quietly borrowed an unembedded NotoSansJP for two characters and the PDF's
    font-embedding gate failed on a book that otherwise had none.
    """
    for chunk in filter(None, CJK_RE.split(text)):
        font = FONT_ZH if CJK_RE.fullmatch(chunk) else FONT_GREEK
        set_run_font(paragraph.add_run(chunk), font, size, color=color if color is not None else INK)


def add_plain_greek(document: Document, text: str, size: float = INTERLINEAR_GREEK_PT) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    add_greek_run(p, text, size)


def add_vocabulary(document: Document, lesson: dict) -> None:
    document.add_heading(
        f"生詞　{lesson['vocabularyCount']} 個　{lesson['vocabularySource']}", level=2)
    rows = lesson["vocabulary"]
    table = document.add_table(rows=1, cols=4)
    widths = [8.0, 46.0, 24.0, USABLE_WIDTH_MM - 78.0]
    set_table_geometry(table, widths)
    set_borders(table, color=RULE)
    header = table.rows[0]
    set_repeat_header(header)
    for cell, title in zip(header.cells, ["#", "詞條", "詞類", "繁體中文詞義"]):
        shade(cell, PALE)
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(title), FONT_UI, LABEL_PT, bold=True, color=ACCENT)
    for entry in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.2
            if index == 0:
                set_run_font(paragraph.add_run(str(entry["ordinal"])), FONT_UI, CAPTION_PT, color=MUTED)
            elif index == 1:
                add_greek_run(paragraph, entry["printedEntry"], TABLE_SIZE_PT + 1.6)
                if entry.get("isProperName"):
                    set_run_font(paragraph.add_run("　專名"), FONT_UI, LABEL_PT, color=ACCENT)
            elif index == 2:
                set_run_font(paragraph.add_run(entry.get("pos") or "—"), FONT_ZH, TABLE_SIZE_PT, color=MUTED)
            else:
                add_mixed_script_text(paragraph, entry["glossZh"] or "—", FONT_ZH, TABLE_SIZE_PT, color=INK)


CORPUS_LABELS = {
    "new-testament": "新約",
    "septuagint": "七十士譯本",
    "deuterocanonical": "次經",
    "pseudepigrapha": "偽經",
}


def memory_source_label(unit: dict) -> str:
    """Where a memory unit came from, in Chinese.

    上冊's units carry a corpus code; 下冊's carry the reading they were cut
    from.  Printing the raw code put "new-testament" on the page.
    """
    corpus = unit.get("corpus")
    if corpus:
        return CORPUS_LABELS.get(corpus, corpus)
    return unit.get("readingTitleZh") or ""


def add_memory(document: Document, lesson: dict) -> None:
    kind = lesson["memoryUnits"][0].get("kind") if lesson["memoryUnits"] else "verse"
    document.add_heading("背誦　兩句" if kind == "sentence" else "背誦　兩節", level=2)
    for verse in lesson["memoryUnits"]:
        caption = document.add_paragraph()
        caption.paragraph_format.space_before = Pt(5)
        caption.paragraph_format.space_after = Pt(2)
        set_run_font(caption.add_run(verse["ref"]), FONT_UI, LABEL_PT, bold=True, color=ACCENT)
        set_run_font(
            caption.add_run(
                f"　{memory_source_label(verse)}　命中本課生詞 {verse['matchCount']}"
            ),
            FONT_UI, LABEL_PT, color=MUTED,
        )
        set_keep(caption, next_paragraph=True)
        tokens = verse.get("tokens") or []
        if tokens:
            add_interlinear(document, tokens, sense=verse.get("translationZh", ""), greek_pt=MEMORY_GREEK_PT)
        else:
            add_plain_greek(document, verse["text"], MEMORY_GREEK_PT)
            if verse.get("translationZh"):
                add_body(document, verse["translationZh"], size=TRANSLATION_PT, color=INK)


def add_reading(document: Document, lesson: dict, interlinear: dict) -> None:
    reading = lesson["reading"]
    is_scripture = reading["kind"] == "scripture_chapter"
    label = "讀文　" + (reading.get("corpusLabel") or reading.get("categoryLabel") or "")
    if reading.get("completeness") == "excerpt":
        label += f"　節錄・{reading.get('extent', '')}"
    # 每一課的讀物另起一頁：詞表與背誦是準備，讀物是這一課的正事，
    # 讓它從頁首開始，翻到就是整篇。
    page_break(document)
    document.add_heading(label, level=2)
    add_body(document, reading["source"], size=CAPTION_PT, color=MUTED)
    if reading.get("numberingNote"):
        add_body(document, reading["numberingNote"], size=CAPTION_PT, color=MUTED)

    segments = (reading.get("verses") if is_scripture else reading.get("segments")) or []
    for segment in segments:
        unit_id = (
            f"scripture:{segment['ref']}"
            if is_scripture
            else f"patristic:{reading['ordinal']}:{segment['ref']}"
        )
        record = interlinear.get(unit_id) or {}
        tokens = record.get("tokens") or []
        sense = segment.get("translationZh") or record.get("translationZh") or ""
        lead = str(segment.get("verse") or segment.get("ref", ""))
        if tokens:
            add_interlinear(document, tokens, lead=lead, sense=sense)
        else:
            add_plain_greek(document, segment.get("displayText", ""))
            if sense:
                add_body(document, sense, size=TRANSLATION_PT, color=INK)
    for absent in reading.get("absentVerses") or []:
        add_body(document, f"{absent['ref']}：{absent['note']}", size=CAPTION_PT, color=MUTED)


def add_lesson(document: Document, lesson: dict, interlinear: dict) -> None:
    # 版式比照希伯來那本：眉標（另起一頁）→ 課次 → Heading 1 課題 → 金線。
    kind = "scripture chapter" if lesson["reading"]["kind"] == "scripture_chapter" else "church reading"
    add_label(document, f"Lesson {lesson['lesson']:02d}  ·  {kind}", page_break_before=True)
    number = document.add_paragraph()
    number.paragraph_format.space_after = Pt(1)
    set_run_font(number.add_run(f"第 {lesson['lesson']:02d} 課"), FONT_UI, 11, bold=True, color=ACCENT)
    heading = document.add_heading(lesson["reading"]["titleZh"], level=1)
    paragraph_rule(heading, color=GOLD, size="14")
    greek_title = document.add_paragraph()
    greek_title.paragraph_format.space_after = Pt(6)
    add_greek_run(greek_title, lesson["reading"]["titleGrc"], TRANSLATION_PT, color=MUTED)
    add_vocabulary(document, lesson)
    add_memory(document, lesson)
    add_reading(document, lesson, interlinear)


def add_liturgy(document: Document, liturgy: dict, interlinear: dict) -> None:
    page_break(document)
    add_label(document, "Appendix  ·  divine liturgy")
    heading = document.add_heading(liturgy["title"], level=1)
    paragraph_rule(heading, color=GOLD, size="14")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    add_greek_run(subtitle, liturgy["titleGrc"], TRANSLATION_PT, color=MUTED)
    add_body(document, liturgy["placement"], size=CAPTION_PT, color=MUTED)
    add_body(document, liturgy["roleDerivationNote"], size=CAPTION_PT, color=MUTED)

    current = ""
    for step in liturgy["steps"]:
        if step["section"] != current:
            current = step["section"]
            document.add_heading(step["sectionLabel"], level=3)
        record = interlinear.get(f"liturgy:{step['ordinal']}") or {}
        caption = document.add_paragraph()
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(1)
        set_run_font(caption.add_run(step["roleLabel"]), FONT_UI, LABEL_PT, bold=True, color=ACCENT)
        if step.get("repeatCount"):
            set_run_font(caption.add_run(f"　重複 {step['repeatCount']} 次"), FONT_UI, LABEL_PT, color=MUTED)
        set_keep(caption, next_paragraph=True)
        tokens = record.get("tokens") or []
        if tokens:
            add_interlinear(document, tokens, sense=record.get("translationZh", ""))
        else:
            add_plain_greek(document, step["displayText"])




def add_latin_and_cjk(paragraph, text: str, size: float, *, color=MUTED) -> None:
    """Set the Latin in a Latin face and the Chinese in a Chinese one.

    One run in "Noto Serif" carrying both leaves LibreOffice to find the Chinese
    somewhere, and it picks a font it does not then embed — the title page's
    「（上冊新約部分）」 came out in an unembedded NotoSansJP-Thin, which fails the
    PDF gate even though it looks fine on screen.
    """
    for chunk in filter(None, CJK_RE.split(text)):
        font = FONT_ZH if CJK_RE.fullmatch(chunk) else "Noto Serif"
        set_run_font(paragraph.add_run(chunk), font, size, color=color)


# 封面上的希臘文題辭，一冊一句：上冊出自福音書，下冊出自尼西亞信經。
COVER_GREEK = {
    1: "Ἡ ΚΑΙΝΗ ΔΙΑΘΗΚΗ",
    2: "ΤΩΝ ΠΑΤΕΡΩΝ ΤΑ ΚΕΙΜΕΝΑ",
}


def add_cover(document: Document, master: dict, volume: dict) -> None:
    """深色橫幅封面，版式與希伯來、拉丁那兩本相同。

    三本並排時封面要看得出是同一套書：同一條深色橫幅、同一行金色眉標、
    同一條金線與規格行，換的只有書名與原文題辭。
    """
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [USABLE_WIDTH_MM])
    set_borders(table, outside=False, inside=False)
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=500, bottom=500, start=350, end=350)
    palette = cover_colors("grc")
    shade(cell, palette["banner"])

    eyebrow = cell.paragraphs[0]
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(eyebrow.add_run("ORIGINAL-LANGUAGE READER"), FONT_UI, 8,
                 color=palette["rule"], bold=True)
    name = cell.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_mixed_script_text(name, master["title"], FONT_ZH, 25, bold=True, color="FFF8ED")
    motto = cell.add_paragraph()
    motto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_greek_run(motto, COVER_GREEK[volume["volume"]], 18, color="FFF8ED")

    document.add_paragraph().paragraph_format.space_after = Pt(26)
    volume_line = document.add_paragraph()
    volume_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_mixed_script_text(volume_line, volume["title"], FONT_ZH, 12, bold=True, color=INK)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_mixed_script_text(subtitle, volume["subtitle"], FONT_ZH, 10.5, color=ACCENT)

    document.add_paragraph().paragraph_format.space_after = Pt(26)
    spec = document.add_paragraph()
    spec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_rule(spec, color=cover_colors("grc")["rule"], size="24")
    set_run_font(spec.add_run("JIS B5  182 × 257 mm  ·  私人研讀"), FONT_UI, 8.5, color=MUTED)
    # `master["textbook"]` 只講得到上冊（它寫的就是「（上冊新約部分）」），印在下冊
    # 封面上是錯的。下冊的詞表來自語料頻率，不出自哪一本教科書，就照實那樣寫。
    textbook = document.add_paragraph()
    textbook.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line = (master["textbook"] if volume["volume"] == 1
            else "詞表：教父文獻與希臘教會文獻語料詞頻，與上冊不重複")
    add_latin_and_cjk(textbook, line, CAPTION_PT)


def add_front_matter(document: Document, master: dict, volume: dict) -> None:
    add_cover(document, master, volume)
    page_break(document)
    document.add_heading("體例與來源", level=1)
    for key, value in master["textPolicy"].items():
        add_body(document, f"{key}：{value}", size=CAPTION_PT, color=INK)
    counts = volume["counts"]
    unit_word = "句背誦" if volume["memoryUnitKind"] == "sentence" else "節背誦"
    add_body(
        document,
        f"本冊五十課・{counts['vocabulary']} 詞・{counts['memoryUnits']} {unit_word}・"
        f"{counts['readings']} 篇讀文；全書兩冊合計 {master['counts']['vocabulary']} 詞、"
        f"{master['counts']['memoryUnits']} 則背誦、連續正文 {master['counts']['totalRunningWords']} 詞。",
        size=CAPTION_PT,
        color=MUTED,
    )
    for half, label in volume["corpusByHalf"].items():
        add_body(document, f"第 {half} 課：{label}", size=CAPTION_PT, color=MUTED)
    add_body(document, f"發布狀態：{master['releaseStatus']}", size=CAPTION_PT, color=MUTED)
    add_body(document, f"音訊：{master['audio']['status']}　{master['audio']['policy']}", size=CAPTION_PT, color=MUTED)
    page_break(document)
    add_contents(
        document,
        [
            (
                f"{lesson['lesson']:02d}",
                lesson["reading"]["titleZh"],
                "完整章" if lesson["reading"]["kind"] == "scripture_chapter" else "教父讀文",
            )
            for lesson in volume["lessons"]
        ],
        title=f"{volume['title']}目錄",
        accent=cover_colors("grc")["accent"],
    )


def appendix_groups(table: dict) -> list[tuple[str, list[dict]]]:
    """把一張附錄表切成印得出來的小節。

    專名表用 `category`（九類，見 scripts/proper_name_categories.py），其餘四張表
    用資料裡本來就有的 `group`。兩者都照 PRINT_ORDER／首次出現排序，沒有分組欄位
    的表就整張當一節。
    """
    field = "category" if any(e.get("category") for e in table["entries"]) else "group"
    buckets: dict[str, list[dict]] = {}
    for entry in table["entries"]:
        buckets.setdefault((entry.get(field) or "").strip(), []).append(entry)
    if len(buckets) <= 1:
        return [("", table["entries"])]
    known = [n for n in PRINT_ORDER if n in buckets]
    rest = [n for n in buckets if n not in known]
    return [(name, buckets[name]) for name in known + rest]


def add_appendix_entry(document: Document, entry: dict) -> None:
    row = document.add_paragraph()
    row.paragraph_format.space_after = Pt(0)
    row.paragraph_format.line_spacing = 1.25
    add_greek_run(row, entry.get("headword") or entry["lemma"], TABLE_SIZE_PT + 1.2)
    chinese = (entry.get("zh") or "").strip()
    if chinese:
        add_mixed_script_text(row, f"　{chinese}", FONT_ZH, TABLE_SIZE_PT, color=INK)
    else:
        # An empty cell is the honest state for a name no register
        # covers; it is marked rather than filled with a guess.
        set_run_font(row.add_run("　（中文待定）"), FONT_UI, CAPTION_PT, color=MUTED)
    if entry.get("frequency"):
        set_run_font(row.add_run(f"　{entry['frequency']}"), FONT_UI, CAPTION_PT, color=MUTED)


def add_appendix_tables(document: Document, master: dict) -> None:
    """The five reference tables, printed at the back of both volumes.

    They index the whole work rather than one volume, and a volume being read on
    its own still needs the numerals and the kinship terms, so they are repeated
    rather than split between the two books.

    專名表按九類分節印：查「彼得是誰」時翻到〈使徒與門徒〉一節就找得到，而不是在
    四百條按字母排的名字裡一條條看過去。
    """
    for table in master["appendices"]:
        page_break(document)
        add_label(document, "Appendix  ·  reference table")
        heading = document.add_heading(table["title"], level=1)
        paragraph_rule(heading, color=GOLD, size="14")
        if table.get("note"):
            add_body(document, table["note"], size=CAPTION_PT, color=MUTED)
        for group_name, entries in appendix_groups(table):
            if group_name:
                document.add_heading(f"{group_name}　{len(entries)} 條", level=2)
            for entry in entries:
                add_appendix_entry(document, entry)


def retitle(document: Document, master: dict, volume: dict) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    for run in list(header.runs):
        run._element.getparent().remove(run._element)
    set_run_font(
        header.add_run(f"{master['title']}  ·  {volume['title']}"), FONT_UI, 7.5, color=MUTED
    )
    properties = document.core_properties
    properties.title = f"{master['title']}：{volume['title']}"
    properties.subject = volume["subtitle"]
    properties.language = "grc"


def build(volume_number: int) -> Path:
    master = json.loads(MASTER_PATH.read_text(encoding="utf-8"))
    interlinear = json.loads(INTERLINEAR_PATH.read_text(encoding="utf-8"))["units"]
    volume = next((item for item in master["volumes"] if item["volume"] == volume_number), None)
    if volume is None:
        raise SystemExit(f"主檔沒有第 {volume_number} 冊")

    document = Document()
    configure(document)
    # configure() is the Hebrew volume's, so it stamps that volume's running
    # header and document title.  Retitle both, or every page of the Greek
    # reader says it is the Hebrew one.
    retitle(document, master, volume)
    add_front_matter(document, master, volume)
    for lesson in volume["lessons"]:
        add_lesson(document, lesson, interlinear)
    if any(item["kind"] == "divine-liturgy" for item in volume["appendices"]):
        liturgy = json.loads((CACHE / "liturgy-chrysostom.json").read_text(encoding="utf-8"))
        add_liturgy(document, liturgy, interlinear)
    add_appendix_tables(document, master)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / f"{OUTPUT_STEM}{volume_number}.docx"
    document.save(path)
    return path


def main() -> None:
    parser = argparse.ArgumentParser(description="排版希臘文讀本 B5 DOCX（兩冊）")
    parser.add_argument("--volume", type=int, choices=(1, 2), help="只排某一冊")
    args = parser.parse_args()
    for number in ([args.volume] if args.volume else [1, 2]):
        path = build(number)
        print(f"已寫出 {path}（{path.stat().st_size / 1_048_576:.1f} MB）")


if __name__ == "__main__":
    main()
